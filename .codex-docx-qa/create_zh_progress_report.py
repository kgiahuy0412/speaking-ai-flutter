from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(
    r"C:\Users\Windows\Documents\ai-speaking-flutter-app\deliverables"
    r"\AI_Speaking_Project_Status_Report_ZH_CN_2026-07-30.docx"
)

FONT = "SimSun"
FONT_CN = "宋体"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6470"
LIGHT_GRAY = "F2F4F7"
BORDER = "CBD2D9"
WHITE = "FFFFFF"
GREEN = "2F6B3C"
AMBER = "8A5A00"


overview_rows = [
    ("Flutter 版本", "1.0.3+5"),
    ("支持平台", "Android、Flutter Web/PWA、iOS 模拟器框架"),
    ("后端", "部署于 Railway 的 Next.js 服务，区域为新加坡"),
    ("测试", "158 项测试全部通过"),
    ("代码分析", "flutter analyze：无错误"),
    ("当前 Release APK", "约 78.5 MB"),
    ("代码规模", "92 个 Dart 文件、31 个测试文件、12 张 Golden Test 基准图"),
]


progress_rows = [
    ("架构", "将 Flutter 与 Next.js 拆分为两个独立项目，Flutter 仅调用后端 API。", "完成"),
    ("API 安全", "OpenAI API 密钥仅保存在后端，不会打包进 APK。", "完成"),
    ("后端配置", "Flutter 在运行或构建时读取 BACKEND_BASE_URL。", "完成"),
    ("后端部署", "Next.js 已在 Railway 上通过生产域名运行。", "完成"),
    ("服务器区域优化", "Railway 区域已从 US West 迁移至新加坡。", "已应用"),
    ("对话主界面", "单一主界面包含麦克风状态、越南语句子、英语句子、结果评价与录音按钮。", "完成"),
    ("Android Streaming", "使用 Android SpeechRecognizer，支持临时结果与最终结果。", "完成"),
    ("越南语识别", "使用 vi-VN，最多返回 3 个候选结果，并包含置信度。", "完成"),
    ("截断句修复", "当 Android 返回的最终句子不完整时，优先采用更长的转写结果。", "完成"),
    ("麦克风录音", "PCM16、单声道、24 kHz；每轮录音上限约 12 秒。", "完成"),
    ("语音活动检测", "已加入 Adaptive VAD，可自学习背景噪声水平，并区分人声与风扇等稳定噪声。", "已实现，尚未提交"),
    ("无语音处理", "在未确认检测到人声时，不会向后端上传音频。", "完成"),
    ("无语音等待时间", "当前代码实际等待约 3 秒，并非 2 秒。", "如需调整"),
    ("OpenAI Realtime", "已实现 WebSocket 连接、流式转写与短期 client secret。", "已构建"),
    ("Realtime 状态", "目前不是默认模式，并且在用户选项中处于隐藏状态。", "具备但未开放"),
    ("Realtime 缓冲", "等待 socket 建立期间持续录音，连接后补发开头的音频。", "完成"),
    ("Batch Chunks", "将音频分成多个小片段上传，并由后端合并识别结果。", "完成"),
    ("降级机制", "Realtime 失败时切换到 Batch Chunks；Batch 失败时上传完整 WAV。", "完成"),
    ("Web 网络优化", "延迟创建上传会话，避免不必要的请求。", "完成"),
    ("音频播放", "使用 just_audio 播放英语句子和已生成音频。", "完成"),
    ("音频缓存", "设备端最多保留约 128 个文件，并支持预加载与预热。", "完成"),
    ("播放优化", "预加载最长等待 500 ms，并跟踪加载耗时与实际开始播放时间。", "完成"),
    ("Safari 音频", "已处理用户手势要求，避免 Safari 阻止自动播放。", "完成"),
    ("HFP 耳机", "支持扫描、选择设备、开启 SCO 音频路由并进行 4 秒麦克风测试。", "已构建"),
    ("Web 端 HFP", "允许选择浏览器可识别的蓝牙麦克风。", "已构建"),
    ("INNOTRIK BLE 设备", "支持扫描、连接、MTU 247、自动重连、读取数据包并提取 Opus 音频。", "已构建"),
    ("Offline Intent", "在高置信度时于设备端识别部分意图，以减少 OpenAI 调用。", "已构建"),
    ("历史记录", "支持查看、搜索、按日期筛选、回放、删除单条记录或清空全部记录。", "完成"),
    ("正确/错误反馈", "用户可评价结果，后端会同步更新历史记录。", "完成"),
    ("设置", "可选择麦克风来源、ASR 模式、VAD、HFP/BLE 与界面语言。", "完成"),
    ("多语言界面", "支持越南语和简体中文。", "完成"),
    ("听力学习内容", "包含 5 个年龄组、50 个主题、101 节课程和 11 首歌曲。", "已有"),
    ("句子内容总量", "课程和歌曲合计约 713 个句子。", "已有"),
    ("引导音频", "应用内已打包 212 个引导/歌曲 MP3 文件。", "已有"),
    ("在线课程音频", "713 个句子中有 574 个已配备英语和越南语音频。", "尚未达到 100%"),
    ("课程介绍音频", "101 节课程中有 96 节已配备介绍音频。", "接近完成"),
    ("逐句练习", "支持听示范、录音、回听、重录、继续和返回上一句。", "完成"),
    ("录音保存", "每个句子最多保留 3 条最近且质量较好的录音。", "完成"),
    ("儿童开口提醒", "每隔 5 秒提醒一次，共提醒两次，之后允许跳过。", "完成"),
    ("积极反馈", "包含表扬音频、烟花动画与完成状态。", "完成"),
    ("卡拉 OK", "歌词按时间滚动，歌曲播放后 3 秒自动开始。", "已构建"),
    ("进度保存", "在 Native/Web 端保存已学课程、已跳过句子与历史记录。", "完成"),
    ("视觉设计", "采用机器人吉祥物、靛蓝/薰衣草紫配色、绿色/红色状态及儿童友好界面。", "完成"),
    ("学习场景", "已加入田野背景、火车、主题图片与课程介绍页图片。", "已实现，尚未提交"),
    ("无障碍支持", "已测试 200% 字体、较小屏幕与减少动态效果模式。", "已测试"),
    ("Flutter Web/PWA", "已具备 Manifest、图标、iPhone 主屏幕安装及浏览器数据持久化。", "完成"),
    ("Web 自动更新", "检查 version.json，显示更新提示并重新加载新版本。", "完成"),
    ("APK 更新", "检查后端版本，可执行强制更新并打开 APK 下载页面。", "客户端已完成"),
    ("应用图标", "已为 Android 与 Web 生成应用图标。", "完成"),
    ("APK 下载页", "页面指向 Cloudflare R2 上的最新 APK。", "完成"),
    ("Release 构建", "支持通用 Release APK 以及按 ABI 分包构建。", "完成"),
    ("iOS", "已具备 iOS scaffold 与面向模拟器的 Codemagic 工作流。", "尚非 App Store 版本"),
    ("Golden Test", "在多种手机屏幕尺寸下验证界面。", "12 组基准"),
    ("音频测试", "覆盖 VAD、WAV、缓存、音频播放完成、降级机制与转写。", "通过"),
    ("硬件测试", "覆盖 HFP、BLE 数据包、Offline Intent 与重连逻辑。", "自动化层面通过"),
    ("业务测试", "覆盖历史记录、设置、课程、卡拉 OK、进度与更新门禁。", "通过"),
]


remaining_rows = [
    ("在多种真实设备上测试 HFP 与 BLE", "重要"),
    ("为全部 713 个句子补齐音频", "仍有缺失"),
    ("使用正式 keystore 进行生产发布签名", "正式发布前必须完成"),
    ("确认 Flutter Web/PWA 已公开部署", "尚未确认"),
    ("加强后端身份验证、限流与日志保留策略", "大规模生产前需要"),
    ("验证 Safari/iPhone、弱网、锁屏与后台运行场景", "需要真机测试"),
    ("决定是否向用户开放 OpenAI Realtime，或继续使用 Android Streaming", "待决定"),
    ("提交 Adaptive VAD、课程界面与新 Golden Test 的本地变更", "仍在本地"),
]


key_files = [
    (
        "conversation_controller.dart",
        r"C:\Users\Windows\Documents\ai-speaking-flutter-app\lib\features\conversation\presentation\conversation_controller.dart",
        "对话流程、录音、识别、降级与播放控制。",
    ),
    (
        "adaptive_voice_activity_detector.dart",
        r"C:\Users\Windows\Documents\ai-speaking-flutter-app\lib\core\audio\adaptive_voice_activity_detector.dart",
        "自适应背景噪声学习与语音活动判定。",
    ),
    (
        "listening_lessons.json",
        r"C:\Users\Windows\Documents\ai-speaking-flutter-app\assets\data\listening_lessons.json",
        "听力课程、主题与句子内容数据。",
    ),
    (
        "pubspec.yaml",
        r"C:\Users\Windows\Documents\ai-speaking-flutter-app\pubspec.yaml",
        "应用版本、依赖项与资源配置。",
    ),
]


def set_run_font(run, size=None, bold=None, color=None, italic=None, latin_font=FONT):
    run.font.name = latin_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_fonts.set(qn("w:cs"), latin_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size, color, bold=False):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    r_pr.append(r_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    r_pr.append(color)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    r_pr.append(size)
    field_run.append(r_pr)
    field_run.append(begin)
    field_run.append(instr)
    field_run.append(separate)
    field_run.append(display)
    field_run.append(end)
    paragraph._p.append(field_run)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def style_table_text(table, header=True, body_size=9.5, status_col=None):
    for row_idx, row in enumerate(table.rows):
        is_header = header and row_idx == 0
        if is_header:
            set_repeat_table_header(row)
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, LIGHT_GRAY if is_header else WHITE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.12
                if is_header or (status_col is not None and col_idx == status_col):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    color = NAVY if is_header else "20242A"
                    if status_col is not None and col_idx == status_col and not is_header:
                        text = paragraph.text
                        color = AMBER if any(
                            token in text
                            for token in (
                                "未",
                                "调整",
                                "需要",
                                "决定",
                                "本地",
                                "基准",
                                "100%",
                                "接近",
                            )
                        ) else GREEN
                    set_run_font(
                        run,
                        size=10 if is_header else body_size,
                        bold=is_header or (status_col is not None and col_idx == status_col),
                        color=color,
                    )


def add_table(doc, headers, rows, widths_dxa, status_col=None, body_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(data):
            cells[idx].text = value
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    style_table_text(table, body_size=body_size, status_col=status_col)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_spacer(doc, points):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run("")
    set_run_font(run, size=1)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 11, "20242A")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        set_style_font(style, size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style.paragraph_format.keep_with_next = True

    if "Report Title" not in styles:
        title_style = styles.add_style("Report Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title_style = styles["Report Title"]
    set_style_font(title_style, 24, NAVY, bold=True)
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(5)
    title_style.paragraph_format.line_spacing = 1.05
    title_style.paragraph_format.keep_with_next = True

    if "Report Subtitle" not in styles:
        subtitle_style = styles.add_style("Report Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle_style = styles["Report Subtitle"]
    set_style_font(subtitle_style, 12.5, MUTED)
    subtitle_style.paragraph_format.space_before = Pt(0)
    subtitle_style.paragraph_format.space_after = Pt(14)
    subtitle_style.paragraph_format.line_spacing = 1.12
    subtitle_style.paragraph_format.keep_with_next = True

    core = doc.core_properties
    core.title = "AI口语学习应用项目进展报告"
    core.subject = "截至2026年7月30日的代码、功能、测试与发布准备情况"
    core.author = "项目团队"
    core.keywords = "Flutter, Android, Web, PWA, 语音识别, 听力学习, 项目进展"

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    header = section.header
    header_p = header.paragraphs[0]
    header_p.paragraph_format.space_after = Pt(0)
    header_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header_p.add_run("AI口语学习应用 · 项目进展报告")
    set_run_font(left, size=8.5, bold=True, color=MUTED)
    right = header_p.add_run("\t2026年7月30日")
    set_run_font(right, size=8.5, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    add_page_field(footer_p)

    add_spacer(doc, 9)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(6)
    kicker.paragraph_format.keep_with_next = True
    run = kicker.add_run("管理层汇报")
    set_run_font(run, size=9.5, bold=True, color=BLUE)

    title = doc.add_paragraph(style="Report Title")
    run = title.add_run("AI口语学习应用")
    set_run_font(run, size=24, bold=True, color=NAVY)
    run.add_break(WD_BREAK.LINE)
    run = title.add_run("项目进展报告")
    set_run_font(run, size=24, bold=True, color=NAVY)

    subtitle = doc.add_paragraph(style="Report Subtitle")
    run = subtitle.add_run("基于截至 2026 年 7 月 30 日的当前代码状态，包含部分仍在本地、尚未提交的变更。")
    set_run_font(run, size=12.5, color=MUTED)

    for label, value in (
        ("报告日期", "2026 年 7 月 30 日"),
        ("当前版本", "Flutter 1.0.3+5"),
        ("覆盖范围", "Android、Flutter Web/PWA、iOS 模拟器框架与 Next.js 后端"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}：")
        set_run_font(label_run, size=10.5, bold=True, color=NAVY)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color="20242A")

    add_spacer(doc, 8)
    summary = doc.add_paragraph()
    summary.paragraph_format.space_before = Pt(0)
    summary.paragraph_format.space_after = Pt(10)
    summary.paragraph_format.line_spacing = 1.18
    summary.paragraph_format.left_indent = Inches(0.14)
    summary.paragraph_format.right_indent = Inches(0.14)
    set_paragraph_shading(summary, LIGHT_GRAY)
    label_run = summary.add_run("执行摘要：")
    set_run_font(label_run, size=10.5, bold=True, color=NAVY)
    text_run = summary.add_run(
        "跨平台架构、核心语音链路、听力学习内容、历史记录与更新机制已经建立。"
        "正式发布前的主要工作集中在真实设备验证、音频覆盖补齐、正式签名、公开部署确认及生产安全强化。"
    )
    set_run_font(text_run, size=10.5, color="20242A")

    add_heading(doc, "1. 总览")
    add_table(doc, ("信息", "当前情况"), overview_rows, (2880, 6480), body_size=10)

    add_heading(doc, "2. 已完成及当前进展")
    intro = doc.add_paragraph("以下内容按当前代码状态汇总；“尚未提交”表示功能已存在于本地工作区，但尚未进入正式版本控制记录。")
    intro.paragraph_format.keep_with_next = True
    for run in intro.runs:
        set_run_font(run, size=10.5, color=MUTED)
    add_table(
        doc,
        ("项目", "已完成内容 / 当前实现", "状态"),
        progress_rows,
        (2160, 5760, 1440),
        status_col=2,
        body_size=9.2,
    )

    add_heading(doc, "3. 待办事项")
    add_table(doc, ("待办工作", "优先级 / 当前状态"), remaining_rows, (6912, 2448), body_size=9.8)

    add_heading(doc, "4. 关键代码与文件")
    key_table = doc.add_table(rows=1, cols=2)
    key_table.style = "Table Grid"
    key_table.rows[0].cells[0].text = "文件"
    key_table.rows[0].cells[1].text = "作用与位置"
    for filename, path, purpose in key_files:
        cells = key_table.add_row().cells
        p0 = cells[0].paragraphs[0]
        p0.clear()
        r0 = p0.add_run(filename)
        set_run_font(r0, size=9.3, bold=True, color=NAVY, latin_font="Consolas")
        p1 = cells[1].paragraphs[0]
        p1.clear()
        purpose_run = p1.add_run(purpose)
        set_run_font(purpose_run, size=9.5, color="20242A")
        path_p = cells[1].add_paragraph()
        path_p.paragraph_format.space_before = Pt(2)
        path_p.paragraph_format.space_after = Pt(0)
        path_run = path_p.add_run(path)
        set_run_font(path_run, size=7.5, color=MUTED, latin_font="Consolas")
    set_table_geometry(key_table, (2880, 6480))
    set_table_borders(key_table)
    style_table_text(key_table, body_size=9.3)
    for row_idx, row in enumerate(key_table.rows):
        if row_idx == 0:
            continue
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if "C:\\" in run.text:
                        set_run_font(run, size=7.5, color=MUTED, latin_font="Consolas")

    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(10)
    closing.paragraph_format.space_after = Pt(0)
    closing.paragraph_format.line_spacing = 1.15
    run = closing.add_run("用途说明：")
    set_run_font(run, size=10.5, bold=True, color=NAVY)
    run = closing.add_run("本报告可作为项目进度汇报材料，也可作为正式发布前的检查清单。")
    set_run_font(run, size=10.5, color="20242A")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(output)
