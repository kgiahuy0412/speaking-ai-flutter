from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def clean(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split())


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    document = Document(args.input)
    lines: list[str] = []
    table_number = 0

    for block in document.iter_inner_content():
        if isinstance(block, Paragraph):
            text = clean(block.text)
            if text:
                style = clean(block.style.name) if block.style else ""
                lines.append(f"\n## P [{style}]\n{text}\n")
            continue
        if not isinstance(block, Table):
            continue
        table_number += 1
        rows: list[list[str]] = []
        width = 0
        for row in block.rows:
            values = [clean(cell.text) for cell in row.cells]
            if not any(values):
                continue
            width = max(width, len(values))
            rows.append(values)
        if not rows:
            continue
        lines.append(f"\n## TABLE {table_number}\n")
        for row in rows:
            padded = row + [""] * (width - len(row))
            lines.append(" | ".join(padded))
        lines.append("")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text("\n".join(lines), encoding="utf-8")
    print(
        f"extracted paragraphs={len(document.paragraphs)} "
        f"tables={len(document.tables)} output={args.output}"
    )


if __name__ == "__main__":
    main()
