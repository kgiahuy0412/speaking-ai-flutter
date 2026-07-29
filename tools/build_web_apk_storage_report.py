from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Phuong_an_luu_tru_dong_bo_Web_APK_V1.docx"

NAVY = "16325C"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6574"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF7EF"
PALE_GOLD = "FFF7E2"
PALE_RED = "FDEEEE"
WHITE = "FFFFFF"
BORDER = "C9D2DF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_rule(paragraph, color=BLUE, size=10):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def add_numbering(document, kind):
    numbering = document.part.numbering_part.element
    existing_abs = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abs, default=-1) + 1
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def add_list_item(document, text, num_id, bold_prefix=None):
    paragraph = document.add_paragraph()
    apply_num(paragraph, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_font(run, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_font(rest)
    else:
        run = paragraph.add_run(text)
        set_font(run)
    return paragraph


def add_body(document, text, bold_lead=None, italic=False):
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_font(lead, bold=True)
        tail = paragraph.add_run(text[len(bold_lead):])
        set_font(tail, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_font(run, italic=italic)
    return paragraph


def add_callout(document, label, text, fill=LIGHT_BLUE):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size=4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, color=NAVY, bold=True)
    body_run = paragraph.add_run(text)
    set_font(body_run, color=INK)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def populate_cell(cell, text, header=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.5):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_font(run, size=size, color=NAVY if header else INK, bold=header)


def add_table(document, headers, rows, widths, center_cols=()):
    table = document.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        populate_cell(cell, header, header=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.CENTER if index in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            populate_cell(row.cells[index], str(value), align=align)
        set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_code_block(document, text):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D7DEE8", size=4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_font(run, name="Consolas", size=9, color="263238")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def setup_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def setup_page(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("AI SPEAKING — KIẾN TRÚC DỮ LIỆU WEB/PWA & APK")
    set_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [7200, 2160], indent=0)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    left_run = left.add_run("Đề xuất V1 • 29/07/2026")
    set_font(left_run, size=9, color=MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    right_run = right.add_run("Trang ")
    set_font(right_run, size=9, color=MUTED)
    add_page_field(right)
    # Quiet footer: remove visible table borders.
    set_table_borders(table, color=WHITE, size=0)


def build_document():
    document = Document()
    setup_styles(document)
    setup_page(document)
    bullet_id = add_numbering(document, "bullet")
    number_id = add_numbering(document, "decimal")

    props = document.core_properties
    props.title = "Phương án lưu trữ, đồng bộ và khôi phục dữ liệu Web/PWA và APK"
    props.subject = "Kiến trúc dữ liệu V1 cho AI Speaking"
    props.author = "Nhóm dự án AI Speaking"
    props.keywords = "PWA, Flutter, APK, IndexedDB, backend, R2, cập nhật APK"

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("PHƯƠNG ÁN LƯU TRỮ, ĐỒNG BỘ\nVÀ KHÔI PHỤC DỮ LIỆU")
    set_font(title_run, size=24, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run("Web/PWA và ứng dụng Flutter APK — Đề xuất kiến trúc V1")
    set_font(subtitle_run, size=13.5, color=MUTED)

    metadata = [
        ("Phạm vi", "Tiến độ học, bản ghi âm luyện nghe, khôi phục Web và cập nhật APK"),
        ("Nguyên tắc", "Không bắt buộc tài khoản; local-first; backend là nguồn sao lưu và đồng bộ"),
        ("Trạng thái", "Đề xuất để thống nhất trước khi triển khai backend và Flutter"),
        ("Ngày", "29/07/2026"),
    ]
    for label, value in metadata:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_font(label_run, bold=True, color=NAVY)
        value_run = p.add_run(value)
        set_font(value_run)
    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    set_paragraph_rule(rule)

    add_callout(
        document,
        "Kết luận đề xuất",
        "V1 nên để Web và APK lưu riêng theo định danh ẩn danh của từng thiết bị. Web chuyển từ blob/localStorage sang IndexedDB và đồng bộ tiến độ lên backend; audio lưu trên R2. QR chỉ là cơ chế khôi phục Web khi toàn bộ dữ liệu trình duyệt bị xóa. Cập nhật APK cài đè bình thường không cần QR.",
        fill=LIGHT_BLUE,
    )

    document.add_heading("1. Mục tiêu và quyết định kiến trúc", level=1)
    add_body(
        document,
        "Tài liệu này thống nhất cách lưu dữ liệu học tập trên Web/PWA và APK trong điều kiện sản phẩm chưa triển khai tài khoản cá nhân. Mục tiêu là tránh mất tiến độ, không làm phụ huynh phải thực hiện nhiều bước, đồng thời giữ khả năng mở rộng sang đồng bộ đa thiết bị trong tương lai.",
    )
    for item in (
        "Web/PWA hoạt động ngay mà không yêu cầu đăng ký hoặc ghép nối APK.",
        "APK tiếp tục hoạt động độc lập; cập nhật chính thức giữ nguyên dữ liệu nếu cài đè đúng chữ ký.",
        "Tiến độ được lưu local trước để giao diện nhanh, sau đó đồng bộ nền lên backend.",
        "File ghi âm không lưu trực tiếp trong PostgreSQL; local dùng IndexedDB/file, cloud dùng R2.",
        "Đồng bộ Web–APK là tính năng riêng của giai đoạn sau, không phải điều kiện của V1.",
    ):
        add_list_item(document, item, bullet_id)

    document.add_heading("2. Hiện trạng được xác nhận", level=1)
    add_table(
        document,
        ["Khu vực", "Cách lưu hiện tại", "Rủi ro", "Định hướng V1"],
        [
            (
                "Bản ghi Web/PWA",
                "Blob URL và danh sách tạm trong bộ nhớ",
                "Reload/đóng trình duyệt có thể làm mất file",
                "Lưu Blob WebM/Opus vào IndexedDB",
            ),
            (
                "Lịch sử bản ghi Web",
                "Metadata trong localStorage",
                "Metadata còn nhưng file blob có thể không còn",
                "Metadata và file cùng quản lý trong IndexedDB",
            ),
            (
                "Tiến độ Web",
                "localStorage: innotrik.listening-progress.v1",
                "Không đồng bộ backend; xóa site data sẽ mất",
                "IndexedDB local-first + API đồng bộ tiến độ",
            ),
            (
                "Tiến độ APK",
                "File JSON trong Application Support Directory",
                "Không chia sẻ với Web; mất khi gỡ/xóa dữ liệu",
                "Giữ local + đồng bộ backend theo Android clientId",
            ),
            (
                "Audio APK",
                "File riêng trong vùng ứng dụng",
                "Không khôi phục khi đổi máy nếu chưa upload",
                "Giữ local; upload R2 nền nếu được bật",
            ),
        ],
        [1450, 2440, 2600, 2870],
    )
    add_body(
        document,
        "Nhận định: vấn đề chính không phải dung lượng metadata mà là vòng đời file audio và danh tính ẩn danh. IndexedDB giải quyết độ bền local của Web; backend giải quyết sao lưu; R2 giải quyết lưu file lớn. Ba lớp này bổ sung cho nhau, không thay thế hoàn toàn cho nhau.",
        bold_lead="Nhận định: ",
    )

    document.add_heading("3. Kiến trúc mục tiêu V1", level=1)
    add_code_block(
        document,
        "WEB/PWA                         BACKEND                         APK\n"
        "IndexedDB  ── sync queue ──>   PostgreSQL   <── sync queue ──  Local JSON/File\n"
        "Blob WebM  ── signed upload ─> Cloudflare R2 <─ signed upload ─ Audio local\n"
        "web clientId                   API + auth                      android clientId",
    )
    add_table(
        document,
        ["Tầng", "Trách nhiệm", "Không nên làm"],
        [
            ("Local Web/APK", "Phản hồi nhanh, hoạt động khi mất mạng, giữ hàng đợi chưa đồng bộ", "Dùng local làm nguồn dữ liệu duy nhất"),
            ("PostgreSQL", "Profile ẩn danh, tiến độ, metadata, trạng thái upload, phiên bản dữ liệu", "Lưu Blob audio trực tiếp"),
            ("Cloudflare R2", "File ghi âm nén, giới hạn 3 bản/câu, vòng đời lưu trữ", "Public URL không giới hạn cho giọng trẻ em"),
            ("Next.js API", "Xác thực, cấp signed URL, hợp nhất tiến độ, idempotency", "Nhận toàn bộ file qua RAM nếu có thể upload thẳng R2"),
        ],
        [1600, 4100, 3660],
    )

    document.add_heading("4. Web/PWA: lưu local và đồng bộ tự động", level=1)
    document.add_heading("4.1. Danh tính ẩn danh", level=2)
    add_body(
        document,
        "Lần đầu mở Web, backend tạo một anonymousProfileId và webDeviceId. Web không yêu cầu email, số điện thoại hoặc mật khẩu. profileId chỉ là mã nhận diện; quyền truy cập phải dựa trên session/token đã ký.",
    )
    add_code_block(
        document,
        "anonymousProfileId: profile_8f72ab...\n"
        "webDeviceId:        web_device_a91c...\n"
        "localStorage key:   innotrik.profile-id.v1\n"
        "IndexedDB:          innotrik-pwa-v1 / identity\n"
        "Cookie:             innotrik_session (HttpOnly, Secure)",
    )
    add_callout(
        document,
        "Bảo mật",
        "Không dùng profileId làm mật khẩu. Cookie session nên là HttpOnly/Secure; LocalStorage và IndexedDB chỉ giữ định danh và dữ liệu local. Nếu frontend và Railway khác domain, cần custom domain hoặc cấu hình cookie/CORS/CSRF đúng chuẩn.",
        fill=PALE_GOLD,
    )

    document.add_heading("4.2. IndexedDB", level=2)
    add_table(
        document,
        ["Object store", "Dữ liệu", "Mục đích"],
        [
            ("identity", "profileId, deviceId, lastSyncedAt", "Khôi phục phiên local và đối chiếu backend"),
            ("learning_progress", "lessonId, câu hiện tại, số câu hoàn thành, skipped", "Chạy nhanh và offline"),
            ("recordings", "Blob WebM/Opus, duration, uploadStatus", "Không mất bản ghi khi reload"),
            ("sync_queue", "mutationId, operation, payload, retryCount", "Gửi lại an toàn khi mạng trở lại"),
            ("settings", "ngôn ngữ, nhóm tuổi, giới hạn cache", "Giữ trải nghiệm cá nhân trên Web"),
        ],
        [1900, 3640, 3820],
    )
    add_body(
        document,
        "Ứng dụng nên gọi navigator.storage.persist() sau một thao tác có chủ đích như hoàn thành bài đầu hoặc cài PWA. Đồng thời theo dõi navigator.storage.estimate(), giới hạn local khoảng 50–100 MB, không xóa bản chưa upload và luôn giữ tối đa 3 bản ghi thành công gần nhất cho mỗi câu.",
    )

    document.add_heading("4.3. Đồng bộ tiến độ", level=2)
    for step in (
        "Ghi thay đổi vào IndexedDB ngay lập tức.",
        "Cập nhật giao diện, không chờ mạng.",
        "Tạo mutationId và đưa tác vụ vào sync_queue.",
        "Gửi backend khi mở app, đổi foreground, mạng trở lại hoặc sau thay đổi quan trọng.",
        "Nhận revision mới từ backend rồi đánh dấu tác vụ đã hoàn thành.",
    ):
        add_list_item(document, step, number_id)
    add_body(
        document,
        "Quy tắc hợp nhất đề xuất: completedSentences lấy giá trị lớn hơn; bài đã hoàn thành không tự lùi; currentSentenceIndex lấy bản cập nhật mới nhất; trạng thái từng câu dùng updatedAt/revision. mutationId giúp request lặp lại không tạo dữ liệu trùng.",
    )

    document.add_heading("5. Backend và lưu file audio", level=1)
    document.add_heading("5.1. Mô hình dữ liệu", level=2)
    add_table(
        document,
        ["Bảng", "Khóa chính/quan hệ", "Nội dung chính"],
        [
            ("anonymous_profiles", "profile_id", "Hồ sơ ẩn danh, trạng thái, thời điểm tạo"),
            ("profile_devices", "device_id → profile_id", "Web/APK device, token hash, last_seen"),
            ("learning_progress", "profile_id + lesson_id", "Tiến độ, revision, updated_at"),
            ("lesson_recordings", "recording_id", "R2 object key, câu, duration, trạng thái"),
            ("recovery_tokens", "token_hash → profile_id", "Mã khôi phục một lần, hạn sử dụng"),
        ],
        [2350, 2850, 4160],
    )
    document.add_heading("5.2. API đề xuất", level=2)
    for endpoint in (
        "POST /api/anonymous-profiles — tạo profile/session lần đầu.",
        "GET /api/learning-progress — tải snapshot tiến độ.",
        "PATCH /api/learning-progress/:lessonId — upsert tiến độ idempotent.",
        "POST /api/lesson-recordings/upload-url — cấp signed URL upload R2.",
        "POST /api/lesson-recordings/complete — xác nhận file và metadata.",
        "GET /api/lesson-recordings — tải danh sách bản ghi theo quyền.",
        "DELETE /api/lesson-recordings/:id — xóa metadata và R2 object.",
        "POST /api/profile/recovery-code — tạo mã/QR khôi phục.",
        "POST /api/profile/recover — dùng mã khôi phục và cấp session mới.",
    ):
        add_list_item(document, endpoint, bullet_id)
    add_body(
        document,
        "Audio nên dùng WebM/Opus hoặc AAC. Một đoạn Opus 5 giây ở 24 kbps chỉ khoảng 15 KB, trong khi WAV PCM 24 kHz khoảng 240 KB. Upload trực tiếp lên R2 bằng signed URL giúp Next.js không phải giữ file trong RAM và giảm tải rõ rệt.",
    )

    document.add_heading("6. QR khôi phục Web/PWA", level=1)
    add_body(
        document,
        "QR khôi phục không phải bước onboarding và không bắt buộc để sử dụng Web. Nó chỉ là phao cứu dữ liệu khi cookie, LocalStorage và IndexedDB đều bị xóa hoặc người dùng đổi máy/trình duyệt.",
    )
    add_table(
        document,
        ["Tình huống", "Có thể tự phục hồi?", "Xử lý"],
        [
            ("Chỉ xóa biểu tượng PWA", "Thường có", "Mở lại cùng domain và trình duyệt; dữ liệu site thường còn"),
            ("Cookie mất, IndexedDB còn", "Có", "Khôi phục session thiết bị rồi tải backend"),
            ("IndexedDB mất, session còn", "Có", "Tải lại tiến độ/metadata từ backend"),
            ("Toàn bộ site data bị xóa", "Không tự động", "Quét QR hoặc nhập mã khôi phục"),
            ("Đổi điện thoại/trình duyệt", "Không tự động", "Dùng QR khôi phục"),
            ("Không lưu QR và mất toàn bộ site data", "Không", "Tạo hồ sơ mới; không thể xác định hồ sơ cũ"),
        ],
        [2600, 1900, 4860],
        center_cols=(1,),
    )
    document.add_heading("6.1. Trải nghiệm phụ huynh", level=2)
    for step in (
        "Sau khi hoàn thành bài đầu, hiển thị gợi ý nhẹ: ‘Lưu mã khôi phục để không mất tiến độ khi đổi máy’. Có nút Để sau.",
        "Khi chọn lưu, backend tạo token ngẫu nhiên; QR chứa URL có token ở phần fragment (#code=...).",
        "Phụ huynh lưu ảnh, gửi vào Zalo/WeChat của chính mình hoặc lưu mã chữ dự phòng.",
        "Khi khôi phục, token được dùng một lần, backend cấp session mới và yêu cầu tạo QR mới.",
    ):
        add_list_item(document, step, number_id)
    add_callout(
        document,
        "Giới hạn bắt buộc",
        "Nếu không có tài khoản, QR/mã khôi phục hoặc mã thiết bị riêng thì backend không thể nhận biết người cũ sau khi toàn bộ dữ liệu trình duyệt bị xóa. IndexedDB và persistent storage chỉ giảm rủi ro, không loại bỏ hoàn toàn.",
        fill=PALE_RED,
    )

    document.add_heading("7. APK: lưu dữ liệu và cập nhật bắt buộc", level=1)
    add_body(
        document,
        "APK dùng Android clientId, lưu tiến độ và audio trong vùng riêng của ứng dụng. V1 có thể đồng bộ backend riêng cho APK mà không liên kết với Web. QR không cần thiết cho cập nhật APK thông thường.",
    )
    add_table(
        document,
        ["Trường hợp APK", "Dữ liệu local", "Có cần QR?"],
        [
            ("Cài đè bản release mới", "Được giữ", "Không"),
            ("Cập nhật bắt buộc qua APK", "Được giữ", "Không"),
            ("Gỡ ứng dụng rồi cài lại", "Có thể mất", "Có ích nếu cần phục hồi"),
            ("Xóa dữ liệu ứng dụng", "Mất", "Có ích"),
            ("Đổi điện thoại", "Không tự chuyển", "Có ích hoặc dùng cơ chế chuyển thiết bị"),
            ("Debug → release chính thức", "Phải gỡ bản Debug", "Chỉ cần nếu muốn giữ dữ liệu test"),
        ],
        [3450, 3000, 2910],
        center_cols=(1, 2),
    )
    document.add_heading("7.1. Điều kiện cài đè", level=2)
    for item in (
        "Giữ application ID com.innotrik.aispeaking.",
        "Tất cả bản chính thức dùng cùng một release keystore.",
        "Build number sau dấu + luôn tăng: 1.0.0+1 → 1.0.1+2 → 1.0.2+3.",
        "Không hướng dẫn người dùng gỡ ứng dụng trong màn hình cập nhật.",
        "Sao lưu release keystore ở ít nhất hai vị trí bảo mật và lưu mật khẩu trong CI secrets.",
    ):
        add_list_item(document, item, bullet_id)
    document.add_heading("7.2. Quy trình phát hành", level=2)
    for step in (
        "Tăng version/build number.",
        "Build bằng release keystore chính thức.",
        "Cài thử đè lên bản release cũ và xác nhận dữ liệu còn nguyên.",
        "Kiểm tra nói, dịch, phát audio, lịch sử, luyện nghe và BLE nếu có.",
        "Upload APK có tên phiên bản lên R2.",
        "Cập nhật latestBuild và trang tải cố định.",
        "Theo dõi lỗi; cuối cùng mới tăng minimumSupportedBuild để bắt buộc cập nhật.",
    ):
        add_list_item(document, step, number_id)

    document.add_heading("8. Quan hệ Web và APK trong V1", level=1)
    add_callout(
        document,
        "Quyết định V1",
        "Web và APK lưu/sync riêng. Không yêu cầu phụ huynh quét QR để liên kết hai nền tảng. QR Web chỉ phục vụ khôi phục Web; QR APK chỉ cân nhắc sau nếu có nhu cầu đổi máy hoặc đồng bộ chéo.",
        fill=PALE_GREEN,
    )
    add_table(
        document,
        ["Nội dung", "Web/PWA V1", "APK V1"],
        [
            ("Định danh", "web clientId + anonymous profile", "android clientId + anonymous profile"),
            ("Local progress", "IndexedDB", "JSON/SQLite/file trong app"),
            ("Local audio", "IndexedDB Blob WebM/Opus", "File trong vùng ứng dụng"),
            ("Cloud progress", "PostgreSQL theo web profile", "PostgreSQL theo APK profile"),
            ("Cloud audio", "R2 nếu bật upload", "R2 nếu bật upload"),
            ("Khôi phục", "QR/mã khôi phục tùy chọn", "Cài đè không cần QR; gỡ/đổi máy mới cần"),
            ("Đồng bộ Web–APK", "Chưa bật", "Chưa bật"),
        ],
        [2300, 3530, 3530],
    )
    add_body(
        document,
        "Giai đoạn sau có thể tạo childProfileId chung và liên kết Web–APK bằng QR một lần. Đây là tính năng đồng bộ đa nền tảng, không nên trộn với nhiệm vụ sao lưu riêng của từng client trong V1.",
    )

    document.add_heading("9. Bảo mật và quyền riêng tư", level=1)
    for item in (
        "Giọng nói trẻ em là dữ liệu nhạy cảm: cần thông báo rõ mục đích, thời hạn lưu và quyền xóa.",
        "R2 bucket để private; nghe lại bằng signed URL có thời hạn.",
        "Không đặt API key Cloudflare/OpenAI trong Web hoặc APK.",
        "Token khôi phục phải đủ ngẫu nhiên, backend chỉ lưu hash và hỗ trợ thu hồi.",
        "Có API xóa profile, tiến độ và toàn bộ file audio liên quan.",
        "Telemetry không ghi nội dung token, URL ký tạm hoặc dữ liệu âm thanh thô.",
        "Áp dụng rate limit cho tạo profile, upload audio và khôi phục.",
    ):
        add_list_item(document, item, bullet_id)

    document.add_heading("10. Lộ trình triển khai đề xuất", level=1)
    add_table(
        document,
        ["Ưu tiên", "Hạng mục", "Kết quả bàn giao", "Điều kiện nghiệm thu"],
        [
            ("P0", "IndexedDB Web", "Lưu Blob, metadata, progress, sync_queue", "Reload/đóng mở PWA không mất bản ghi"),
            ("P0", "Progress API", "Profile ẩn danh và GET/PATCH tiến độ", "Web tải lại đúng câu sau đăng nhập phiên"),
            ("P0", "Release APK", "Keystore, versionCode, cài đè", "Dữ liệu local còn nguyên sau update"),
            ("P1", "Audio R2", "Signed upload, metadata, max 3 bản/câu", "Không upload trùng; xóa bản cũ đúng quy tắc"),
            ("P1", "Khôi phục Web", "QR/mã một lần và session mới", "Xóa site data vẫn phục hồi được bằng QR"),
            ("P1", "Version API", "Latest/minimum/update type/message/URL", "Cập nhật tùy chọn và bắt buộc đúng luồng"),
            ("P2", "Đồng bộ Web–APK", "childProfileId và ghép nối", "Hai nền tảng dùng chung tiến độ khi chủ động liên kết"),
        ],
        [900, 2180, 3180, 3100],
        center_cols=(0,),
    )

    document.add_heading("11. Tiêu chí nghiệm thu V1", level=1)
    checklist = (
        "Web ghi âm xong, reload trang vẫn nghe lại được bản mới nhất.",
        "Mỗi câu giữ tối đa 3 bản ghi thành công; lỗi upload không xóa bản local trước.",
        "Mất mạng vẫn học và ghi âm; khi có mạng, hàng đợi đồng bộ tự chạy lại.",
        "Tiến độ backend không bị lùi khi request đến sai thứ tự.",
        "Xóa riêng cookie hoặc riêng IndexedDB vẫn có đường phục hồi tự động phù hợp.",
        "Xóa toàn bộ site data phục hồi được bằng QR hợp lệ; QR đã dùng không dùng lại được.",
        "APK release mới cài đè bản cũ, giữ nguyên tiến độ, audio và thiết lập.",
        "APK bắt buộc cập nhật không yêu cầu gỡ ứng dụng và không dùng Debug key.",
        "Không có API key hoặc secret R2 trong bundle Web/APK.",
        "Có chức năng xóa dữ liệu trẻ em và file R2 tương ứng.",
    )
    for item in checklist:
        add_list_item(document, f"□ {item}", bullet_id)

    document.add_heading("12. Kết luận", level=1)
    add_body(
        document,
        "Kiến trúc phù hợp nhất cho V1 là local-first và lưu riêng theo từng client: Web dùng IndexedDB, APK dùng vùng lưu trữ riêng, cả hai có thể đồng bộ backend độc lập. PostgreSQL chỉ lưu tiến độ và metadata; R2 lưu audio nén. QR là cơ chế khôi phục tùy chọn dành cho Web khi mất toàn bộ site data, không phải bước bắt buộc và không liên quan đến cập nhật APK cài đè.",
    )
    add_body(
        document,
        "Khi sản phẩm cần học tiếp xuyên suốt giữa Web và APK, nhóm dự án mới bổ sung childProfileId chung và ghép nối một lần. Việc tách hai giai đoạn giúp V1 dễ sử dụng, ít rủi ro và triển khai nhanh hơn.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
