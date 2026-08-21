from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("output/docx/Ke_hoach_ca_nhan_7_ngay_H20_iOS_Native_20-26_08_2026.docx")

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
NAVY = "17365D"
TEXT = "243447"
MUTED = "5E6B7A"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FB"
PALE_GREEN = "E8F4EC"
GREEN = "26734D"
PALE_AMBER = "FFF3D6"
AMBER = "9A6700"
PALE_RED = "FCE8E6"
RED = "B42318"
LIGHT_GRAY = "F5F6F8"
MID_GRAY = "D7DCE2"
WHITE = "FFFFFF"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths: list[int], header=True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        if header and row_index == 0:
            set_repeat_table_header(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text: str, *, bold=False, color=TEXT, size=9.5, align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_cell_lines(cell, lines: list[tuple[str, bool, str]], size=9.25) -> None:
    cell.text = ""
    for index, (text, bold, color) in enumerate(lines):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2 if index < len(lines) - 1 else 0)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_run(paragraph, text: str, *, bold=False, color=TEXT, size=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(color)
    if size is not None:
        run.font.size = Pt(size)
    return run


def add_bullet(doc, text: str, *, bold_prefix: str | None = None, color=TEXT):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True, color=color)
        add_run(p, text[len(bold_prefix):], color=color)
    else:
        add_run(p, text, color=color)
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    add_run(p, text)
    return p


def add_callout(doc, heading: str, body: str, *, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=2)
    configure_table(table, [240, 9120], header=False)
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_shading(table.cell(0, 1), fill)
    set_cell_margins(table.cell(0, 0), 0, 0, 0, 0)
    set_cell_margins(table.cell(0, 1), 130, 170, 130, 170)
    table.cell(0, 0).text = ""
    add_cell_lines(table.cell(0, 1), [(heading, True, accent), (body, False, TEXT)], size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text: str, *, color=TEXT, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True, color=color)
        add_run(p, text[len(bold_prefix):], color=color)
    else:
        add_run(p, text, color=color)
    return p


def set_table_header(table, labels: list[str]) -> None:
    for index, label in enumerate(labels):
        cell = table.cell(0, index)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, label, bold=True, color=DEEP_BLUE, size=9.5)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 24, NAVY, 0, 10),
        ("Subtitle", 12, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DEEP_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(TEXT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Eyebrow" not in styles:
        style = styles.add_style("Eyebrow", 1)
    eyebrow = styles["Eyebrow"]
    eyebrow.font.name = "Calibri"
    eyebrow.font.size = Pt(9)
    eyebrow.font.bold = True
    eyebrow.font.color.rgb = RGBColor.from_string(BLUE)
    eyebrow.paragraph_format.space_after = Pt(6)

    if "Small Note" not in styles:
        style = styles.add_style("Small Note", 1)
    note = styles["Small Note"]
    note.font.name = "Calibri"
    note.font.size = Pt(8.5)
    note.font.color.rgb = RGBColor.from_string(MUTED)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.1


def setup_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, width in enumerate((4680, 4680)):
        set_cell_width(table.cell(0, index), width)
        set_cell_margins(table.cell(0, index), 0, 0, 0, 0)
    set_cell_text(table.cell(0, 0), "INNOTRIK • H20 • KẾ HOẠCH CÁ NHÂN", bold=True, color=DEEP_BLUE, size=8.5)
    set_cell_text(
        table.cell(0, 1),
        "20–26/08/2026",
        bold=True,
        color=MUTED,
        size=8.5,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    if header.paragraphs:
        header.paragraphs[0]._element.getparent().remove(header.paragraphs[0]._element)

    footer = section.footer
    ft = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    ft.autofit = False
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, width in enumerate((3300, 2760, 3300)):
        set_cell_width(ft.cell(0, index), width)
        set_cell_margins(ft.cell(0, index), 0, 0, 0, 0)
    set_cell_text(ft.cell(0, 0), "Nội bộ • Bản 1.0", color=MUTED, size=8)
    set_cell_text(ft.cell(0, 1), "Phạm vi: công việc cá nhân", color=MUTED, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    right = ft.cell(0, 2)
    right.text = ""
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Trang ", color=MUTED, size=8)
    add_page_field(p)
    if footer.paragraphs:
        footer.paragraphs[0]._element.getparent().remove(footer.paragraphs[0]._element)


def build_document() -> Document:
    doc = Document()
    setup_styles(doc)
    setup_header_footer(doc)

    properties = doc.core_properties
    properties.title = "Kế hoạch cá nhân 7 ngày H20 và iOS Native"
    properties.subject = "Ưu tiên P0/P1/P2, tồn đọng, iOS Native, App Store và báo cáo quản trị"
    properties.author = "INNOTRIK"
    properties.keywords = "H20, iOS Native, App Store, BLE, HFP, MAIN, kế hoạch 7 ngày"

    eyebrow = doc.add_paragraph(style="Eyebrow")
    add_run(eyebrow, "BẢN KẾ HOẠCH CÁ NHÂN • 7 NGÀY • ƯU TIÊN THEO RỦI RO")

    title = doc.add_paragraph(style="Title")
    add_run(title, "H20, iOS Native & App Store", bold=True, color=NAVY, size=24)

    subtitle = doc.add_paragraph(style="Subtitle")
    add_run(
        subtitle,
        "Đóng băng nền Android đã kết nối thiết bị, xử lý tồn đọng quan trọng, "
        "tạo bằng chứng iOS Native và chốt báo cáo điều hành cho sếp.",
        color=MUTED,
        size=12,
    )

    meta = doc.add_table(rows=2, cols=4)
    configure_table(meta, [1650, 3030, 1650, 3030], header=False)
    labels = [
        ("Giai đoạn", "20–26/08/2026"),
        ("Code gốc", "main @ 2ec38ab"),
        ("Trạng thái", "Analyze xanh • BLE unit test 10/10"),
        ("Phạm vi", "Chỉ công việc cá nhân; không tính việc đồng nghiệp"),
    ]
    for idx, (label, value) in enumerate(labels):
        row = idx // 2
        col = (idx % 2) * 2
        set_cell_shading(meta.cell(row, col), LIGHT_BLUE)
        set_cell_shading(meta.cell(row, col + 1), PALE_BLUE)
        set_cell_text(meta.cell(row, col), label, bold=True, color=DEEP_BLUE, size=9)
        set_cell_text(meta.cell(row, col + 1), value, color=TEXT, size=9)

    add_heading(doc, "Kết luận điều hành", 1)
    add_callout(
        doc,
        "Mục tiêu thực tế sau 7 ngày",
        "Một Android/H20 baseline có bằng chứng, một PoC iOS Native chứng minh BLE MAIN + đường âm thanh H20, "
        "một bản signed/TestFlight-ready nếu đủ Mac và tài khoản Apple, cùng báo cáo P0/P1 để sếp ra quyết định. "
        "Không cam kết Apple duyệt App Store trong 7 ngày vì còn phụ thuộc tài khoản, thiết bị, ODM và thời gian review.",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_bullet(doc, "Khẩn cấp nhất: khóa phiên bản đang chạy được và đo lại trên H20 thật trước khi tiếp tục thay đổi nhiều luồng.", bold_prefix="Khẩn cấp nhất:")
    add_bullet(doc, "Quan trọng nhất: iOS phải là Flutter + native bridge, không phải đóng gói PWA rồi kỳ vọng nhận MAIN/HFP như APK.", bold_prefix="Quan trọng nhất:")
    add_bullet(doc, "Nguyên tắc làm việc: mỗi lỗi có bước tái hiện, log, video, commit và tiêu chí đạt; không vá nhiều tầng cùng lúc.", bold_prefix="Nguyên tắc làm việc:")
    add_bullet(doc, "Ranh giới: không nhận thay phần nội dung, nhạc, UI hay module của đồng nghiệp; chỉ tiếp nhận chúng như đầu vào tích hợp.", bold_prefix="Ranh giới:")

    add_heading(doc, "Ma trận ưu tiên", 2)
    priorities = doc.add_table(rows=1, cols=4)
    configure_table(priorities, [900, 2300, 3560, 2600])
    set_table_header(priorities, ["Mức", "Thời hạn", "Trọng tâm cá nhân", "Điều kiện hoàn tất"])
    priority_rows = [
        ("P0", "0–48 giờ", "Đóng băng Android/H20; tái kiểm tra HFP, BLE MAIN, reconnect; chốt protocol ODM; mở đường iOS Native.", "Có bằng chứng và không còn lỗi chặn luồng chính."),
        ("P1", "Ngày 3–6", "PoC CoreBluetooth + AVAudioSession; xử lý interruption/mạng/background; chuẩn bị signing và privacy.", "Chạy được trên iPhone thật hoặc ghi rõ blocker có chủ sở hữu."),
        ("P2", "Ngày 7 / sau đó", "Tối ưu hiệu năng, mở rộng ma trận thiết bị, hoàn thiện metadata và cải tiến không chặn release.", "Không làm trễ P0/P1; có lịch sau 7 ngày."),
    ]
    for level, deadline, focus, done in priority_rows:
        row = priorities.add_row()
        fill = PALE_RED if level == "P0" else PALE_AMBER if level == "P1" else PALE_GREEN
        accent = RED if level == "P0" else AMBER if level == "P1" else GREEN
        set_cell_shading(row.cells[0], fill)
        set_cell_text(row.cells[0], level, bold=True, color=accent, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], deadline, bold=True, color=TEXT, size=9.25)
        set_cell_text(row.cells[2], focus, color=TEXT, size=9.25)
        set_cell_text(row.cells[3], done, color=TEXT, size=9.25)

    add_heading(doc, "1. Điểm xuất phát đã xác nhận", 1)
    status = doc.add_table(rows=1, cols=3)
    configure_table(status, [1800, 3300, 4260])
    set_table_header(status, ["Hạng mục", "Hiện trạng", "Ý nghĩa / việc còn phải chứng minh"])
    status_rows = [
        ("Android/H20", "HFP mic/loa và BLE Control đã hình thành; volume hoạt động; MAIN có raw hex.", "Có nền tích hợp thật. Cần regression có kiểm soát, không chỉ thử thủ công một lần."),
        ("Tự kết nối BLE", "Code đã ưu tiên thiết bị đã lưu, service 9E3B0001 và có retry khi tiến trình APP còn sống.", "Cần test force-stop, reboot, out-of-range và nhiều hãng Android; HFP lần đầu vẫn do hệ điều hành ghép đôi."),
        ("Chất lượng code", "flutter analyze: không lỗi; test AIV0 BLE: 10/10 xanh ngày 20/08/2026.", "Native HFP/GATT chưa có kiểm thử tự động trực tiếp; vẫn cần hardware evidence."),
        ("Protocol ODM", "Service 9E3B0001; MAIN 9E3B0002 Indicate; APP state 9E3B0003 Write; pin/firmware đã quan sát.", "Cần văn bản chốt packet, sequence/checksum, duplicate, ACK và hành vi reconnect."),
        ("iOS hiện tại", "AppDelegate mới có client identity; chưa có CoreBluetooth H20 hoặc bridge AVAudioSession riêng.", "iOS Web/PWA không thể thay vai trò iOS Native cho MAIN vật lý và xác nhận route H20."),
        ("Nút điều khiển", "Sản phẩm đã bỏ long press; luồng chủ lực dùng short MAIN + lời nói theo state.", "Phải dọn checklist/test cũ còn long press để tránh ODM, QA và APP hiểu khác nhau."),
    ]
    for i, values in enumerate(status_rows):
        row = status.add_row()
        if i % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col, value in enumerate(values):
            set_cell_text(row.cells[col], value, bold=(col == 0), color=TEXT, size=9.15)

    add_heading(doc, "2. Tồn đọng cần cá nhân tôi chịu trách nhiệm", 1)
    backlog = doc.add_table(rows=1, cols=5)
    configure_table(backlog, [760, 2320, 2900, 1880, 1500])
    set_table_header(backlog, ["ID", "Vấn đề", "Rủi ro nếu chưa xử lý", "Bằng chứng cần có", "Ưu tiên"])
    backlog_rows = [
        ("A-01", "Chưa có iOS Native bridge cho BLE MAIN và audio route H20.", "Không thể tuyên bố App Store hỗ trợ thiết bị.", "PoC iPhone thật: 20/20 MAIN; mic/loa H20 xác nhận.", "P0"),
        ("A-02", "Chưa khóa đầy đủ protocol ODM/ACK APP state; test cũ còn long press.", "Sai state, duplicate hoặc hiểu sai firmware.", "Protocol versioned + bảng packet + phản hồi ODM bằng văn bản.", "P0"),
        ("A-03", "Auto reconnect chưa được chứng minh qua force-stop/reboot/nhiều hãng.", "Người dùng vẫn phải vào cài đặt hoặc kết nối lại.", "Ma trận 10 vòng/scenario, log reconnect và video lỗi.", "P0"),
        ("A-04", "Luồng nói liên tục/từng câu, lời dẫn và điều hướng đã từng hồi quy qua nhiều commit.", "Demo thất bại dù kết nối thiết bị tốt.", "Smoke suite luồng chính + tag commit ổn định.", "P0"),
        ("A-05", "Background/lock/call/notification/audio-focus còn phụ thuộc nền tảng.", "Ghi âm/học bị ngắt hoặc route sai.", "Kịch bản interruption và resume; không mất state.", "P1"),
        ("A-06", "Mất mạng/khôi phục mạng và thông báo lỗi cần được đo lại end-to-end.", "Báo nhầm thành công, mất tiến độ hoặc kẹt UI.", "10/10 recovery; lỗi nói rõ; retry không nhân đôi request.", "P1"),
        ("A-07", "Thiếu benchmark giọng trẻ, vùng miền, nhiễu và latency P50/P95.", "Không biết lỗi do ASR, mạng, HFP hay state machine.", "Dataset test có nhãn; phân loại lỗi; báo cáo số liệu.", "P1"),
        ("A-08", "App Store privacy, consent, quyền mic/Bluetooth và child-data chưa chốt.", "Bị review từ chối hoặc rủi ro dữ liệu trẻ em.", "Data map, retention/deletion, permission copy, privacy decision.", "P1"),
        ("A-09", "Chưa có release gate và báo cáo một nguồn duy nhất.", "Sếp/ODM nhận thông tin khác nhau; khó go/no-go.", "Báo cáo điều hành 1 trang.", "P1"),
    ]
    for idx, values in enumerate(backlog_rows):
        row = backlog.add_row()
        for col, value in enumerate(values):
            if col == 4:
                fill = PALE_RED if value == "P0" else PALE_AMBER
                accent = RED if value == "P0" else AMBER
                set_cell_shading(row.cells[col], fill)
                set_cell_text(row.cells[col], value, bold=True, color=accent, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                if idx % 2 == 1:
                    set_cell_shading(row.cells[col], LIGHT_GRAY)
                set_cell_text(row.cells[col], value, bold=(col == 0), color=TEXT, size=8.75)

    add_heading(doc, "3. Kiến trúc iOS Native cần chốt", 1)
    add_body(doc, "Không “build web lên App Store”. Giữ Flutter UI, domain, backend và nội dung; bổ sung native iOS bridge ở đúng lớp phần cứng:")
    architecture = doc.add_table(rows=1, cols=3)
    configure_table(architecture, [2300, 3200, 3860])
    set_table_header(architecture, ["Lớp", "Giữ lại", "Làm mới / kiểm chứng trên iOS"])
    architecture_rows = [
        ("Flutter dùng chung", "UI, navigation, lesson/vocabulary, controller, rule/cache, API.", "Tách capability theo platform; không dùng nhánh web để giả lập BLE/HFP."),
        ("BLE Control", "State machine MAIN ở Dart nếu phù hợp.", "CoreBluetooth: scan 9E3B0001, subscribe 9E3B0002, write 9E3B0003, pin/firmware/reconnect."),
        ("Audio", "VAD, PCM/batch fallback, playback/cache.", "AVAudioSession playAndRecord/voice processing; xác nhận current route H20; interruption và route change."),
        ("ASR", "Cloudflare Batch fallback và backend hiện tại.", "PoC native streaming/Apple Speech chỉ khi KPI tốt; một audio pipeline, không mở hai recorder tranh mic."),
        ("Phát hành", "Bundle ID và source iOS có sẵn.", "Mac/Xcode, signing, Info.plist, Privacy, TestFlight, review notes và phone fallback."),
    ]
    for idx, values in enumerate(architecture_rows):
        row = architecture.add_row()
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col, value in enumerate(values):
            set_cell_text(row.cells[col], value, bold=(col == 0), color=TEXT, size=9.15)

    add_callout(
        doc,
        "Điều kiện chặn bắt buộc",
        "Muốn build/ký iOS cần Mac, Xcode, iPhone thật và Apple Developer đang hoạt động. Nếu thiếu một trong các đầu vào này, "
        "kết quả tuần đầu là tài liệu kiến trúc + branch iOS + blocker, không được báo là đã có bản App Store.",
        fill=PALE_AMBER,
        accent=AMBER,
    )

    doc.add_page_break()
    add_heading(doc, "4. Kế hoạch cá nhân 7 ngày", 1)
    add_body(doc, "Mỗi ngày chỉ chốt một nhóm rủi ro chính. Cuối ngày phải có bằng chứng và một quyết định rõ ràng, không chỉ có mô tả “đã thử”.")

    schedule = doc.add_table(rows=1, cols=5)
    configure_table(schedule, [800, 920, 2460, 3320, 1860])
    set_table_header(schedule, ["Ngày", "Mốc", "Mục tiêu", "Việc cá nhân thực hiện", "Đầu ra / Gate"])
    schedule_rows = [
        ("1", "20/08", "Đóng băng baseline", "Tạo branch/tag ổn định; ghi commit, firmware, thiết bị; build APK; chạy smoke HFP/BLE/MAIN; lập issue A-01…A-09; xóa long-press khỏi phạm vi sản phẩm.", "APK baseline + issue register + thư mục evidence."),
        ("2", "21/08", "Chốt P0 Android/ODM", "Test HFP in/out, BLE+HFP, MAIN 20 lượt, reconnect, BT off/on, out-of-range, lock/background, cuộc gọi, mạng yếu; phân loại APP/firmware/OS; gửi ODM bảng packet cần xác nhận.", "P0 pass/fail + báo cáo ODM có log/video."),
        ("3", "22/08", "Mở dự án iOS Native", "Xác nhận Mac/Xcode/iPhone/Apple Developer; tạo branch iOS; chạy app trên iPhone; chốt architecture Flutter + Swift; rà bundle ID, display name, permission và capability tối thiểu.", "Build debug iPhone + architecture decision record."),
        ("4", "23/08", "PoC BLE MAIN", "Viết CoreBluetooth bridge; scan service 9E3B0001; subscribe Indicate 9E3B0002; log short MAIN; đọc pin/firmware; thử reconnect và chống duplicate.", "20/20 MAIN; 10 vòng reconnect; log packet."),
        ("5", "24/08", "PoC H20 audio + flow", "Cấu hình AVAudioSession; xác nhận input/output route; record/playback H20; xử lý route change/interruption; nối MAIN → prompt → nghe → intent → topic/translation; giữ phone fallback.", "Video end-to-end + route diagnostics + lỗi còn lại."),
        ("6", "25/08", "Release dry-run", "Tăng version/build; signed archive; TestFlight internal nếu tài khoản sẵn sàng; rà privacy/data map, consent, retention/deletion, review mode; đo P50/P95 và first-attempt capture.", "IPA/archive hợp lệ hoặc blocker đỏ; compliance checklist."),
        ("7", "26/08", "Đóng việc & báo cáo sếp", "Rerun regression Android + iOS; đối chiếu P0/P1; tổng hợp ODM response; ra quyết định go/no-go; chốt 14 ngày tiếp theo, chi phí/rủi ro và người cần quyết định.", "Báo cáo 1 trang + phụ lục evidence + decision request."),
    ]
    for idx, values in enumerate(schedule_rows):
        row = schedule.add_row()
        fill = PALE_RED if idx < 2 else PALE_BLUE if idx < 5 else PALE_GREEN
        for col, value in enumerate(values):
            if col == 0:
                set_cell_shading(row.cells[col], fill)
            elif idx % 2 == 1:
                set_cell_shading(row.cells[col], LIGHT_GRAY)
            set_cell_text(row.cells[col], value, bold=(col in (0, 1)), color=TEXT, size=8.65, align=WD_ALIGN_PARAGRAPH.CENTER if col in (0, 1) else None)

    doc.add_page_break()
    add_heading(doc, "5. Release gate và tiêu chí đo", 1)
    gates = doc.add_table(rows=1, cols=4)
    configure_table(gates, [1850, 3700, 2100, 1710])
    set_table_header(gates, ["Gate", "Tiêu chí cá nhân phải chứng minh", "Kết quả cần lưu", "Quyết định"])
    gate_rows = [
        ("G0 • Baseline", "Analyze không lỗi; test nền xanh; build từ commit đã ghi; không có thay đổi chưa kiểm soát.", "Log CI/local + APK/hash.", "Cho phép test H20."),
        ("G1 • Android/H20", "HFP mic/loa đúng; short MAIN 20/20, không duplicate; BLE/HFP đồng thời; reconnect đạt ma trận.", "Video, raw hex, route và reconnect log.", "Giữ hoặc mở issue ODM."),
        ("G2 • iOS BLE", "iPhone nhận MAIN qua 9E3B0002 và reconnect; không phụ thuộc PWA.", "Xcode log + video 20 lượt.", "Cho phép tích hợp flow."),
        ("G3 • iOS Audio", "Mic/loa H20 được xác nhận; interruption/resume đúng; phone fallback hoạt động.", "Route dump + audio sample/video.", "Cho phép TestFlight."),
        ("G4 • Release", "Không crash/stuck; lỗi mạng rõ; privacy/permission khớp; signed archive; review mode dùng không cần H20.", "Checklist ký + privacy + build.", "Submit TestFlight / hoãn."),
    ]
    for idx, values in enumerate(gate_rows):
        row = gates.add_row()
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col, value in enumerate(values):
            set_cell_text(row.cells[col], value, bold=(col == 0), color=TEXT, size=9)

    add_heading(doc, "KPI thử nghiệm nội bộ", 2)
    add_bullet(doc, "Short MAIN: 20/20 sự kiện hợp lệ; 0 duplicate; 0 kích hoạt ghi âm ngoài ý muốn.")
    add_bullet(doc, "Reconnect BLE: 10/10 cho từng kịch bản mục tiêu; nếu fail phải có phân loại OS/APP/firmware.")
    add_bullet(doc, "HFP route: 100% phiên test đã chọn H20 phải xác nhận đúng input và output.")
    add_bullet(doc, "Mạng mất/khôi phục: không báo nhầm thành công, không mất trạng thái, không gửi trùng.")
    add_bullet(doc, "Latency: đo end-of-speech → playback; mục tiêu thử nghiệm P50 ≤ 2,5 giây, P95 ≤ 4 giây; chưa dùng làm cam kết marketing.")
    add_bullet(doc, "Phải nói lại: mục tiêu ≤ 5% trên tập test đã kiểm soát; lưu nguyên nhân ASR/mạng/audio/state.")

    add_heading(doc, "6. Rủi ro và quyết định cần xin sếp", 1)
    risks = doc.add_table(rows=1, cols=4)
    configure_table(risks, [2550, 2500, 2360, 1950])
    set_table_header(risks, ["Rủi ro / quyết định", "Tác động", "Đề xuất cá nhân", "Hạn chốt"])
    risk_rows = [
        ("Chưa có Mac/Xcode/iPhone/Developer active", "Không build/ký/TestFlight được.", "Cấp hoặc thuê Mac; xác nhận quyền tài khoản ngay Ngày 1–3.", "22/08"),
        ("ODM chưa chốt packet/ACK", "APP state/reconnect có thể sai firmware.", "Chỉ dispatch packet đã biết; unknown để diagnostic; yêu cầu spec versioned.", "21/08"),
        ("Kids/Education và lưu audio", "Ảnh hưởng consent, privacy và App Review.", "Sếp chốt đối tượng phát hành và data retention trước TestFlight external.", "25/08"),
        ("H20 là bắt buộc hay tùy chọn", "Reviewer có thể không có thiết bị.", "Giữ phone fallback và video hardware; H20 là optional trong review mode.", "25/08"),
        ("Phạm vi release 7 ngày", "Dễ báo quá mức thực tế.", "Cam kết PoC/signed internal build; App Store approval là mốc ngoài 7 ngày.", "20/08"),
    ]
    for idx, values in enumerate(risk_rows):
        row = risks.add_row()
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col, value in enumerate(values):
            set_cell_text(row.cells[col], value, bold=(col == 0), color=TEXT, size=9.1)

    add_heading(doc, "7. Mẫu báo cáo cuối ngày cho sếp", 1)
    add_number(doc, "Đã hoàn thành: tối đa 3 dòng, ghi đúng kết quả và commit/build.")
    add_number(doc, "Bằng chứng: link video/log/screenshot/issue; không dùng mô tả cảm tính.")
    add_number(doc, "P0 còn mở: lỗi nào chặn demo/release, ai/đầu vào nào đang chờ.")
    add_number(doc, "Quyết định cần sếp: nêu 1–2 lựa chọn, tác động thời gian/chi phí/rủi ro.")
    add_number(doc, "Kế hoạch 24 giờ tới: 3 việc cá nhân, có điều kiện hoàn tất.")

    add_callout(
        doc,
        "Mẫu câu kết luận",
        "Hôm nay baseline [commit/build] đạt [x/y] tiêu chí. P0 còn [n] lỗi: […]. Tôi đề xuất [go/hold] cho bước […]. "
        "Cần sếp quyết định […] trước [thời hạn]. Ngày mai tôi sẽ hoàn tất […], với gate […].",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "8. Phạm vi loại trừ và đầu vào phụ thuộc", 1)
    exclusions = doc.add_table(rows=1, cols=2)
    configure_table(exclusions, [4680, 4680])
    set_table_header(exclusions, ["Không tính vào khối lượng cá nhân", "Đầu vào tôi cần nhận để hoàn thành"])
    row = exclusions.add_row()
    add_cell_lines(
        row.cells[0],
        [
            ("• Sáng tác/nạp bài hát, nội dung và asset UI.", False, TEXT),
            ("• Phát triển module đồng nghiệp đang sở hữu.", False, TEXT),
            ("• Sửa firmware ODM hoặc sản xuất khuôn/phần cứng.", False, TEXT),
            ("• Refactor lớn không gắn với lỗi P0/P1.", False, TEXT),
        ],
        size=9.25,
    )
    add_cell_lines(
        row.cells[1],
        [
            ("• ODM: packet spec, firmware/version và câu trả lời reconnect/ACK.", False, TEXT),
            ("• Sếp: Apple account, mô hình phát hành, privacy/retention và go/no-go.", False, TEXT),
            ("• Đồng nghiệp: build/module đã ổn định để tôi tích hợp và test; không tính là việc tôi.", False, TEXT),
            ("• Apple: thời gian xử lý TestFlight/App Review là phụ thuộc ngoài nhóm.", False, TEXT),
        ],
        size=9.25,
    )

    note = doc.add_paragraph(style="Small Note")
    add_run(
        note,
        "Căn cứ lập kế hoạch: code main @ 2ec38ab; flutter analyze và AIV0 BLE unit test ngày 20/08/2026; "
        "tài liệu nghiệm thu AIV1; phương án iOS Native/App Store nội bộ; các kết quả H20 thật đã xác nhận. "
        "Yêu cầu App Store và SDK phải được kiểm tra lại theo tài liệu Apple tại ngày submit.",
        color=MUTED,
        size=8.5,
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
