from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docs" / "AIVO_V1_YEU_CAU_TICH_HOP_ODM_VI_ZH.docx"

# rfi_response preset with named overrides for an Asian partner document:
# A4 portrait, monochrome, Arial + Microsoft YaHei, compact bilingual tables.
CONTENT_WIDTH_DXA = 9860
TABLE_INDENT_DXA = 100
CELL_MARGIN_DXA = {"top": 70, "bottom": 70, "start": 100, "end": 100}
BLACK = "000000"
DARK = "333333"
MUTED = "666666"
LIGHT = "F2F2F2"
BORDER = "B7B7B7"


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGIN_DXA.items():
        tag = qn(f"w:{edge}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    assert sum(widths) == CONTENT_WIDTH_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), BORDER)


def set_run(run, *, size=10.5, bold=False, italic=False, color=BLACK, lang="vi-VN") -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    lang_node = r_pr.find(qn("w:lang"))
    if lang_node is None:
        lang_node = OxmlElement("w:lang")
        r_pr.append(lang_node)
    lang_node.set(qn("w:val"), lang)
    lang_node.set(qn("w:eastAsia"), "zh-CN" if lang == "zh-CN" else lang)


def add_text(doc, text: str, *, bold=False, italic=False, size=10.5, after=5, lang="vi-VN", keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_with_next = keep
    set_run(p.add_run(text), size=size, bold=bold, italic=italic, lang=lang)
    return p


def add_labeled(doc, label: str, text: str, *, lang="vi-VN", after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    set_run(p.add_run(label), bold=True, lang=lang)
    set_run(p.add_run(text), lang=lang)
    return p


def add_heading(doc, text: str, level: int, *, lang="vi-VN"):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = level == 1 and text.startswith(("PHẦN B", "B 部分"))
    r = p.add_run(text)
    set_run(
        r,
        size={1: 15, 2: 12.5, 3: 11}[level],
        bold=True,
        color=BLACK if level < 3 else DARK,
        lang=lang,
    )
    return p


def _next_num_id(numbering) -> int:
    ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    return max(ids, default=0) + 1


def _next_abstract_id(numbering) -> int:
    ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    return max(ids, default=0) + 1


def create_bullet_numbering(doc) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = _next_abstract_id(numbering)
    num_id = _next_num_id(numbering)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), "A1B2C3D4")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "260")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, text: str, num_id: int, *, lang="vi-VN", bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True, lang=lang)
        set_run(p.add_run(text[len(bold_prefix):]), lang=lang)
    else:
        set_run(p.add_run(text), lang=lang)
    return p


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], *, lang="vi-VN", font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    hdr_pr = hdr._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    hdr_pr.append(repeat)
    for index, text in enumerate(headers):
        set_cell_shading(hdr.cells[index], LIGHT)
        p = hdr.cells[index].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_run(p.add_run(text), size=font_size, bold=True, lang=lang)
    for values in rows:
        row = table.add_row()
        for index, text in enumerate(values):
            p = row.cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run(p.add_run(text), size=font_size, lang=lang)
    set_table_geometry(table, widths)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_code_block(doc, lines: list[str], *, lang="vi-VN"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run(run, size=8.7, lang=lang)
        run.font.name = "Consolas"
        r_pr = run._element.get_or_add_rPr()
        r_pr.rFonts.set(qn("w:ascii"), "Consolas")
        r_pr.rFonts.set(qn("w:hAnsi"), "Consolas")
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    heading_specs = {
        "Heading 1": (15, 14, 7),
        "Heading 2": (12.5, 10, 5),
        "Heading 3": (11, 7, 3),
    }
    for style_name, (size, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(
        hp.add_run("AIVO V1 | ODM Integration Requirements / ODM 集成要求"),
        size=8.5,
        color=MUTED,
        lang="zh-CN",
    )
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(fp.add_run("AIVO - Technical draft / 技术草案  |  "), size=8, color=MUTED, lang="zh-CN")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    fp._p.append(field)


def add_title_page(doc: Document) -> None:
    for _ in range(4):
        add_text(doc, "", after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("AIVO V1"), size=24, bold=True, lang="vi-VN")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run("YÊU CẦU TÍCH HỢP PHẦN CỨNG VÀ GIAO THỨC APP"), size=17, bold=True, lang="vi-VN")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    set_run(p.add_run("硬件与应用程序集成要求"), size=17, bold=True, lang="zh-CN")

    add_table(
        doc,
        ["Thông tin / 项目", "Nội dung / 内容"],
        [
            ["Đối tượng / 接收方", "Nhà sản xuất ODM mới / 新 ODM 制造商"],
            ["Phạm vi / 范围", "Android và iOS Native / Android 与 iOS 原生应用"],
            ["Trạng thái / 状态", "Dự thảo kỹ thuật - chờ ODM xác nhận / 技术草案 - 待 ODM 确认"],
            ["Phiên bản / 版本", "V0.2 - 03/09/2026"],
        ],
        [2600, 7260],
        lang="zh-CN",
        font_size=10,
    )
    add_text(
        doc,
        "Lưu ý / 注意：Tài liệu này cung cấp trực tiếp hợp đồng giao tiếp APP đề xuất. Các khả năng HFP, BLE và OTA vẫn phải được ODM xác nhận bằng firmware và mẫu thật trước khi chốt sản xuất. / 本文件已提供应用侧建议通信协议。HFP、BLE 与 OTA 能力仍须由 ODM 通过固件和实物样机确认后方可定版生产。",
        bold=True,
        size=10,
        after=4,
        lang="zh-CN",
    )
    doc.add_page_break()


def add_vietnamese(doc: Document, bullet_num: int) -> None:
    add_heading(doc, "PHẦN A - BẢN TIẾNG VIỆT", 1)
    add_heading(doc, "1. Kết luận từ hai tài liệu ODM cung cấp", 2)
    add_table(
        doc,
        ["Tài liệu", "Thông tin sử dụng được", "Giới hạn cần làm rõ"],
        [
            [
                "BT8952U Audio Player Microcontroller - V0.0.5",
                "Có RF 2,4 GHz proprietary 2 Mbps; ADC/DAC 24-bit; hai đầu vào micro; ENC/ANC; VAD/KWS; GPIO; bộ sạc; Flash 16 Mbit; RAM 892 KB; công suất phát tại chip điển hình 9 dBm, tối đa 11 dBm.",
                "Không ghi rõ Bluetooth Classic, HFP, BLE GATT, A2DP/AVRCP hoặc OTA. Không được dùng tài liệu này làm bằng chứng rằng điện thoại có thể kết nối HFP + BLE.",
            ],
            [
                "天地盖.pdf - bản vẽ hộp nắp/đáy",
                "Nắp trên khoảng 120 x 90 x 18 mm; đáy dưới khoảng 116 x 86 x 26 mm.",
                "Chưa có vật liệu, độ dày, dung sai, khay giữ thiết bị, phụ kiện, vùng in và đường bleed.",
            ],
        ],
        [2350, 3900, 3610],
        font_size=8.8,
    )
    add_text(
        doc,
        "Điểm cần xác nhận đầu tiên: nếu BT8952U là chip RF duy nhất trên bo mạch, ODM phải cung cấp bằng chứng chính thức về HFP + BLE. Nếu có thêm chip/module Bluetooth, ODM phải cung cấp mã chip, datasheet và sơ đồ kết nối giữa hai chip.",
        bold=True,
        size=10,
    )

    add_heading(doc, "2. Kiến trúc tích hợp bắt buộc", 2)
    for text in [
        "Bluetooth Classic HFP hai chiều: micro thiết bị truyền tiếng trẻ vào APP; âm thanh từ APP phát qua loa thiết bị.",
        "BLE chỉ truyền nút bấm, pin, firmware và trạng thái APP. Không truyền PCM/Opus hoặc âm thanh thời gian thực qua BLE.",
        "HFP và BLE phải hoạt động đồng thời, không làm mất Indication khi mở hoặc đóng đường âm thanh HFP/SCO.",
        "Thiết bị không tự nhận dạng, dịch, chấm điểm hoặc lựa chọn nội dung. APP chịu trách nhiệm toàn bộ logic học và AI.",
        "Áp dụng cho Android và iOS Native. iOS Web/Safari không thuộc phạm vi điều khiển BLE.",
    ]:
        add_bullet(doc, text, bullet_num)

    add_heading(doc, "3. Chức năng nút đã chốt", 2)
    add_table(
        doc,
        ["Nút", "Thao tác", "Kết quả"],
        [
            ["Giảm âm lượng", "Bấm ngắn", "Giảm âm lượng cục bộ"],
            ["Giảm âm lượng", "Bấm giữ", "APP chuyển sang câu tiếp theo trong bài học"],
            ["Tăng âm lượng", "Bấm ngắn", "Tăng âm lượng cục bộ"],
            ["Tăng âm lượng", "Bấm giữ", "APP quay về câu trước trong bài học"],
            ["Nguồn", "Bấm ngắn", "Luyện lại câu hiện tại, chỉ trong chủ đề luyện nghe"],
            ["Nguồn", "Bấm giữ 2-3 giây", "Bật hoặc tắt thiết bị; xử lý cục bộ"],
            ["MAIN", "Bấm ngắn", "Gọi hoặc tiếp tục trợ lý AI"],
            ["MAIN", "Bấm giữ", "Tạm dừng toàn bộ hoạt động hiện tại của APP"],
        ],
        [2200, 2100, 5560],
        font_size=9.1,
    )
    add_text(
        doc,
        "Firmware phải phát đúng một sự kiện cho mỗi thao tác. Sự kiện LONG không được phát thêm SHORT. Các lệnh POWER_SHORT và MAIN phải là sự kiện BLE riêng, không được biến thành Siri, Google Assistant hoặc media key của hệ điều hành.",
        bold=True,
        size=9.8,
    )

    doc.add_page_break()
    add_heading(doc, "4. Giao thức BLE do phía APP cung cấp", 2)
    add_heading(doc, "4.1 UUID và thuộc tính", 3)
    add_table(
        doc,
        ["Đối tượng", "UUID", "Thuộc tính yêu cầu"],
        [
            ["Control Service", "9E3B0001-4A7C-4D6F-8B21-5C17A2D94010", "Primary Service; advertise UUID"],
            ["Button Event", "9E3B0002-4A7C-4D6F-8B21-5C17A2D94010", "Indicate bắt buộc; CCCD 2902"],
            ["APP State", "9E3B0003-4A7C-4D6F-8B21-5C17A2D94010", "Write With Response bắt buộc"],
            ["Battery", "180F / 2A19", "Read + Notify; giá trị 0-100"],
            ["Firmware Revision", "180A / 2A26", "Read; UTF-8"],
        ],
        [2200, 4750, 2910],
        font_size=8.6,
    )
    add_heading(doc, "4.2 Button Event - 12 byte, Little Endian", 3)
    add_table(
        doc,
        ["Byte", "Trường", "Giá trị/quy tắc"],
        [
            ["0", "Protocol version", "0x01"],
            ["1", "Button ID", "0x01 MAIN; 0x02 POWER; 0x03 VOLUME_UP; 0x04 VOLUME_DOWN"],
            ["2", "Gesture", "0x01 SHORT; 0x02 LONG; 0x03 RELEASE (tùy chọn)"],
            ["3", "Flags", "0x00 ở V1; dành cho mở rộng"],
            ["4-5", "Sequence", "uint16 LE; tăng sau mỗi sự kiện logic"],
            ["6", "Battery", "0-100; 0xFF nếu chưa biết"],
            ["7", "Reserved", "0x00"],
            ["8-11", "Uptime", "uint32 LE, mili giây"],
        ],
        [1200, 2450, 6210],
        font_size=8.9,
    )
    add_heading(doc, "4.3 Mã trạng thái APP - APP State 8 byte", 3)
    add_table(
        doc,
        ["Byte", "Trường", "Giá trị/quy tắc"],
        [
            ["0", "Protocol version", "0x01"],
            ["1", "APP State", "00 IDLE; 01 RECORDING; 02 PROCESSING; 03 READY; 04 PLAYING; 05 ERROR; 06 PAUSED"],
            ["2", "Result/ACK", "00 ACCEPTED; 01 BUSY; 02 NO_RESULT; 03 MIC_UNAVAILABLE; 04 BT_ROUTE_UNAVAILABLE; 05 DUPLICATE; 06 INTERNAL_ERROR"],
            ["3", "Flags", "0x00 ở V1"],
            ["4-5", "Related sequence", "Sequence của Button Event được phản hồi, uint16 LE"],
            ["6-7", "Reserved", "0x00 0x00"],
        ],
        [1200, 2450, 6210],
        font_size=8.7,
    )
    add_heading(doc, "4.4 Raw Hex mẫu do APP cung cấp", 3)
    add_code_block(
        doc,
        [
            "MAIN_SHORT        01 01 01 00 01 00 64 00 10 27 00 00",
            "MAIN_LONG         01 01 02 00 02 00 64 00 20 4E 00 00",
            "POWER_SHORT       01 02 01 00 03 00 64 00 30 75 00 00",
            "VOLUME_UP_LONG    01 03 02 00 04 00 64 00 40 9C 00 00",
            "VOLUME_DOWN_LONG  01 04 02 00 05 00 64 00 50 C3 00 00",
            "APP RECORDING ACK 01 01 00 00 01 00 00 00",
            "APP PAUSED ACK    01 06 00 00 02 00 00 00",
        ],
    )
    add_text(doc, "Các gói trên là mẫu chuẩn theo hợp đồng đề xuất. ODM phải trả lại Raw Hex thực tế từ firmware mẫu để hai bên đối chiếu trước khi khóa giao thức.", italic=True, size=9.2)

    add_heading(doc, "5. Quy tắc xử lý của APP đã cung cấp", 2)
    add_labeled(doc, "POWER_SHORT trong luyện nghe: ", "Giữ nguyên câu hiện tại; hủy lượt ghi/chấm đang chạy; bỏ qua kết quả cũ trả muộn; phát lại câu tiếng Anh; phát hướng dẫn tiếng Việt nếu bài yêu cầu; mở lại micro và chấm lượt mới. Ngoài luyện nghe, APP bỏ qua sự kiện.")
    add_labeled(doc, "MAIN_LONG: ", "Nếu đang ghi thì dừng và hủy bản ghi; nếu đang xử lý thì hủy hoặc bỏ qua kết quả; nếu đang phát thì dừng phát; giữ nguyên chủ đề và vị trí câu; chuyển trạng thái PAUSED. MAIN_SHORT dùng để tiếp tục.")
    add_labeled(doc, "NEXT/PREVIOUS: ", "Chỉ APP thay đổi câu. Thiết bị không tự đổi nội dung. Ngoài màn hình bài học, APP có thể bỏ qua sự kiện.")
    add_labeled(doc, "Chống trùng: ", "APP dùng Sequence và cửa sổ thời gian để bỏ sự kiện trùng. ODM vẫn phải chống rung nút tại firmware.")

    add_heading(doc, "6. Thông tin và xác nhận cần từ ODM", 2)
    for text in [
        "Xác nhận BT8952U có trực tiếp hỗ trợ Bluetooth Classic HFP + BLE hay có thêm chip/module Bluetooth; cung cấp mã chip, datasheet và sơ đồ khối.",
        "Cung cấp phiên bản HFP, codec CVSD/mSBC, tên Bluetooth Classic và tên BLE advertise; xác nhận một thiết bị có thể duy trì HFP và BLE đồng thời.",
        "Xác nhận UUID/packet ở Mục 4, GATT dump và Raw Hex thực tế của MAIN_SHORT, MAIN_LONG, POWER_SHORT, VOLUME_UP_LONG và VOLUME_DOWN_LONG.",
        "Mỗi thao tác chỉ tạo một sự kiện; LONG không kèm SHORT; mô tả ngưỡng thời gian SHORT/LONG và debounce.",
        "Cung cấp dung lượng pin, thời gian sử dụng/sạc, khả năng vừa bật vừa sạc, cảnh báo pin dưới 20%, Battery Notify và trạng thái sạc nếu có.",
        "Xác nhận chính sách ngủ sau 30 phút, thao tác đánh thức và tự kết nối lại HFP + BLE với điện thoại đã ghép.",
        "Cung cấp công suất RF/EIRP thực tế của sản phẩm hoàn chỉnh, loại anten và báo cáo thử nghiệm; không dùng 9-11 dBm tại chân chip thay cho EIRP.",
        "Cho biết ENC/ANC nào đã được kích hoạt và hiệu chỉnh trên mẫu thật; cung cấp điều kiện kiểm thử micro/loa.",
        "Xác nhận vật liệu, độ dày, dung sai, kích thước trong/ngoài, khay giữ, phụ kiện và vùng in của hộp nắp/đáy.",
    ]:
        add_bullet(doc, text, bullet_num)

    add_heading(doc, "7. Nội dung phía APP cung cấp trong tài liệu này", 2)
    for text in [
        "Kiến trúc HFP audio + BLE control dùng chung cho Android và iOS Native.",
        "Bộ UUID đề xuất, thuộc tính Characteristic, packet Button Event 12 byte và APP State 8 byte.",
        "Mã Button ID, Gesture, APP State, Result/ACK và Raw Hex mẫu.",
        "Quy tắc xử lý POWER_SHORT, MAIN_LONG, chuyển câu và chống sự kiện trùng.",
        "Checklist nghiệm thu ở Mục 9 để ODM tự kiểm tra firmware trước khi gửi mẫu.",
        "Sau khi ODM xác nhận giao thức và gửi firmware mẫu: phía APP sẽ cập nhật parser cho bốn nút, bổ sung PAUSED, cung cấp APK Android và TestFlight iOS có màn hình chẩn đoán HFP/BLE/Raw Hex.",
    ]:
        add_bullet(doc, text, bullet_num)

    add_heading(doc, "8. Yêu cầu OTA", 2)
    for text in [
        "Cung cấp phương thức OTA, tài liệu giao thức/SDK Android và iOS, định dạng file firmware và quy tắc kiểm tra phiên bản.",
        "Firmware phải có CRC/SHA và xác minh chữ ký; không chấp nhận file hỏng hoặc sai model.",
        "Hỗ trợ resume sau mất kết nối và dual-bank/rollback để thiết bị không bị brick.",
        "Quy định pin tối thiểu, hành vi khi đang sạc, thời gian cập nhật và trạng thái tiến trình/lỗi trả về APP.",
        "Cung cấp factory reset, chế độ bootloader/recovery và quy trình nạp firmware tại nhà máy.",
        "Nếu OTA chưa sẵn sàng, ODM phải ghi rõ trạng thái chưa hỗ trợ và kế hoạch bổ sung trước sản xuất hàng loạt.",
    ]:
        add_bullet(doc, text, bullet_num)

    doc.add_page_break()
    add_heading(doc, "9. Checklist nghiệm thu mẫu", 2)
    add_table(
        doc,
        ["Hạng mục", "Điều kiện đạt", "Kết quả ODM"],
        [
            ["HFP micro", "Android và iPhone xác nhận micro thiết bị là nguồn thu thực tế", "Đạt / Không đạt"],
            ["HFP loa", "Âm thanh APP phát đúng qua loa thiết bị", "Đạt / Không đạt"],
            ["BLE đồng thời", "Mở/đóng HFP không làm mất Button Indication", "Đạt / Không đạt"],
            ["Nút", "5 sự kiện mới đúng một lần, đúng Raw Hex, không trùng", "Đạt / Không đạt"],
            ["APP State", "Write With Response thành nhận ACK và ghi thành công", "Đạt / Không đạt"],
            ["Reconnect", "Tắt/bật và ra/vào vùng phủ sóng đều tự kết nối lại", "Đạt / Không đạt"],
            ["Pin/sạc", "Notify chính xác; cảnh báo <20%; xác nhận vừa dùng vừa sạc", "Đạt / Không đạt"],
            ["OTA", "Mất kết nối giữa chừng vẫn resume/rollback, thiết bị không brick", "Đạt / Không đạt"],
        ],
        [2100, 5650, 2110],
        font_size=8.8,
    )
    add_heading(doc, "10. Hồ sơ ODM gửi lại cùng mẫu", 2)
    for text in [
        "Firmware release note, mã PCB, mã chip/module Bluetooth và số serial của mẫu.",
        "GATT dump, Raw Hex của năm sự kiện nút và kết quả checklist đã điền.",
        "Kết quả HFP/BLE trên ít nhất một Android và một iPhone; trạng thái OTA và RF/EIRP.",
        "Người phụ trách kỹ thuật, ngày xác nhận giao thức và ngày dự kiến gửi mẫu.",
    ]:
        add_bullet(doc, text, bullet_num)
    add_table(
        doc,
        ["Xác nhận ODM", "Thông tin điền"],
        [
            ["Tên công ty/người phụ trách", ""],
            ["Firmware/PCB version", ""],
            ["Ngày xác nhận/chữ ký", ""],
        ],
        [3100, 6760],
        font_size=9,
    )


def add_chinese(
    doc: Document,
    bullet_num: int,
    *,
    include_part_heading: bool = True,
    include_page_break: bool = True,
) -> None:
    if include_page_break:
        doc.add_page_break()
    if include_part_heading:
        add_heading(doc, "B 部分 - 简体中文版", 1, lang="zh-CN")
    add_heading(doc, "1. 对 ODM 已提供文件的结论", 2, lang="zh-CN")
    add_table(
        doc,
        ["文件", "可采用的信息", "必须澄清的限制"],
        [
            [
                "BT8952U Audio Player Microcontroller - V0.0.5",
                "包含 2.4 GHz 私有 2 Mbps RF、24 位 ADC/DAC、双麦克风输入、ENC/ANC、VAD/KWS、GPIO、充电模块、16 Mbit Flash、892 KB RAM；芯片端发射功率典型值 9 dBm、最大值 11 dBm。",
                "文件未明确说明 Bluetooth Classic、HFP、BLE GATT、A2DP/AVRCP 或 OTA，不能据此证明手机能够通过 HFP + BLE 连接。",
            ],
            [
                "天地盖.pdf - 天地盖包装刀模图",
                "上盖约 120 x 90 x 18 mm；下盖约 116 x 86 x 26 mm。",
                "缺少材料、厚度、公差、内托、配件位置、印刷区域及出血线。",
            ],
        ],
        [2350, 3900, 3610],
        lang="zh-CN",
        font_size=8.8,
    )
    add_text(
        doc,
        "首要确认事项：如果 BT8952U 是 PCB 上唯一的无线芯片，ODM 必须提供其支持 HFP + BLE 的正式证明；如果另有蓝牙芯片或模组，须提供具体型号、数据手册及两个芯片之间的连接框图。",
        bold=True,
        size=10,
        lang="zh-CN",
    )

    add_heading(doc, "2. 必须采用的集成架构", 2, lang="zh-CN")
    for text in [
        "Bluetooth Classic HFP 双向音频：设备麦克风将儿童语音传给 APP，APP 音频通过设备扬声器播放。",
        "BLE 仅传输按键、剩余电量、固件版本及 APP 状态，不通过 BLE 传输 PCM、Opus 或实时音频。",
        "HFP 与 BLE 必须同时稳定工作；开启或关闭 HFP/SCO 音频链路时不得丢失按键 Indication。",
        "设备不自行识别、翻译、评分或选择播放内容，全部学习与 AI 逻辑由 APP 决定。",
        "适用于 Android 与 iOS 原生应用；iOS Web/Safari 不在 BLE 按键控制范围内。",
    ]:
        add_bullet(doc, text, bullet_num, lang="zh-CN")

    add_heading(doc, "3. 已确认的按键功能", 2, lang="zh-CN")
    add_table(
        doc,
        ["按键", "操作", "结果"],
        [
            ["音量减", "短按", "设备本地降低音量"],
            ["音量减", "长按", "APP 切换到课程下一句"],
            ["音量加", "短按", "设备本地提高音量"],
            ["音量加", "长按", "APP 返回课程上一句"],
            ["电源", "短按", "仅在听力练习中重新练习当前句"],
            ["电源", "长按 2-3 秒", "设备本地开机或关机"],
            ["MAIN", "短按", "呼叫或继续 AI 助手"],
            ["MAIN", "长按", "暂停 APP 当前全部活动"],
        ],
        [2200, 2100, 5560],
        lang="zh-CN",
        font_size=9.1,
    )
    add_text(
        doc,
        "每次物理操作只能产生一个逻辑事件。长按不得再产生短按事件。POWER_SHORT 与 MAIN 必须作为独立 BLE 事件发送，不得映射为 Siri、Google Assistant 或手机系统媒体键。",
        bold=True,
        size=9.8,
        lang="zh-CN",
    )

    doc.add_page_break()
    add_heading(doc, "4. APP 方提供的 BLE 通信协议", 2, lang="zh-CN")
    add_heading(doc, "4.1 UUID 与属性", 3, lang="zh-CN")
    add_table(
        doc,
        ["对象", "UUID", "要求的属性"],
        [
            ["Control Service", "9E3B0001-4A7C-4D6F-8B21-5C17A2D94010", "Primary Service；广播该 UUID"],
            ["Button Event", "9E3B0002-4A7C-4D6F-8B21-5C17A2D94010", "必须支持 Indicate；CCCD 2902"],
            ["APP State", "9E3B0003-4A7C-4D6F-8B21-5C17A2D94010", "必须支持 Write With Response"],
            ["Battery", "180F / 2A19", "Read + Notify；范围 0-100"],
            ["Firmware Revision", "180A / 2A26", "Read；UTF-8"],
        ],
        [2200, 4750, 2910],
        lang="zh-CN",
        font_size=8.6,
    )
    add_heading(doc, "4.2 Button Event - 12 字节，Little Endian", 3, lang="zh-CN")
    add_table(
        doc,
        ["字节", "字段", "数值/规则"],
        [
            ["0", "Protocol version", "0x01"],
            ["1", "Button ID", "0x01 MAIN；0x02 POWER；0x03 VOLUME_UP；0x04 VOLUME_DOWN"],
            ["2", "Gesture", "0x01 SHORT；0x02 LONG；0x03 RELEASE（可选）"],
            ["3", "Flags", "V1 使用 0x00，保留扩展"],
            ["4-5", "Sequence", "uint16 LE；每个逻辑事件递增"],
            ["6", "Battery", "0-100；未知时为 0xFF"],
            ["7", "Reserved", "0x00"],
            ["8-11", "Uptime", "uint32 LE，单位毫秒"],
        ],
        [1200, 2450, 6210],
        lang="zh-CN",
        font_size=8.9,
    )
    add_heading(doc, "4.3 APP State - 8 字节", 3, lang="zh-CN")
    add_table(
        doc,
        ["字节", "字段", "数值/规则"],
        [
            ["0", "Protocol version", "0x01"],
            ["1", "APP State", "00 IDLE；01 RECORDING；02 PROCESSING；03 READY；04 PLAYING；05 ERROR；06 PAUSED"],
            ["2", "Result/ACK", "00 ACCEPTED；01 BUSY；02 NO_RESULT；03 MIC_UNAVAILABLE；04 BT_ROUTE_UNAVAILABLE；05 DUPLICATE；06 INTERNAL_ERROR"],
            ["3", "Flags", "V1 使用 0x00"],
            ["4-5", "Related sequence", "所响应的 Button Event Sequence，uint16 LE"],
            ["6-7", "Reserved", "0x00 0x00"],
        ],
        [1200, 2450, 6210],
        lang="zh-CN",
        font_size=8.7,
    )
    add_heading(doc, "4.4 APP 方提供的 Raw Hex 示例", 3, lang="zh-CN")
    add_code_block(
        doc,
        [
            "MAIN_SHORT        01 01 01 00 01 00 64 00 10 27 00 00",
            "MAIN_LONG         01 01 02 00 02 00 64 00 20 4E 00 00",
            "POWER_SHORT       01 02 01 00 03 00 64 00 30 75 00 00",
            "VOLUME_UP_LONG    01 03 02 00 04 00 64 00 40 9C 00 00",
            "VOLUME_DOWN_LONG  01 04 02 00 05 00 64 00 50 C3 00 00",
            "APP RECORDING ACK 01 01 00 00 01 00 00 00",
            "APP PAUSED ACK    01 06 00 00 02 00 00 00",
        ],
        lang="zh-CN",
    )
    add_text(doc, "以上数据包为建议通信协议的标准示例。ODM 必须提供样机固件实际产生的 Raw Hex，由双方核对后再锁定协议。", italic=True, size=9.2, lang="zh-CN")

    add_heading(doc, "5. APP 已提供的处理规则", 2, lang="zh-CN")
    add_labeled(doc, "听力练习中的 POWER_SHORT：", "保持当前句；取消正在进行的录音或评分；忽略延迟返回的旧结果；重新播放当前英文句；如课程要求则播放越南语提示；重新打开麦克风并重新评分。非听力练习页面忽略该事件。", lang="zh-CN")
    add_labeled(doc, "MAIN_LONG：", "录音中则停止并丢弃当前录音；处理中则取消请求或忽略返回结果；播放中则立即停止；保留当前主题与句子位置；APP 进入 PAUSED。MAIN_SHORT 用于继续。", lang="zh-CN")
    add_labeled(doc, "NEXT/PREVIOUS：", "仅由 APP 切换句子，设备不得自行改变内容；不在课程页面时 APP 可以忽略。", lang="zh-CN")
    add_labeled(doc, "去重：", "APP 使用 Sequence 与时间窗口过滤重复事件；ODM 固件仍须完成按键消抖。", lang="zh-CN")

    add_heading(doc, "6. ODM 必须提供或确认的信息", 2, lang="zh-CN")
    for text in [
        "确认 BT8952U 是否直接支持 Bluetooth Classic HFP + BLE，或另有蓝牙芯片/模组；提供芯片型号、数据手册与系统框图。",
        "提供 HFP 版本、CVSD/mSBC 编解码器、Bluetooth Classic 名称及 BLE 广播名称；确认 HFP 与 BLE 可同时保持连接。",
        "确认第 4 节 UUID/数据包，提供完整 GATT dump，以及 MAIN_SHORT、MAIN_LONG、POWER_SHORT、VOLUME_UP_LONG、VOLUME_DOWN_LONG 的真实 Raw Hex。",
        "每个操作只产生一个事件；LONG 不得附带 SHORT；说明 SHORT/LONG 时间阈值及消抖规则。",
        "提供电池容量、续航、充电时间、开机充电能力、低于 20% 提示、Battery Notify 及充电状态能力。",
        "确认 30 分钟无操作后的休眠策略、唤醒方式，以及与已配对手机自动恢复 HFP + BLE 的行为。",
        "提供整机实际 RF/EIRP、天线类型与测试报告；不得以芯片引脚端 9-11 dBm 代替整机 EIRP。",
        "说明样机实际启用及调试的 ENC/ANC 功能，并提供麦克风/扬声器测试条件。",
        "确认天地盖包装的材料、厚度、公差、内外尺寸、内托、配件位置及印刷区域。",
    ]:
        add_bullet(doc, text, bullet_num, lang="zh-CN")

    add_heading(doc, "7. 本文件已由 APP 方提供的内容", 2, lang="zh-CN")
    for text in [
        "Android 与 iOS 原生应用共用的 HFP Audio + BLE Control 架构。",
        "建议 UUID、Characteristic 属性、12 字节 Button Event 与 8 字节 APP State 格式。",
        "Button ID、Gesture、APP State、Result/ACK 编码及 Raw Hex 示例。",
        "POWER_SHORT、MAIN_LONG、切换句子及重复事件过滤规则。",
        "第 9 节的样机验收清单，供 ODM 发样前完成固件自测。",
        "ODM 确认协议并提供样机固件后：APP 方将扩展四个按键解析、增加 PAUSED 状态，并提供带 HFP/BLE/Raw Hex 诊断页面的 Android APK 与 iOS TestFlight。",
    ]:
        add_bullet(doc, text, bullet_num, lang="zh-CN")

    add_heading(doc, "8. OTA 要求", 2, lang="zh-CN")
    for text in [
        "提供 OTA 方式、Android/iOS 协议或 SDK、固件文件格式及版本检查规则。",
        "固件须具备 CRC/SHA 与签名验证，拒绝损坏文件或错误型号固件。",
        "支持断点续传以及双分区或回滚机制，避免升级失败导致设备变砖。",
        "说明最低电量、充电状态限制、升级时间，以及向 APP 返回的进度和错误状态。",
        "提供恢复出厂设置、Bootloader/Recovery 模式及工厂烧录流程。",
        "如果 OTA 尚未完成，ODM 必须明确标记为未支持，并提供量产前的补充计划。",
    ]:
        add_bullet(doc, text, bullet_num, lang="zh-CN")

    add_heading(doc, "9. 样机验收清单", 2, lang="zh-CN")
    add_table(
        doc,
        ["项目", "通过条件", "ODM 结果"],
        [
            ["HFP 麦克风", "Android 与 iPhone 均确认设备麦克风为实际录音源", "通过 / 不通过"],
            ["HFP 扬声器", "APP 音频正确从设备扬声器播放", "通过 / 不通过"],
            ["BLE 并行", "开启/关闭 HFP 时不丢失 Button Indication", "通过 / 不通过"],
            ["按键", "5 个新事件各触发一次，Raw Hex 正确且无重复", "通过 / 不通过"],
            ["APP State", "Write With Response 成功并收到确认", "通过 / 不通过"],
            ["自动重连", "开关机或离开/返回范围后自动恢复连接", "通过 / 不通过"],
            ["电池/充电", "Notify 正确；低于 20% 提示；确认可开机充电", "通过 / 不通过"],
            ["OTA", "升级中断后可续传/回滚，设备不会变砖", "通过 / 不通过"],
        ],
        [2100, 5650, 2110],
        lang="zh-CN",
        font_size=8.8,
    )
    add_heading(doc, "10. ODM 随样机返还的资料", 2, lang="zh-CN")
    for text in [
        "固件 Release Note、PCB 编号、蓝牙芯片/模组型号及样机序列号。",
        "完整 GATT dump、五个按键事件的真实 Raw Hex，以及已填写的验收清单。",
        "至少一台 Android 与一台 iPhone 的 HFP/BLE 测试结果，以及 OTA 与 RF/EIRP 状态。",
        "技术负责人、协议确认日期及预计寄样日期。",
    ]:
        add_bullet(doc, text, bullet_num, lang="zh-CN")
    add_table(
        doc,
        ["ODM 确认", "填写内容"],
        [
            ["公司/技术负责人", ""],
            ["Firmware/PCB version", ""],
            ["协议确认日期/签字", ""],
        ],
        [3100, 6760],
        lang="zh-CN",
        font_size=9,
    )


def audit_docx(path: Path) -> None:
    check = Document(path)
    assert len(check.sections) == 1
    assert len(check.tables) >= 10
    all_text = "\n".join(p.text for p in check.paragraphs)
    required = [
        "YÊU CẦU TÍCH HỢP PHẦN CỨNG",
        "硬件与应用程序集成要求",
        "POWER_SHORT",
        "MAIN_LONG",
        "9E3B0001-4A7C-4D6F-8B21-5C17A2D94010",
        "BT8952U",
        "APP PAUSED ACK",
        "OTA 要求",
    ]
    table_text = "\n".join(cell.text for table in check.tables for row in table.rows for cell in row.cells)
    combined = all_text + "\n" + table_text
    for token in required:
        assert token in combined, token


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = "AIVO V1 - Yêu cầu tích hợp ODM Việt-Trung"
    doc.core_properties.subject = "HFP, BLE Control, button protocol, OTA and packaging requirements"
    doc.core_properties.author = "AIVO"
    doc.core_properties.keywords = "AIVO, ODM, HFP, BLE, OTA, Android, iOS"
    style_document(doc)
    bullet_num = create_bullet_numbering(doc)
    add_title_page(doc)
    add_vietnamese(doc, bullet_num)
    add_chinese(doc, bullet_num)
    doc.save(OUTPUT)
    audit_docx(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
