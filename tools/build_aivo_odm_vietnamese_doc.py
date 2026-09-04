from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

import build_aivo_odm_bilingual_doc as base


OUTPUT = Path(__file__).resolve().parents[1] / "output" / "docs" / "AIVO_V1_YEU_CAU_TICH_HOP_ODM_VI.docx"


def set_vietnamese_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = True
    section.different_first_page_header_footer = True

    for header_part in [section.header, section.even_page_header, section.first_page_header]:
        header = header_part.paragraphs[0]
        header.clear()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        base.set_run(
            header.add_run("AIVO V1 | Yêu cầu tích hợp ODM"),
            size=8.5,
            color=base.MUTED,
            lang="vi-VN",
        )

    for footer_part in [section.footer, section.even_page_footer, section.first_page_footer]:
        footer = footer_part.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        base.set_run(
            footer.add_run("AIVO - Dự thảo kỹ thuật  |  "),
            size=8,
            color=base.MUTED,
            lang="vi-VN",
        )
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)


def add_title_page(doc: Document) -> None:
    for _ in range(4):
        base.add_text(doc, "", after=8, lang="vi-VN")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    base.set_run(p.add_run("AIVO V1"), size=24, bold=True, lang="vi-VN")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    base.set_run(
        p.add_run("YÊU CẦU TÍCH HỢP PHẦN CỨNG VÀ GIAO THỨC APP"),
        size=17,
        bold=True,
        lang="vi-VN",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    base.set_run(
        p.add_run("Dùng cho ODM làm mẫu phát triển firmware và phối hợp kiểm thử"),
        size=12.5,
        lang="vi-VN",
    )

    base.add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ["Đối tượng", "Nhà sản xuất ODM mới"],
            ["Phạm vi", "Android và iOS Native"],
            ["Trạng thái", "Dự thảo kỹ thuật - chờ ODM xác nhận"],
            ["Phiên bản", "V0.2 - 04/09/2026"],
        ],
        [2600, 7260],
        lang="vi-VN",
        font_size=10,
    )
    base.add_text(
        doc,
        "Lưu ý: Tài liệu đã cung cấp giao thức đề xuất từ phía APP. Các khả năng HFP, BLE và OTA vẫn phải được ODM xác nhận bằng firmware và mẫu thật trước khi chốt sản xuất.",
        bold=True,
        size=10,
        after=4,
        lang="vi-VN",
    )
    doc.add_page_break()


def audit_docx(path: Path) -> None:
    check = Document(path)
    text = "\n".join(p.text for p in check.paragraphs)
    text += "\n" + "\n".join(
        cell.text for table in check.tables for row in table.rows for cell in row.cells
    )
    required = [
        "YÊU CẦU TÍCH HỢP PHẦN CỨNG VÀ GIAO THỨC APP",
        "7. Nội dung phía APP cung cấp trong tài liệu này",
        "9E3B0001-4A7C-4D6F-8B21-5C17A2D94010",
        "MAIN_LONG",
        "APP PAUSED ACK",
        "Yêu cầu OTA",
    ]
    for token in required:
        assert token in text, token
    assert len(check.tables) >= 8


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = "AIVO V1 Yêu cầu tích hợp phần cứng và giao thức APP"
    doc.core_properties.subject = "HFP BLE Control giao thức nút OTA và yêu cầu đóng gói"
    doc.core_properties.author = "AIVO"
    doc.core_properties.keywords = "AIVO, ODM, HFP, BLE, OTA, Android, iOS"
    base.style_document(doc)
    set_vietnamese_header_footer(doc)
    bullet_num = base.create_bullet_numbering(doc)
    add_title_page(doc)
    base.add_vietnamese(doc, bullet_num, include_part_heading=False)
    compact_section = False
    for paragraph in doc.paragraphs:
        if paragraph.text in {
            "2. Kiến trúc tích hợp bắt buộc",
            "6. Thông tin và xác nhận cần từ ODM",
        }:
            compact_section = True
            continue
        if paragraph.text in {
            "3. Chức năng nút đã chốt",
            "7. Nội dung phía APP cung cấp trong tài liệu này",
        }:
            compact_section = False
        if compact_section and paragraph.text.strip():
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.size = Pt(9.8)
        if paragraph.text.startswith("Firmware phải phát đúng một sự kiện"):
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.size = Pt(9.3)
    doc.save(OUTPUT)
    audit_docx(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
