from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "output" / "docx" / "Ghi_chu_chuyen_PWA_sang_iOS_native_App_Store.docx"

NAVY = "153B6B"
BLUE = "2E74B5"
LIGHT_BLUE = "E8F1FA"
LIGHT_GRAY = "F2F4F7"
MUTED = "5E6B78"
GREEN = "1B7F4C"


def set_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_font(first, bold=True, color=NAVY)
        rest = p.add_run(text[len(bold_prefix):])
        set_font(rest)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_matrix(doc):
    rows = [
        ("Giao diện Flutter", "GIỮ", "Giữ màn hình, mascot, chủ đề, bài học, từ vựng và lịch sử."),
        ("Logic ứng dụng", "GIỮ", "Giữ điều hướng AI, trạng thái bài học, xử lý câu tiếp/câu trước/dừng."),
        ("Nút MAIN", "GIỮ PHẦN LOGIC", "Giữ MainButtonCoordinator; chỉ thay nguồn sự kiện từ MAIN ảo sang BLE iOS."),
        ("Backend và AI", "GIỮ", "Giữ API Railway, Cloudflare ASR/TTS, rule/cache và dữ liệu hiện tại."),
        ("BLE H20", "VIẾT MỚI iOS", "Thay cầu nối Kotlin Android bằng Swift/CoreBluetooth; nhận MAIN ngắn từ 9E3B0002 và gửi lệnh qua 9E3B0003."),
        ("Mic/loa HFP", "VIẾT MỚI iOS", "Thay điều khiển audio Android/Web bằng AVAudioSession; H20 phải được ghép đôi trong Cài đặt Bluetooth của iPhone."),
        ("Chạy nền", "VIẾT MỚI iOS", "Thay Android Foreground Service bằng bluetooth-central, state restoration và xử lý gián đoạn theo quy tắc iOS."),
        ("PWA/Safari", "BỎ KHỎI BẢN iOS", "Không dùng service worker, hướng dẫn cài PWA, autoplay/unlock Safari; vẫn giữ riêng cho phiên bản Web."),
        ("Phát hành", "LÀM MỚI", "Ký ứng dụng, quyền riêng tư Bluetooth/micro, TestFlight và quy trình App Store."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 1800, 5360])
    headers = ("Hạng mục", "Cách xử lý", "Ghi chú")
    for idx, text in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], LIGHT_BLUE)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_font(run, bold=True, color=NAVY)

    for item, action, note in rows:
        cells = table.add_row().cells
        values = (item, action, note)
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_font(run, size=9.6, bold=(idx < 2), color=(GREEN if action == "GIỮ" and idx == 1 else "000000"))
            if idx == 1 and "VIẾT MỚI" in action:
                run.font.color.rgb = RGBColor.from_string(BLUE)
        if len(table.rows) % 2 == 1:
            for cell in cells:
                shade_cell(cell, "FAFBFC")
    set_table_geometry(table, [2200, 1800, 5360])
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, before, after, color in (
        ("Heading 1", 15, 12, 6, BLUE),
        ("Heading 2", 12.5, 8, 4, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    hrun = header_p.add_run("INNOTRIK | Ghi chú kỹ thuật")
    set_font(hrun, size=8.5, bold=True, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_after = Pt(0)
    frun = footer_p.add_run("Nội bộ - Cập nhật 19/08/2026")
    set_font(frun, size=8.5, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(3)
    title.paragraph_format.space_after = Pt(3)
    trun = title.add_run("GHI CHÚ CHUYỂN TỪ PWA SANG iOS NATIVE")
    set_font(trun, size=21, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    srun = subtitle.add_run("Phạm vi giữ lại, phần viết mới và khả năng dùng nút MAIN BLE trên App Store")
    set_font(srun, size=11.5, color=MUTED)

    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    set_table_geometry(callout, [9360])
    shade_cell(callout.cell(0, 0), LIGHT_BLUE)
    cp = callout.cell(0, 0).paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    cr1 = cp.add_run("Kết luận ngắn: ")
    set_font(cr1, bold=True, color=NAVY)
    cr2 = cp.add_run(
        "Không đưa nguyên bản Web PWA lên App Store. Sẽ dùng chung mã Flutter, giữ phần giao diện và logic, nhưng thay lớp Web/Android bằng lớp iOS native cho BLE, HFP và chạy nền."
    )
    set_font(cr2)

    add_heading(doc, "1. Tình trạng hiện tại", 1)
    add_bullet(doc, "PWA trên Safari sử dụng micro của iPhone và nút MAIN ảo trên màn hình.")
    add_bullet(doc, "PWA hiện chưa nhận trực tiếp nút MAIN vật lý qua BLE H20.")
    add_bullet(doc, "APK Android đã có BLE/HFP và logic MAIN; phần native hiện được viết bằng Kotlin, không tự chạy trên iOS.")
    add_bullet(doc, "Long press đã bỏ; iOS chỉ cần nhận một sự kiện MAIN nhấn ngắn nên giao thức đơn giản và ổn định hơn.")

    add_heading(doc, "2. Phần giữ lại và phần phải làm mới", 1)
    add_matrix(doc)

    doc.add_page_break()
    add_heading(doc, "3. Luồng dự kiến trên bản App Store", 1)
    steps = [
        "Phụ huynh ghép đôi H20 HFP trong Cài đặt Bluetooth của iPhone.",
        "Ứng dụng native quét và kết nối BLE H20 bằng CoreBluetooth.",
        "Ứng dụng đăng ký nhận sự kiện tại 9E3B0002; MAIN nhấn ngắn được chuyển vào logic MAIN hiện tại.",
        "Mic và loa được định tuyến qua H20 bằng AVAudioSession; backend/AI xử lý như hiện nay.",
        "Khi khóa màn hình, iOS nhận sự kiện BLE trong phạm vi hệ điều hành cho phép và khôi phục kết nối bằng state restoration.",
    ]
    for index, text in enumerate(steps, start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_font(run)

    add_heading(doc, "4. Điều kiện để đạt tỉ lệ thành công 85-90%", 1)
    add_bullet(doc, "ODM xác nhận packet MAIN nhấn ngắn ổn định và đúng service/characteristic.")
    add_bullet(doc, "H20 ghép đôi HFP hai chiều được trên iPhone: mic vào và loa ra.")
    add_bullet(doc, "Viết cầu nối Swift/CoreBluetooth và AVAudioSession, không chỉ build lại Web.")
    add_bullet(doc, "Kiểm thử trên iPhone thật qua TestFlight: foreground, khóa màn hình, cuộc gọi, mất kết nối và tự nối lại.")

    note = doc.add_table(rows=1, cols=1)
    note.style = "Table Grid"
    set_table_geometry(note, [9360])
    shade_cell(note.cell(0, 0), LIGHT_GRAY)
    np = note.cell(0, 0).paragraphs[0]
    np.paragraph_format.space_after = Pt(0)
    nr1 = np.add_run("Lưu ý quản lý mã nguồn: ")
    set_font(nr1, bold=True, color=NAVY)
    nr2 = np.add_run(
        "Không xóa phần PWA khỏi dự án chung. Tách lớp nền tảng Web, Android và iOS để Web vẫn hoạt động, còn App Store dùng adapter iOS native."
    )
    set_font(nr2)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
