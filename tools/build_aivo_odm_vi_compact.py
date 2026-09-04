from __future__ import annotations

from hashlib import sha256
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "tmp" / "docx" / "aivo-odm-vi-compact" / "reference-copy.docx"
SOURCE = ROOT / "output" / "docs" / "AIVO_V1_YEU_CAU_TICH_HOP_ODM_VI.docx"
OUTPUT = ROOT / "output" / "docs" / "AIVO_V1_YEU_CAU_TICH_HOP_ODM_VI_RUT_GON.docx"
EXPECTED_SOURCE_SHA256 = "2e8f4c9e8b202764568c9eca81ecd954230aad149ed6b624182c97f131a79d94"


def digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def set_font(run, size: float, *, bold: bool = False, italic: bool = False, name: str = "Arial") -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_keep(paragraph, *, next_: bool = False, lines: bool = True) -> None:
    paragraph.paragraph_format.keep_together = lines
    paragraph.paragraph_format.keep_with_next = next_


def add_title(doc: Document, text: str, *, size: float, space_after: float) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    set_keep(paragraph, next_=True)
    set_font(paragraph.add_run(text), size, bold=True)


def add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(4)
    set_keep(paragraph, next_=True)
    set_font(paragraph.add_run(text), 12.5 if level == 1 else 10.8, bold=True)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.03
    if bold_prefix and text.startswith(bold_prefix):
        set_font(paragraph.add_run(bold_prefix), 9.7, bold=True)
        set_font(paragraph.add_run(text[len(bold_prefix):]), 9.7, italic=italic)
    else:
        set_font(paragraph.add_run(text), 9.7, italic=italic)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.first_line_indent = Inches(-0.16)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    set_font(paragraph.add_run(text), 9.4)


def shade_cell(cell, fill: str = "E7E6E6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 45, start: int = 65, bottom: int = 45, end: int = 65) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    *,
    font_size: float = 8.3,
    body_font_name: str = "Arial",
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    repeat_table_header(table.rows[0])
    for index, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        shade_cell(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        set_font(paragraph.add_run(header), font_size, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, values, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            set_font(paragraph.add_run(value), font_size, name=body_font_name)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    return table


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break()
    paragraph.runs[-1]._element.getparent().remove(paragraph.runs[-1]._element)
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:pageBreakBefore"))


def build() -> None:
    if digest(REFERENCE) != EXPECTED_SOURCE_SHA256:
        raise RuntimeError("Reference copy does not match the inspected Vietnamese source.")

    doc = Document(REFERENCE)
    clear_body(doc)
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    add_title(doc, "AIVO V1", size=20, space_after=2)
    add_title(doc, "YÊU CẦU TÍCH HỢP PHẦN CỨNG VÀ GIAO THỨC APP", size=14.5, space_after=8)
    add_body(
        doc,
        "Mục đích: xác định các yêu cầu firmware và phần cứng cần ODM triển khai hoặc xác nhận để mẫu AIVO hoạt động thống nhất với Android và iOS Native. Bộ UUID, packet và Raw Hex trong tài liệu là đề xuất của phía APP; chỉ được xem là chốt sau khi ODM xác nhận và cung cấp dữ liệu thực tế từ firmware mẫu.",
    )

    add_heading(doc, "1. Kiến trúc tích hợp bắt buộc")
    for item in [
        "Bluetooth Classic HFP dùng cho âm thanh hai chiều: micro thiết bị truyền tiếng nói đến APP; âm thanh từ APP phát qua loa thiết bị.",
        "BLE chỉ truyền sự kiện nút, pin, phiên bản firmware và trạng thái APP; không truyền PCM, Opus hoặc âm thanh thời gian thực qua BLE.",
        "HFP và BLE phải hoạt động đồng thời, ổn định; việc mở hoặc đóng đường HFP/SCO không được làm mất Button Event Indication.",
        "Thiết bị không tự nhận dạng, dịch, chấm điểm, tải âm thanh lên cloud hoặc tự chọn nội dung phát. Toàn bộ logic học và AI do APP quyết định.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Chức năng nút")
    add_table(
        doc,
        ["Nút", "Thao tác", "Kết quả yêu cầu"],
        [
            ["Giảm âm lượng", "Bấm ngắn", "Thiết bị giảm âm lượng cục bộ"],
            ["Giảm âm lượng", "Bấm giữ", "Gửi VOLUME_DOWN_LONG; APP chuyển sang câu tiếp theo"],
            ["Tăng âm lượng", "Bấm ngắn", "Thiết bị tăng âm lượng cục bộ"],
            ["Tăng âm lượng", "Bấm giữ", "Gửi VOLUME_UP_LONG; APP trở về câu trước"],
            ["Nguồn", "Bấm ngắn", "Gửi POWER_SHORT; APP phát lại câu tiếng Anh hiện tại trong luyện nghe"],
            ["Nguồn", "Giữ 2–3 giây", "Thiết bị bật hoặc tắt nguồn cục bộ"],
            ["MAIN", "Bấm ngắn", "Gửi MAIN_SHORT; gọi hoặc tiếp tục trợ lý AI"],
            ["MAIN", "Bấm giữ", "Gửi MAIN_LONG; tạm dừng hoạt động hiện tại của APP"],
        ],
        [1.35, 1.25, 4.2],
        font_size=8.5,
    )
    add_body(
        doc,
        "Quy tắc bắt buộc: mỗi thao tác vật lý chỉ tạo một sự kiện logic; LONG không được phát kèm SHORT. POWER_SHORT và MAIN phải là các sự kiện BLE độc lập, không ánh xạ sang Siri, Google Assistant hoặc phím media của hệ điều hành. Nếu không ở màn hình luyện nghe hoặc chưa có câu tiếng Anh hiện tại, APP có thể bỏ qua POWER_SHORT.",
        bold_prefix="Quy tắc bắt buộc:",
    )

    add_page_break(doc)
    add_heading(doc, "3. Giao thức BLE đề xuất từ phía APP")
    add_heading(doc, "3.1 UUID và thuộc tính", level=2)
    add_table(
        doc,
        ["Đối tượng", "UUID", "Thuộc tính yêu cầu"],
        [
            ["Control Service", "9E3B0001-4A7C-4D6F-8B21-5C17A2D94010", "Primary Service; advertise UUID"],
            ["Button Event", "9E3B0002-4A7C-4D6F-8B21-5C17A2D94010", "Indicate bắt buộc; CCCD 2902"],
            ["APP State", "9E3B0003-4A7C-4D6F-8B21-5C17A2D94010", "Write With Response bắt buộc"],
            ["Battery", "180F / 2A19", "Read + Notify; giá trị 0–100"],
            ["Firmware Revision", "180A / 2A26", "Read; UTF-8"],
        ],
        [1.45, 3.25, 2.1],
        font_size=8.1,
    )

    add_heading(doc, "3.2 Button Event — 12 byte, Little Endian", level=2)
    add_table(
        doc,
        ["Byte", "Trường", "Giá trị hoặc quy tắc"],
        [
            ["0", "Protocol version", "0x01"],
            ["1", "Button ID", "0x01 MAIN; 0x02 POWER; 0x03 VOLUME_UP; 0x04 VOLUME_DOWN"],
            ["2", "Gesture", "0x01 SHORT; 0x02 LONG; 0x03 RELEASE (tùy chọn)"],
            ["3", "Flags", "0x00 ở V1; dành cho mở rộng"],
            ["4–5", "Sequence", "uint16 LE; tăng sau mỗi sự kiện logic"],
            ["6", "Battery", "0–100; 0xFF nếu chưa biết"],
            ["7", "Reserved", "0x00"],
            ["8–11", "Uptime", "uint32 LE; đơn vị mili giây"],
        ],
        [0.8, 1.7, 4.3],
        font_size=8.0,
    )

    add_heading(doc, "3.3 APP State — 8 byte", level=2)
    add_table(
        doc,
        ["Byte", "Trường", "Giá trị hoặc quy tắc"],
        [
            ["0", "Protocol version", "0x01"],
            ["1", "APP State", "00 IDLE; 01 RECORDING; 02 PROCESSING; 03 READY; 04 PLAYING; 05 ERROR; 06 PAUSED"],
            ["2", "Result/ACK", "00 ACCEPTED; 01 BUSY; 02 NO_RESULT; 03 MIC_UNAVAILABLE; 04 BT_ROUTE_UNAVAILABLE; 05 DUPLICATE; 06 INTERNAL_ERROR"],
            ["3", "Flags", "0x00 ở V1"],
            ["4–5", "Related sequence", "Sequence của Button Event được phản hồi; uint16 LE"],
            ["6–7", "Reserved", "0x00 0x00"],
        ],
        [0.8, 1.7, 4.3],
        font_size=7.9,
    )

    add_page_break(doc)
    add_heading(doc, "4. Quy tắc phản hồi của APP")
    for item in [
        "Sau khi nhận Button Event hợp lệ, APP phản hồi APP State qua 9E3B0003, kèm Related Sequence và Result/ACK tương ứng.",
        "APP dùng Sequence và cửa sổ thời gian để bỏ sự kiện trùng; firmware ODM vẫn phải chống rung nút và không phát hai sự kiện cho một thao tác.",
        "MAIN_LONG làm dừng hoặc hủy hoạt động hiện tại, dừng phát âm thanh và đưa APP vào PAUSED; MAIN_SHORT dùng để bắt đầu hoặc tiếp tục.",
        "Chuyển câu do APP quyết định. Thiết bị chỉ gửi sự kiện và không tự thay đổi nội dung học.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. Thông tin ODM cần xác nhận hoặc cung cấp")
    for item in [
        "Chip hoặc module Bluetooth thực tế; bằng chứng hỗ trợ HFP + BLE đồng thời; phiên bản HFP, codec CVSD/mSBC, tên Bluetooth Classic, tên BLE advertise, firmware và mã PCB. Nếu có chip/module phụ, cung cấp model, datasheet và sơ đồ kết nối.",
        "Xác nhận UUID, thuộc tính Characteristic và packet ở Mục 3; cung cấp GATT dump đầy đủ và Raw Hex thực tế của MAIN_SHORT, MAIN_LONG, POWER_SHORT, VOLUME_UP_LONG và VOLUME_DOWN_LONG.",
        "Xác nhận mỗi thao tác chỉ tạo một Indication; LONG không kèm SHORT; cung cấp ngưỡng thời gian SHORT/LONG và quy tắc debounce.",
        "Dung lượng pin, thời gian sử dụng và sạc, khả năng vừa bật vừa sạc, cảnh báo dưới 20%, Battery Notify và khả năng gửi trạng thái đang sạc hoặc đầy pin.",
        "Chính sách ngủ sau 30 phút, thao tác đánh thức và hành vi tự khôi phục HFP + BLE với điện thoại đã ghép đôi.",
        "Công suất RF/EIRP thực tế của sản phẩm hoàn chỉnh, loại anten và báo cáo thử nghiệm; không dùng công suất tại chân chip thay cho EIRP của cả thiết bị.",
        "Các chức năng ENC/ANC thực tế đã bật và hiệu chỉnh; cung cấp điều kiện kiểm tra micro, loa, âm lượng và méo tiếng.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Yêu cầu OTA trước sản xuất hàng loạt")
    for item in [
        "Cung cấp phương thức OTA, giao thức hoặc SDK cho Android và iOS, định dạng firmware và quy tắc kiểm tra phiên bản.",
        "Firmware phải có CRC/SHA, xác minh chữ ký và kiểm tra đúng model/bo mạch trước khi nâng cấp.",
        "Hỗ trợ tiếp tục truyền sau mất kết nối và cơ chế dual-bank hoặc rollback để tránh thiết bị bị brick.",
        "Quy định mức pin tối thiểu, giới hạn khi đang sạc, thời gian nâng cấp và trạng thái tiến trình/lỗi trả về APP.",
        "Cung cấp factory reset, chế độ bootloader/recovery và quy trình nạp firmware tại nhà máy.",
        "Nếu mẫu hiện tại chưa hỗ trợ OTA, ODM phải ghi rõ 'chưa hỗ trợ' và cung cấp kế hoạch hoàn thiện trước sản xuất hàng loạt; nội dung này không chặn việc đánh giá mẫu ban đầu.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Phụ lục A. Raw Hex tham chiếu", level=2)
    add_table(
        doc,
        ["Sự kiện", "Raw Hex đề xuất"],
        [
            ["MAIN_SHORT", "01 01 01 00 01 00 64 00 10 27 00 00"],
            ["MAIN_LONG", "01 01 02 00 02 00 64 00 20 4E 00 00"],
            ["POWER_SHORT", "01 02 01 00 03 00 64 00 30 75 00 00"],
            ["VOLUME_UP_LONG", "01 03 02 00 04 00 64 00 40 9C 00 00"],
            ["VOLUME_DOWN_LONG", "01 04 02 00 05 00 64 00 50 C3 00 00"],
            ["APP RECORDING ACK", "01 01 00 00 01 00 00 00"],
            ["APP PAUSED ACK", "01 06 00 00 02 00 00 00"],
        ],
        [2.0, 4.8],
        font_size=8.0,
        body_font_name="Consolas",
    )
    add_body(
        doc,
        "Các gói trên chỉ dùng để đối chiếu. Raw Hex thực tế từ firmware mẫu và văn bản xác nhận của ODM mới là cơ sở khóa giao thức.",
        italic=True,
    )

    doc.core_properties.title = "AIVO V1 Yêu cầu tích hợp phần cứng và giao thức APP"
    doc.core_properties.subject = "Yêu cầu HFP BLE nút bấm xác nhận ODM và OTA"
    doc.core_properties.author = "AIVO"
    doc.core_properties.keywords = "AIVO, ODM, HFP, BLE, OTA, Android, iOS"
    doc.save(OUTPUT)

    if digest(REFERENCE) != EXPECTED_SOURCE_SHA256:
        raise RuntimeError("Reference copy changed during generation.")
    check = Document(OUTPUT)
    all_text = "\n".join(p.text for p in check.paragraphs)
    all_text += "\n" + "\n".join(
        cell.text for table in check.tables for row in table.rows for cell in row.cells
    )
    required = [
        "9E3B0001-4A7C-4D6F-8B21-5C17A2D94010",
        "9E3B0002-4A7C-4D6F-8B21-5C17A2D94010",
        "9E3B0003-4A7C-4D6F-8B21-5C17A2D94010",
        "MAIN_LONG",
        "POWER_SHORT",
        "APP PAUSED ACK",
        "Write With Response",
        "Yêu cầu OTA trước sản xuất hàng loạt",
    ]
    for token in required:
        if token not in all_text:
            raise RuntimeError(f"Missing required content: {token}")
    forbidden = [
        "iOS Web/Safari",
        "Nội dung phía APP cung cấp trong tài liệu này",
        "BT8952U Audio Player Microcontroller",
        "天地盖.pdf",
    ]
    for token in forbidden:
        if token in all_text:
            raise RuntimeError(f"Unnecessary content remains: {token}")
    if len(check.tables) != 5:
        raise RuntimeError(f"Expected 5 compact tables, found {len(check.tables)}")
    print(OUTPUT)


if __name__ == "__main__":
    build()
