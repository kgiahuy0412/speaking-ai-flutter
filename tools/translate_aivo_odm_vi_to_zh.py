from copy import deepcopy
from hashlib import sha256
from pathlib import Path
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "docs" / "AIVO_V1_YEU_CAU_TICH_HOP_ODM_VI.docx"
ZH_REFERENCE = ROOT / "output" / "docs" / "AIVO_V1_ODM_INTEGRATION_REQUIREMENTS_ZH_CN.docx"
OUTPUT = ROOT / "output" / "docs" / "AIVO_V1_YEU_CAU_TICH_HOP_ODM_ZH_CN_CAP_NHAT.docx"
EXPECTED_SOURCE_SHA256 = "2e8f4c9e8b202764568c9eca81ecd954230aad149ed6b624182c97f131a79d94"


def digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def replace_paragraph_content(target, translated) -> None:
    # Keep the compact typography of the user's revised Vietnamese document.
    # The older Chinese reference used larger direct run sizes in several body
    # paragraphs, which unnecessarily expanded the updated document to 5 pages.
    target_size = next(
        (run.font.size for run in target.runs if run.font.size is not None),
        None,
    )
    target_element = target._p
    for child in list(target_element):
        if child.tag.endswith("}pPr"):
            continue
        target_element.remove(child)
    for child in translated._p:
        if child.tag.endswith("}pPr"):
            continue
        target_element.append(deepcopy(child))
    if target_size is not None:
        for run in target.runs:
            run.font.size = target_size


def replace_story_content(target_story, translated_story) -> None:
    target_element = target_story._element
    for child in list(target_element):
        target_element.remove(child)
    for child in translated_story._element:
        target_element.append(deepcopy(child))


def select_chinese_paragraphs(doc: Document):
    selected = []
    in_body = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text in {"AIVO V1", "硬件与 APP 协议集成要求"}:
            selected.append(paragraph)
            continue
        if text == "1. 对 ODM 已提供文件的结论":
            in_body = True
        if text.startswith("9. 样机验收清单"):
            break
        if in_body:
            selected.append(paragraph)
    return selected


def translate_document() -> None:
    if digest(SOURCE) != EXPECTED_SOURCE_SHA256:
        raise RuntimeError("The Vietnamese source changed after inspection; inspect it again before translating.")

    source_doc = Document(SOURCE)
    zh_doc = Document(ZH_REFERENCE)

    source_paragraphs = [p for p in source_doc.paragraphs if p.text.strip()]
    translated_paragraphs = select_chinese_paragraphs(zh_doc)
    if len(source_paragraphs) != len(translated_paragraphs):
        raise RuntimeError(
            f"Paragraph mismatch: source={len(source_paragraphs)}, translated={len(translated_paragraphs)}"
        )
    for target, translated in zip(source_paragraphs, translated_paragraphs):
        replace_paragraph_content(target, translated)
        # The revised source ends at section 8, so avoid a dangling reference
        # to the removed section 9 while retaining the checklist requirement.
        if target.text == "第 9 节的样机验收清单，供 ODM 发样前完成固件自测。":
            for run in target.runs:
                run.text = run.text.replace("第 9 节的", "")

    translated_tables = zh_doc.tables[1:7]
    if len(source_doc.tables) != len(translated_tables):
        raise RuntimeError(
            f"Table mismatch: source={len(source_doc.tables)}, translated={len(translated_tables)}"
        )
    for target_table, translated_table in zip(source_doc.tables, translated_tables):
        if len(target_table.rows) != len(translated_table.rows):
            raise RuntimeError("Table row mismatch")
        for target_row, translated_row in zip(target_table.rows, translated_table.rows):
            if len(target_row.cells) != len(translated_row.cells):
                raise RuntimeError("Table column mismatch")
            for target_cell, translated_cell in zip(target_row.cells, translated_row.cells):
                if len(target_cell.paragraphs) != len(translated_cell.paragraphs):
                    raise RuntimeError("Table paragraph mismatch")
                for target_paragraph, translated_paragraph in zip(
                    target_cell.paragraphs,
                    translated_cell.paragraphs,
                ):
                    replace_paragraph_content(target_paragraph, translated_paragraph)

    source_section = source_doc.sections[0]
    zh_section = zh_doc.sections[0]
    for target_story, translated_story in [
        (source_section.header, zh_section.header),
        (source_section.even_page_header, zh_section.even_page_header),
        (source_section.first_page_header, zh_section.first_page_header),
        (source_section.footer, zh_section.footer),
        (source_section.even_page_footer, zh_section.even_page_footer),
        (source_section.first_page_footer, zh_section.first_page_footer),
    ]:
        replace_story_content(target_story, translated_story)

    source_doc.core_properties.title = "AIVO V1 硬件与 APP 协议集成要求"
    source_doc.core_properties.subject = "HFP BLE 控制 按键协议与 OTA 要求"
    source_doc.core_properties.author = "AIVO"
    source_doc.core_properties.keywords = "AIVO, ODM, HFP, BLE, OTA, Android, iOS"
    source_doc.save(OUTPUT)

    if digest(SOURCE) != EXPECTED_SOURCE_SHA256:
        raise RuntimeError("The Vietnamese source was modified during translation.")
    check = Document(OUTPUT)
    visible_text = "\n".join(p.text for p in check.paragraphs)
    visible_text += "\n" + "\n".join(
        cell.text for table in check.tables for row in table.rows for cell in row.cells
    )
    required = [
        "硬件与 APP 协议集成要求",
        "7. 本文件已由 APP 方提供的内容",
        "9E3B0001-4A7C-4D6F-8B21-5C17A2D94010",
        "MAIN_LONG",
        "APP PAUSED ACK",
        "8. OTA 要求",
    ]
    for token in required:
        if token not in visible_text:
            raise RuntimeError(f"Missing translated content: {token}")
    if re.search(r"[ăâđêôơưĂÂĐÊÔƠƯ]", visible_text):
        raise RuntimeError("Vietnamese body text remains in the Chinese output")
    if len(check.tables) != 6:
        raise RuntimeError(f"Expected 6 tables, found {len(check.tables)}")
    print(OUTPUT)


if __name__ == "__main__":
    translate_document()
