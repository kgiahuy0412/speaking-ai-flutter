from __future__ import annotations

from hashlib import sha256
from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "docs" / "AIVO_V1_YEU_CAU_TICH_HOP_ODM_VI_RUT_GON.docx"
OUTPUT = ROOT / "output" / "docs" / "AIVO_V1_YEU_CAU_TICH_HOP_ODM_ZH_CN_RUT_GON.docx"
EXPECTED_SOURCE_SHA256 = "75fb709755336af43f15e220feed68f2a3c32cf64b1ae8f6685e042968ccef9d"


PARAGRAPH_TRANSLATIONS = {
    "AIVO V1": "AIVO V1",
    "YÊU CẦU TÍCH HỢP PHẦN CỨNG VÀ GIAO THỨC APP": "硬件与 APP 协议集成要求",
    "Mục đích: xác định các yêu cầu firmware và phần cứng cần ODM triển khai hoặc xác nhận để mẫu AIVO hoạt động thống nhất với Android và iOS Native. Bộ UUID, packet và Raw Hex trong tài liệu là đề xuất của phía APP; chỉ được xem là chốt sau khi ODM xác nhận và cung cấp dữ liệu thực tế từ firmware mẫu.": (
        "目的：明确 ODM 必须实施或确认的固件与硬件要求，确保 AIVO 样机能够统一适配 Android 与 iOS 原生应用。"
        "本文中的 UUID、数据包和 Raw Hex 由 APP 方提出；只有在 ODM 确认并提供样机固件的实际数据后，方可正式锁定。"
    ),
    "1. Kiến trúc tích hợp bắt buộc": "1. 强制集成架构",
    "Bluetooth Classic HFP dùng cho âm thanh hai chiều: micro thiết bị truyền tiếng nói đến APP; âm thanh từ APP phát qua loa thiết bị.": (
        "Bluetooth Classic HFP 用于双向音频：设备麦克风将语音传送至 APP；APP 音频通过设备扬声器播放。"
    ),
    "BLE chỉ truyền sự kiện nút, pin, phiên bản firmware và trạng thái APP; không truyền PCM, Opus hoặc âm thanh thời gian thực qua BLE.": (
        "BLE 仅传输按键事件、电量、固件版本和 APP 状态；不得通过 BLE 传输 PCM、Opus 或实时音频。"
    ),
    "HFP và BLE phải hoạt động đồng thời, ổn định; việc mở hoặc đóng đường HFP/SCO không được làm mất Button Event Indication.": (
        "HFP 与 BLE 必须同时稳定工作；开启或关闭 HFP/SCO 音频通道时，不得丢失 Button Event Indication。"
    ),
    "Thiết bị không tự nhận dạng, dịch, chấm điểm, tải âm thanh lên cloud hoặc tự chọn nội dung phát. Toàn bộ logic học và AI do APP quyết định.": (
        "设备不得自行进行识别、翻译、评分、向云端上传音频或选择播放内容。全部学习与 AI 逻辑由 APP 决定。"
    ),
    "2. Chức năng nút": "2. 按键功能",
    "Quy tắc bắt buộc: mỗi thao tác vật lý chỉ tạo một sự kiện logic; LONG không được phát kèm SHORT. POWER_SHORT và MAIN phải là các sự kiện BLE độc lập, không ánh xạ sang Siri, Google Assistant hoặc phím media của hệ điều hành. Nếu không ở màn hình luyện nghe hoặc chưa có câu tiếng Anh hiện tại, APP có thể bỏ qua POWER_SHORT.": (
        "强制规则：每次物理操作只能产生一个逻辑事件；LONG 不得同时产生 SHORT。POWER_SHORT 与 MAIN 必须作为独立 BLE 事件发送，"
        "不得映射为 Siri、Google Assistant 或系统媒体键。若当前不在听力练习页面，或尚无当前英文句子，APP 可忽略 POWER_SHORT。"
    ),
    "3. Giao thức BLE đề xuất từ phía APP": "3. APP 方建议的 BLE 协议",
    "3.1 UUID và thuộc tính": "3.1 UUID 与属性",
    "3.2 Button Event — 12 byte, Little Endian": "3.2 Button Event - 12 字节，Little Endian",
    "3.3 APP State — 8 byte": "3.3 APP State - 8 字节",
    "4. Quy tắc phản hồi của APP": "4. APP 响应规则",
    "Sau khi nhận Button Event hợp lệ, APP phản hồi APP State qua 9E3B0003, kèm Related Sequence và Result/ACK tương ứng.": (
        "收到有效 Button Event 后，APP 通过 9E3B0003 返回 APP State，并携带对应的 Related Sequence 与 Result/ACK。"
    ),
    "APP dùng Sequence và cửa sổ thời gian để bỏ sự kiện trùng; firmware ODM vẫn phải chống rung nút và không phát hai sự kiện cho một thao tác.": (
        "APP 使用 Sequence 和时间窗口过滤重复事件；ODM 固件仍须完成按键消抖，并确保一次操作不会产生两个事件。"
    ),
    "MAIN_LONG làm dừng hoặc hủy hoạt động hiện tại, dừng phát âm thanh và đưa APP vào PAUSED; MAIN_SHORT dùng để bắt đầu hoặc tiếp tục.": (
        "MAIN_LONG 用于停止或取消当前活动、停止音频播放并使 APP 进入 PAUSED；MAIN_SHORT 用于开始或继续。"
    ),
    "Chuyển câu do APP quyết định. Thiết bị chỉ gửi sự kiện và không tự thay đổi nội dung học.": (
        "句子切换由 APP 决定。设备只发送事件，不得自行更改学习内容。"
    ),
    "5. Thông tin ODM cần xác nhận hoặc cung cấp": "5. ODM 需要确认或提供的信息",
    "Chip hoặc module Bluetooth thực tế; bằng chứng hỗ trợ HFP + BLE đồng thời; phiên bản HFP, codec CVSD/mSBC, tên Bluetooth Classic, tên BLE advertise, firmware và mã PCB. Nếu có chip/module phụ, cung cấp model, datasheet và sơ đồ kết nối.": (
        "实际使用的 Bluetooth 芯片或模块；支持 HFP + BLE 同时工作的证明；HFP 版本、CVSD/mSBC 编解码器、"
        "Bluetooth Classic 名称、BLE 广播名称、固件版本和 PCB 编号。如使用额外芯片或模块，请提供型号、数据手册和连接框图。"
    ),
    "Xác nhận UUID, thuộc tính Characteristic và packet ở Mục 3; cung cấp GATT dump đầy đủ và Raw Hex thực tế của MAIN_SHORT, MAIN_LONG, POWER_SHORT, VOLUME_UP_LONG và VOLUME_DOWN_LONG.": (
        "确认第 3 节中的 UUID、Characteristic 属性和数据包；提供完整 GATT dump，以及 MAIN_SHORT、MAIN_LONG、"
        "POWER_SHORT、VOLUME_UP_LONG 和 VOLUME_DOWN_LONG 的实际 Raw Hex。"
    ),
    "Xác nhận mỗi thao tác chỉ tạo một Indication; LONG không kèm SHORT; cung cấp ngưỡng thời gian SHORT/LONG và quy tắc debounce.": (
        "确认每次操作只产生一个 Indication；LONG 不得附带 SHORT；提供 SHORT/LONG 的时间阈值和 debounce 规则。"
    ),
    "Dung lượng pin, thời gian sử dụng và sạc, khả năng vừa bật vừa sạc, cảnh báo dưới 20%, Battery Notify và khả năng gửi trạng thái đang sạc hoặc đầy pin.": (
        "提供电池容量、使用与充电时间、开机充电能力、低于 20% 的提示、Battery Notify，以及上报正在充电或已充满状态的能力。"
    ),
    "Chính sách ngủ sau 30 phút, thao tác đánh thức và hành vi tự khôi phục HFP + BLE với điện thoại đã ghép đôi.": (
        "提供 30 分钟无操作后的休眠策略、唤醒方式，以及与已配对手机自动恢复 HFP + BLE 的行为说明。"
    ),
    "Công suất RF/EIRP thực tế của sản phẩm hoàn chỉnh, loại anten và báo cáo thử nghiệm; không dùng công suất tại chân chip thay cho EIRP của cả thiết bị.": (
        "提供整机实际 RF/EIRP、天线类型和测试报告；不得以芯片引脚端功率代替整机 EIRP。"
    ),
    "Các chức năng ENC/ANC thực tế đã bật và hiệu chỉnh; cung cấp điều kiện kiểm tra micro, loa, âm lượng và méo tiếng.": (
        "说明样机实际启用和调试的 ENC/ANC 功能，并提供麦克风、扬声器、音量和失真测试条件。"
    ),
    "6. Yêu cầu OTA trước sản xuất hàng loạt": "6. 量产前 OTA 要求",
    "Cung cấp phương thức OTA, giao thức hoặc SDK cho Android và iOS, định dạng firmware và quy tắc kiểm tra phiên bản.": (
        "提供 OTA 方式、Android 和 iOS 的协议或 SDK、固件格式和版本检查规则。"
    ),
    "Firmware phải có CRC/SHA, xác minh chữ ký và kiểm tra đúng model/bo mạch trước khi nâng cấp.": (
        "固件必须具备 CRC/SHA、签名验证，并在升级前校验正确的型号/PCB。"
    ),
    "Hỗ trợ tiếp tục truyền sau mất kết nối và cơ chế dual-bank hoặc rollback để tránh thiết bị bị brick.": (
        "支持断点续传，以及 dual-bank 或 rollback 机制，避免设备因升级失败而变砖。"
    ),
    "Quy định mức pin tối thiểu, giới hạn khi đang sạc, thời gian nâng cấp và trạng thái tiến trình/lỗi trả về APP.": (
        "规定最低电量、充电状态限制、升级时长，以及返回 APP 的进度和错误状态。"
    ),
    "Cung cấp factory reset, chế độ bootloader/recovery và quy trình nạp firmware tại nhà máy.": (
        "提供 factory reset、bootloader/recovery 模式和工厂烧录流程。"
    ),
    "Nếu mẫu hiện tại chưa hỗ trợ OTA, ODM phải ghi rõ 'chưa hỗ trợ' và cung cấp kế hoạch hoàn thiện trước sản xuất hàng loạt; nội dung này không chặn việc đánh giá mẫu ban đầu.": (
        "如当前样机尚未支持 OTA，ODM 必须明确标注“暂不支持”，并提供量产前的完成计划；该项不影响初始样机评估。"
    ),
    "Phụ lục A. Raw Hex tham chiếu": "附录 A. Raw Hex 参考",
    "Các gói trên chỉ dùng để đối chiếu. Raw Hex thực tế từ firmware mẫu và văn bản xác nhận của ODM mới là cơ sở khóa giao thức.": (
        "以上数据包仅用于对照。只有样机固件产生的实际 Raw Hex 和 ODM 的书面确认，才能作为锁定协议的依据。"
    ),
}


TABLE_TRANSLATIONS = [
    [
        ["按键", "操作", "要求结果"],
        ["音量减", "短按", "设备本地降低音量"],
        ["音量减", "长按", "发送 VOLUME_DOWN_LONG；APP 切换至下一句"],
        ["音量加", "短按", "设备本地提高音量"],
        ["音量加", "长按", "发送 VOLUME_UP_LONG；APP 返回上一句"],
        ["电源", "短按", "发送 POWER_SHORT；APP 在听力练习中重播当前英文句子"],
        ["电源", "长按 2–3 秒", "设备本地开机或关机"],
        ["MAIN", "短按", "发送 MAIN_SHORT；调用或继续 AI 助手"],
        ["MAIN", "长按", "发送 MAIN_LONG；暂停 APP 当前活动"],
    ],
    [
        ["对象", "UUID", "要求属性"],
        ["Control Service", "9E3B0001-4A7C-4D6F-8B21-5C17A2D94010", "Primary Service；广播该 UUID"],
        ["Button Event", "9E3B0002-4A7C-4D6F-8B21-5C17A2D94010", "必须支持 Indicate；CCCD 2902"],
        ["APP State", "9E3B0003-4A7C-4D6F-8B21-5C17A2D94010", "必须支持 Write With Response"],
        ["Battery", "180F / 2A19", "Read + Notify；范围 0–100"],
        ["Firmware Revision", "180A / 2A26", "Read；UTF-8"],
    ],
    [
        ["字节", "字段", "数值或规则"],
        ["0", "Protocol version", "0x01"],
        ["1", "Button ID", "0x01 MAIN；0x02 POWER；0x03 VOLUME_UP；0x04 VOLUME_DOWN"],
        ["2", "Gesture", "0x01 SHORT；0x02 LONG；0x03 RELEASE（可选）"],
        ["3", "Flags", "V1 使用 0x00；保留扩展"],
        ["4–5", "Sequence", "uint16 LE；每个逻辑事件后递增"],
        ["6", "Battery", "0–100；未知时为 0xFF"],
        ["7", "Reserved", "0x00"],
        ["8–11", "Uptime", "uint32 LE；单位为毫秒"],
    ],
    [
        ["字节", "字段", "数值或规则"],
        ["0", "Protocol version", "0x01"],
        ["1", "APP State", "00 IDLE；01 RECORDING；02 PROCESSING；03 READY；04 PLAYING；05 ERROR；06 PAUSED"],
        ["2", "Result/ACK", "00 ACCEPTED；01 BUSY；02 NO_RESULT；03 MIC_UNAVAILABLE；04 BT_ROUTE_UNAVAILABLE；05 DUPLICATE；06 INTERNAL_ERROR"],
        ["3", "Flags", "V1 使用 0x00"],
        ["4–5", "Related sequence", "被响应的 Button Event Sequence；uint16 LE"],
        ["6–7", "Reserved", "0x00 0x00"],
    ],
    [
        ["事件", "建议 Raw Hex"],
        ["MAIN_SHORT", "01 01 01 00 01 00 64 00 10 27 00 00"],
        ["MAIN_LONG", "01 01 02 00 02 00 64 00 20 4E 00 00"],
        ["POWER_SHORT", "01 02 01 00 03 00 64 00 30 75 00 00"],
        ["VOLUME_UP_LONG", "01 03 02 00 04 00 64 00 40 9C 00 00"],
        ["VOLUME_DOWN_LONG", "01 04 02 00 05 00 64 00 50 C3 00 00"],
        ["APP RECORDING ACK", "01 01 00 00 01 00 00 00"],
        ["APP PAUSED ACK", "01 06 00 00 02 00 00 00"],
    ],
]


def digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def set_east_asian_font(run, font_name: str = "Microsoft YaHei") -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font_name)


def replace_paragraph_text(paragraph, translated: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(translated)
    else:
        paragraph.runs[0].text = translated
        for run in paragraph.runs[1:]:
            run.text = ""
    for run in paragraph.runs:
        set_east_asian_font(run)


def translate_story_label(story) -> None:
    replacements = {
        "AIVO V1 | Yêu cầu tích hợp ODM": "AIVO V1 | ODM 集成要求",
        "AIVO - Dự thảo kỹ thuật": "AIVO - 技术草案",
    }
    for paragraph in story.paragraphs:
        for run in paragraph.runs:
            for source, translated in replacements.items():
                if source in run.text:
                    run.text = run.text.replace(source, translated)
            set_east_asian_font(run)


def translate() -> None:
    if digest(SOURCE) != EXPECTED_SOURCE_SHA256:
        raise RuntimeError("The compact Vietnamese source changed after inspection.")

    doc = Document(SOURCE)

    translated_count = 0
    for paragraph in doc.paragraphs:
        source_text = paragraph.text.strip()
        if not source_text:
            continue
        translated = PARAGRAPH_TRANSLATIONS.get(source_text)
        if translated is None:
            raise RuntimeError(f"Unmapped Vietnamese paragraph: {source_text}")
        replace_paragraph_text(paragraph, translated)
        translated_count += 1

    if translated_count != len(PARAGRAPH_TRANSLATIONS):
        raise RuntimeError(
            f"Paragraph coverage mismatch: translated={translated_count}, expected={len(PARAGRAPH_TRANSLATIONS)}"
        )

    if len(doc.tables) != len(TABLE_TRANSLATIONS):
        raise RuntimeError(
            f"Table count mismatch: source={len(doc.tables)}, translations={len(TABLE_TRANSLATIONS)}"
        )

    for table, translated_rows in zip(doc.tables, TABLE_TRANSLATIONS):
        if len(table.rows) != len(translated_rows):
            raise RuntimeError("Table row count mismatch")
        for row, translated_cells in zip(table.rows, translated_rows):
            if len(row.cells) != len(translated_cells):
                raise RuntimeError("Table column count mismatch")
            for cell, translated_text in zip(row.cells, translated_cells):
                if not cell.paragraphs:
                    raise RuntimeError("Unexpected table cell without paragraph")
                replace_paragraph_text(cell.paragraphs[0], translated_text)
                for extra in cell.paragraphs[1:]:
                    replace_paragraph_text(extra, "")

    section = doc.sections[0]
    for story in (
        section.header,
        section.even_page_header,
        section.first_page_header,
        section.footer,
        section.even_page_footer,
        section.first_page_footer,
    ):
        translate_story_label(story)

    doc.core_properties.title = "AIVO V1 硬件与 APP 协议集成要求"
    doc.core_properties.subject = "HFP BLE 按键协议 ODM 确认与 OTA 要求"
    doc.core_properties.author = "AIVO"
    doc.core_properties.keywords = "AIVO, ODM, HFP, BLE, OTA, Android, iOS"
    doc.save(OUTPUT)

    if digest(SOURCE) != EXPECTED_SOURCE_SHA256:
        raise RuntimeError("The Vietnamese source was modified during translation.")

    check = Document(OUTPUT)
    visible_text = "\n".join(p.text for p in check.paragraphs)
    visible_text += "\n" + "\n".join(
        cell.text for table in check.tables for row in table.rows for cell in row.cells
    )
    required = [
        "硬件与 APP 协议集成要求",
        "1. 强制集成架构",
        "3. APP 方建议的 BLE 协议",
        "9E3B0001-4A7C-4D6F-8B21-5C17A2D94010",
        "9E3B0002-4A7C-4D6F-8B21-5C17A2D94010",
        "9E3B0003-4A7C-4D6F-8B21-5C17A2D94010",
        "MAIN_LONG",
        "APP PAUSED ACK",
        "必须支持 Write With Response",
        "6. 量产前 OTA 要求",
        "附录 A. Raw Hex 参考",
    ]
    for token in required:
        if token not in visible_text:
            raise RuntimeError(f"Missing translated content: {token}")

    if re.search(r"[ăâđêôơưĂÂĐÊÔƠƯ]", visible_text):
        raise RuntimeError("Vietnamese body text remains in the Chinese output")
    if len(check.tables) != 5 or len(check.sections) != 1:
        raise RuntimeError("The source document structure was not preserved")

    for table, expected_rows in zip(check.tables, TABLE_TRANSLATIONS):
        actual_rows = [[cell.text for cell in row.cells] for row in table.rows]
        if actual_rows != expected_rows:
            raise RuntimeError("Translated table content does not match the approved mapping")

    print(OUTPUT)


if __name__ == "__main__":
    translate()
