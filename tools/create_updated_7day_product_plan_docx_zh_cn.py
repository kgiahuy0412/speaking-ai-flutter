from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from create_updated_7day_product_plan_docx import (
    AMBER,
    BLUE,
    DEEP_BLUE,
    GREEN,
    LIGHT_BLUE,
    LIGHT_GRAY,
    MUTED,
    NAVY,
    PALE_AMBER,
    PALE_BLUE,
    PALE_GREEN,
    PALE_RED,
    RED,
    TEXT,
    add_body,
    add_bullet,
    add_callout,
    add_cell_lines,
    add_heading,
    add_number,
    add_page_field,
    add_run,
    configure_table,
    set_cell_margins,
    set_cell_shading,
    set_cell_text,
    set_cell_width,
    set_table_header,
    setup_styles,
)


OUTPUT = Path(
    "output/docx/Ke_hoach_ca_nhan_7_ngay_H20_iOS_Native_Tieng_Trung_Gian_The_20-26_08_2026.docx"
)
CHINESE_FONT = "Microsoft YaHei"


def setup_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, width in enumerate((5800, 3560)):
        set_cell_width(table.cell(0, index), width)
        set_cell_margins(table.cell(0, index), 0, 0, 0, 0)
    set_cell_text(
        table.cell(0, 0),
        "INNOTRIK • H20 • 产品收口与 iOS NATIVE",
        bold=True,
        color=DEEP_BLUE,
        size=8.5,
    )
    set_cell_text(
        table.cell(0, 1),
        "20–26/08/2026",
        bold=True,
        color=MUTED,
        size=8.5,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    if header.paragraphs:
        header.paragraphs[0]._element.getparent().remove(header.paragraphs[0]._element)

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    footer_table.autofit = False
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, width in enumerate((3120, 3120, 3120)):
        set_cell_width(footer_table.cell(0, index), width)
        set_cell_margins(footer_table.cell(0, index), 0, 0, 0, 0)
    set_cell_text(footer_table.cell(0, 0), "内部资料 • 更新版 2.0", color=MUTED, size=8)
    set_cell_text(
        footer_table.cell(0, 1),
        "范围：个人工作",
        color=MUTED,
        size=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    right = footer_table.cell(0, 2)
    right.text = ""
    paragraph = right.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    add_run(paragraph, "第 ", color=MUTED, size=8)
    add_page_field(paragraph)
    add_run(paragraph, " 页", color=MUTED, size=8)
    if footer.paragraphs:
        footer.paragraphs[0]._element.getparent().remove(footer.paragraphs[0]._element)


def add_day_block(
    doc: Document,
    day: str,
    date: str,
    title: str,
    objective: str,
    morning: list[str],
    afternoon: list[str],
    outputs: list[str],
    gate: str,
    *,
    accent: str,
    fill: str,
) -> None:
    heading = doc.add_heading(f"第 {day} 天 • {date} — {title}", level=2)
    heading.paragraph_format.keep_with_next = True

    add_callout(doc, "当日目标", objective, fill=fill, accent=accent)

    phases = doc.add_table(rows=1, cols=2)
    configure_table(phases, [4680, 4680], header=False)
    set_cell_shading(phases.cell(0, 0), LIGHT_GRAY)
    set_cell_shading(phases.cell(0, 1), LIGHT_GRAY)
    add_cell_lines(
        phases.cell(0, 0),
        [("上午", True, accent)] + [(f"• {item}", False, TEXT) for item in morning],
        size=9.25,
    )
    add_cell_lines(
        phases.cell(0, 1),
        [("下午", True, accent)] + [(f"• {item}", False, TEXT) for item in afternoon],
        size=9.25,
    )

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    add_run(paragraph, "必交付物：", bold=True, color=DEEP_BLUE)
    add_run(paragraph, "；".join(outputs), color=TEXT)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    add_run(paragraph, "当日验收门槛：", bold=True, color=GREEN)
    add_run(paragraph, gate.rstrip("。"), color=TEXT)


def _iter_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested in cell.tables:
                    for nested_row in nested.rows:
                        for nested_cell in nested_row.cells:
                            yield from nested_cell.paragraphs
    for section in doc.sections:
        for part in (section.header, section.footer):
            yield from part.paragraphs
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def apply_chinese_fonts(doc: Document) -> None:
    for style in doc.styles:
        if not hasattr(style, "font"):
            continue
        style.font.name = CHINESE_FONT
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:ascii"), CHINESE_FONT)
        r_fonts.set(qn("w:hAnsi"), CHINESE_FONT)
        r_fonts.set(qn("w:eastAsia"), CHINESE_FONT)
        r_fonts.set(qn("w:cs"), CHINESE_FONT)
        language = r_pr.find(qn("w:lang"))
        if language is None:
            language = OxmlElement("w:lang")
            r_pr.append(language)
        language.set(qn("w:val"), "zh-CN")
        language.set(qn("w:eastAsia"), "zh-CN")

    for paragraph in _iter_paragraphs(doc):
        for run in paragraph.runs:
            run.font.name = CHINESE_FONT
            r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            r_fonts.set(qn("w:ascii"), CHINESE_FONT)
            r_fonts.set(qn("w:hAnsi"), CHINESE_FONT)
            r_fonts.set(qn("w:eastAsia"), CHINESE_FONT)
            r_fonts.set(qn("w:cs"), CHINESE_FONT)
            r_pr = run._element.get_or_add_rPr()
            language = r_pr.find(qn("w:lang"))
            if language is None:
                language = OxmlElement("w:lang")
                r_pr.append(language)
            language.set(qn("w:val"), "zh-CN")
            language.set(qn("w:eastAsia"), "zh-CN")


def build_document() -> Document:
    doc = Document()
    setup_styles(doc)
    setup_header_footer(doc)

    props = doc.core_properties
    props.title = "H20 产品收口与 iOS Native 准备：7 天个人计划"
    props.subject = "UI、AI 助手、延迟、iOS Native、BLE/HFP、测试与汇报"
    props.author = "INNOTRIK"
    props.keywords = "H20, Flutter, iOS Native, BLE, HFP, Speech, UI, AI 助手, 7 天计划"

    eyebrow = doc.add_paragraph(style="Eyebrow")
    add_run(eyebrow, "个人计划 • 7 天 • 按风险排序")

    title = doc.add_paragraph(style="Title")
    add_run(title, "完成 H20 产品收口并启动 iOS Native", bold=True, color=NAVY, size=24)

    subtitle = doc.add_paragraph(style="Subtitle")
    add_run(
        subtitle,
        "完成尚未收口的功能流程与 UI，量化 AI 助手效果，降低连续语音延迟，"
        "同时形成可在 iPhone 上验证的 iOS Native 版本。",
        color=MUTED,
        size=12,
    )

    meta = doc.add_table(rows=2, cols=4)
    configure_table(meta, [1680, 3000, 1680, 3000], header=False)
    metadata = [
        ("时间", "20–26/08/2026"),
        ("负责人", "个人 — 不包含同事工作"),
        ("本周目标", "UI 定稿 • 功能可量化 • iOS Native PoC"),
        ("原则", "每日必须有交付物、数据及通过/未通过门槛"),
    ]
    for index, (label, value) in enumerate(metadata):
        row = index // 2
        col = (index % 2) * 2
        set_cell_shading(meta.cell(row, col), LIGHT_BLUE)
        set_cell_shading(meta.cell(row, col + 1), PALE_BLUE)
        set_cell_text(meta.cell(row, col), label, bold=True, color=DEEP_BLUE, size=9)
        set_cell_text(meta.cell(row, col + 1), value, color=TEXT, size=9)

    add_heading(doc, "执行摘要", 1)
    add_callout(
        doc,
        "已确认的优先顺序",
        "紧急且重要：先处理（1）功能，再处理（2）UI。不紧急但重要：（3）用户体验与 iOS Native，"
        "因为已有基础代码。但 iOS Native 必须在本周启动，避免成为后续发布风险。",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_bullet(doc, "功能必须通过准确率、P50/P95 延迟、每 1,000 次成本和回退率评估，不能只凭主观感受。")
    add_bullet(doc, "UI 于第 2 天结束时冻结；之后仅修复功能、无障碍或 P0/P1 显示问题。")
    add_bullet(
        doc,
        "“追加训练”当前指扩展 intent 的词、短语、表达变体和带标注测试集；在没有数据与儿童数据使用同意前，不进行模型微调。",
    )
    add_bullet(
        doc,
        "7 天内的 iOS 版本是有证据的 PoC/内部版。App Store 审核时长属于计划外依赖，不作为本周完成承诺。",
    )

    add_heading(doc, "1. 标准化待办清单", 1)
    backlog = doc.add_table(rows=1, cols=4)
    configure_table(backlog, [1180, 2100, 3860, 2220])
    set_table_header(backlog, ["类别", "优先级", "待解决问题", "目标结果"])
    backlog_rows = [
        (
            "2 • 功能",
            "P0 — 紧急",
            "整体成本未优化；AI 助手缺少表达变体；连续语音延迟高；范围外回退规则不清晰。",
            "基线、优化版、前后数据及两步回退规则。",
        ),
        (
            "1 • UI",
            "P0 — 功能之后",
            "背景、图标、课程位置、内容流程、使用流程及进入歌曲流程尚未最终确认。",
            "UI 规格 v1.0 + 原型/页面清单 + 冻结决策。",
        ),
        (
            "3 • 体验",
            "P1 — 重要",
            "iOS 尚未完全 Native；BLE/HFP 与音频管理未验证；近/远场降噪不足；尚未完成北部/中部/南部儿童测试。",
            "iOS Native PoC + 音频/地区矩阵 + 优先级问题单。",
        ),
    ]
    for index, values in enumerate(backlog_rows):
        row = backlog.add_row()
        fill = PALE_RED if index == 0 else PALE_AMBER if index == 1 else PALE_GREEN
        accent = RED if index == 0 else AMBER if index == 1 else GREEN
        set_cell_shading(row.cells[0], fill)
        set_cell_shading(row.cells[1], fill)
        for col, value in enumerate(values):
            set_cell_text(
                row.cells[col],
                value,
                bold=col in (0, 1),
                color=accent if col in (0, 1) else TEXT,
                size=9.1,
                align=WD_ALIGN_PARAGRAPH.CENTER if col == 1 else None,
            )

    add_heading(doc, "需确认的范围外输入回退规则", 2)
    add_number(doc, "第 1 次：儿童语音清晰但超出当前 intent；简短回应、重述有效选项并重新打开麦克风。")
    add_number(doc, "第 2 次仍超出范围：说再见，干净结束会话并返回 MAIN 等待状态。")
    add_number(doc, "静音、ASR 置信度低、麦克风/HFP/网络错误或 backend 超时，不计为范围外输入。")
    add_number(doc, "日志必须分开记录：no_speech、low_confidence、out_of_scope、network_error、audio_route_error、cancelled。")

    add_heading(doc, "2. 7 天详细计划", 1)
    add_body(doc, "本计划保留第 1–3 天的原始方向，并补充第 4–7 天的 iOS Native、H20、地区测试与发布门槛。")

    days = [
        dict(
            day="1",
            date="20/08",
            title="设计定稿并扩展识别语料",
            objective="锁定直接影响开发的产品决策，并在优化前建立 AI 助手基线。",
            morning=[
                "确定背景、图标、字体、颜色以及 loading/error/success 状态。",
                "确定课程位置、主题/词汇/翻译/歌曲入口，并绘制 screen/flow map。",
                "确定提示语、屏幕内容及每条流程的入口和出口。",
            ],
            afternoon=[
                "扩展已有 intent 的表达变体：学主题、学新词、翻译英语、下一句/上一句/停止。",
                "建立带标注测试集：标准句、地区表达、缺词句、范围外句。",
                "若 Huyền 姐已确认流程，则在独立 branch 实施并测试；否则继续验证 AI 助手，不阻塞 UI 规格。",
            ],
            outputs=["UI/UX 规格 v1.0", "intent lexicon v1", "准确率/回退基线报告"],
            gate="不再存在未决的 UI P0；每个核心 intent 都有标准句与变体测试集。",
            accent=BLUE,
            fill=PALE_BLUE,
        ),
        dict(
            day="2",
            date="21/08",
            title="降低延迟并冻结 UI",
            objective="减少儿童说完后的等待时间，同时不降低准确率，也不让麦克风过早关闭。",
            morning=[
                "埋点：mic_start、speech_start、end_of_speech、ASR_final、intent_done、translation_done、first_audio_play。",
                "分别测量短句、长句、好/弱网络的 P50/P95，并定位 ASR、网络、backend、TTS 或 state machine 瓶颈。",
                "采用可控优化：prewarm、提前上传、cache、并行处理；不能只凭感觉缩短 timeout。",
            ],
            afternoon=[
                "应用最终 UI，通过主流程 smoke test 后冻结。",
                "评估租借/借用 Mac、所需 Xcode/iPhone 配置、升级便利性及备份方案。",
                "确认 Apple Developer 个人/组织条件、法律权限、Apple Account 2FA、D‑U‑N‑S（组织）和当前费用。",
                "向管理层汇报 UI、功能、延迟数据及 iOS 准备状态。",
            ],
            outputs=["延迟前后对比表", "UI release candidate", "Mac + Apple Developer 备忘录", "一页管理汇报"],
            gate="P50 相比基线至少下降 20%，或达到内部目标 ≤2.5 秒；准确率下降不超过 3 个百分点；UI smoke test 全绿。",
            accent=RED,
            fill=PALE_RED,
        ),
        dict(
            day="3",
            date="22/08",
            title="启动 iOS Native 与 Apple Speech",
            objective="在真机 iPhone 上运行 Flutter Native，复用 UI/domain/backend，并从主识别路径移除 Safari 依赖。",
            morning=[
                "配置 Mac/Xcode/CocoaPods、debug signing 和测试 iPhone；缺账号或设备时记录 blocker。",
                "建立 iOS Native branch；让 Flutter app 直接运行在 iPhone 上，不把 PWA/WebView 当主版本。",
                "检查 Info.plist、麦克风/Speech/Bluetooth 权限和 deployment target。",
            ],
            afternoon=[
                "完成原生 Speech framework PoC；按 deployment target 选择 API，并检查 locale 可用性与权限。",
                "用同一测试集比较 Apple Speech Native 与 Cloudflare Batch fallback。",
                "设计单一 audio pipeline，避免两个 recorder 抢麦；原生服务不可用时保留 fallback。",
            ],
            outputs=["iPhone debug build", "Architecture Decision Record", "20 句识别结果", "signing blocker 清单"],
            gate="App 可在真实 iPhone 运行并输出原生 transcript；主路径不再依赖 Safari Web Speech。",
            accent=BLUE,
            fill=PALE_BLUE,
        ),
        dict(
            day="4",
            date="23/08",
            title="iOS 上的 BLE MAIN 与 H20 音频路由",
            objective="证明 iOS Native 能通过 BLE 接收 MAIN，并通过原生层管理 H20 麦克风/扬声器，而非依赖 PWA。",
            morning=[
                "实现 CoreBluetooth bridge：扫描/连接 service 9E3B0001，订阅 Indicate 9E3B0002，向 9E3B0003 写入 APP state。",
                "保存 peripheral，并在 iOS 允许范围内验证 reconnect/state restoration。",
                "确认产品仅使用 short MAIN；从新测试用例中移除 long press。",
            ],
            afternoon=[
                "配置 AVAudioSession playAndRecord + Bluetooth HFP；检查 currentRoute 的 input/output。",
                "测试 H20 录放音、route 切换、断开/重连、来电/通知和 phone fallback。",
                "保存 raw packet、route log、视频，并区分 APP/OS/firmware 问题。",
            ],
            outputs=["最小 native bridge", "MAIN 20/20", "HFP route matrix", "firmware 阻塞时的 ODM issue"],
            gate="20/20 short MAIN 无重复；10/10 次 H20 input/output 路由正确，或以日志证明 firmware blocker。",
            accent=GREEN,
            fill=PALE_GREEN,
        ),
        dict(
            day="5",
            date="24/08",
            title="端到端集成与安全回退",
            objective="在 iOS Native 串联 MAIN → 助手 → intent → 学习/翻译，并消除常见卡死状态。",
            morning=[
                "把 CoreBluetooth/AVAudioSession/Speech 接入 Flutter 共用 state machine。",
                "端到端运行学主题、学新词、逐句翻译和连续语音。",
                "route 变化、App 进后台、锁屏或 audio interruption 时保持状态一致。",
            ],
            afternoon=[
                "实施两步范围外回退，并区分 no_speech/low_confidence/network/audio_route。",
                "断网时不得误报成功并保留状态；恢复后只重试一次，避免重复 request。",
                "账号具备时准备 signed archive/TestFlight 内测；review mode 不强制要求 H20。",
            ],
            outputs=["iOS Native 演示视频", "10 个端到端场景", "回退测试报告", "内部 archive 或 blocker"],
            gate="无“正在准备”卡死；10/10 主流程正确结束；P0 测试集中范围外句误导航为 0。",
            accent=AMBER,
            fill=PALE_AMBER,
        ),
        dict(
            day="6",
            date="25/08",
            title="音频、地区与成本模型测试",
            objective="在真实条件下量化体验，而不是只在安静环境验证功能。",
            morning=[
                "测试安静、电视/远处人声、儿童近麦；同一脚本测 SNR/ASR/intent。",
                "在取得同意后试测北部/中部/南部儿童；区分声学、地方词和 intent mapping 错误。",
                "分开比较 Android、iOS Native、Safari PWA，不能混合三平台结果。",
            ],
            afternoon=[
                "按每 1,000 次、峰值负载和存储数据，汇总 Railway/Cloudflare/TTS-ASR/cache/log 及自建服务器成本。",
                "只有 benchmark 证明成本/延迟/控制收益时才建议迁移；数据不足不在本周仓促迁移。",
                "修复矩阵中发现的 P0；其余列为 P1/P2，明确 owner 和期限。",
            ],
            outputs=["噪声 + 地区报告", "平台矩阵", "每 1,000 次成本模型", "P0/P1/P2 清单"],
            gate="获得不同条件下的准确率/延迟；未经同意不存储或共享儿童音频；服务器决策有量化依据。",
            accent=GREEN,
            fill=PALE_GREEN,
        ),
        dict(
            day="7",
            date="26/08",
            title="回归、发布门槛与管理汇报",
            objective="用可决策的产品状态结束本周，避免结果散落在聊天或个人日志中。",
            morning=[
                "回归 Android + iOS Native：UI、MAIN、BLE reconnect、HFP route、Speech、网络、回退和内容。",
                "对照 gate；仅修复阻塞演示/发布的 P0，不新增 UI 或范围外重构。",
                "关闭可复现的 branch/tag/build，记录 hash、version、测试设备及 firmware。",
            ],
            afternoon=[
                "完成管理汇报：已达成、未达成、成本、风险、ODM/Apple 依赖和待批决策。",
                "确认后续 14 天路线图：App Store/TestFlight、扩大噪声/地区测试、服务器、privacy、analytics。",
                "作出 go/no-go：内部演示、TestFlight，或维持 Android baseline 并继续完善 iOS。",
            ],
            outputs=["release gate", "build/tag", "执行报告", "14 天路线图", "决策请求"],
            gate="所有未关闭 P0 都有 owner；其余问题均有优先级和期限；管理层具备 go/no-go 所需数据。",
            accent=BLUE,
            fill=PALE_BLUE,
        ),
    ]
    for day in days:
        add_day_block(doc, **day)

    add_heading(doc, "3. KPI 与验收标准", 1)
    kpi = doc.add_table(rows=1, cols=4)
    configure_table(kpi, [1980, 3720, 1980, 1680])
    set_table_header(kpi, ["项目", "测量指标", "内部目标", "验收日"])
    kpi_rows = [
        ("核心 intent", "各 intent 准确率；confusion matrix；误导航。", "安静环境 ≥95%；P0 集误导航为 0。", "第 1/5 天"),
        ("范围外输入", "第 1 次重述；第 2 次结束；不得误判技术错误。", "30 句受控 OOS 测试 100% 正确。", "第 5 天"),
        ("连续语音", "end-of-speech → ASR final → intent → first audio。", "P50 ≤2.5s；P95 ≤4.0s，或比基线改善 ≥20%。", "第 2/7 天"),
        ("UI", "screen/flow 覆盖；显示问题；主流程回归。", "已决 screen/flow 覆盖 100%；P0 为 0。", "第 2 天"),
        ("iOS BLE", "MAIN event、duplicate、reconnect。", "MAIN 20/20；0 duplicate；reconnect 达到矩阵要求。", "第 4 天"),
        ("iOS HFP", "currentRoute input/output；record/play；interruption。", "10/10 route 正确，或 blocker 有完整日志。", "第 4/5 天"),
        ("地区/噪声", "准确率、重说率、各条件错误。", "形成北/中/南与近/远场基线；小样本不做过度推断。", "第 6 天"),
        ("成本", "每 1,000 次及峰值下的 ASR/TTS/backend/cache/log。", "完成现状与自建服务器对比及决策阈值。", "第 6 天"),
    ]
    for index, values in enumerate(kpi_rows):
        row = kpi.add_row()
        if index % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col, value in enumerate(values):
            set_cell_text(
                row.cells[col],
                value,
                bold=col == 0,
                color=TEXT,
                size=8.9,
                align=WD_ALIGN_PARAGRAPH.CENTER if col == 3 else None,
            )

    add_callout(
        doc,
        "KPI 说明",
        "以上阈值属于内部测试 gate，并非市场承诺。若基线或样本不足，报告必须注明样本量、设备、"
        "麦克风距离、环境和 confidence interval/不确定性。",
        fill=PALE_AMBER,
        accent=AMBER,
    )

    add_heading(doc, "4. 依赖、风险与待决事项", 1)
    risks = doc.add_table(rows=1, cols=4)
    configure_table(risks, [2260, 2500, 2820, 1780])
    set_table_header(risks, ["依赖/风险", "影响", "本计划处理方式", "决策期限"])
    risk_rows = [
        ("Huyền 姐的流程未定稿", "内容/state 可能返工。", "分离 UI/intent 基础层；仅按定稿版/version 集成新流程。", "第 1 天"),
        ("缺少 Mac/Xcode/iPhone", "无法运行或签名 iOS Native。", "租借/借用 Mac；确认配置与使用时段；准备备份。", "第 2 天"),
        ("Apple Developer 未就绪", "无法 TestFlight/App Store。", "选择个人/组织；2FA；法律权限；组织需 D‑U‑N‑S；按 Apple 当前价格预留 99 USD/年。", "第 2–3 天"),
        ("ODM protocol/firmware 不清楚", "MAIN/reconnect 可能错误。", "保存 raw packet；仅分发已知 packet；用视频/日志提交 issue。", "第 4 天"),
        ("原生 Speech 不可用", "locale/网络/OS 导致识别中断。", "检查 availability；保留 Cloudflare Batch fallback；禁止两个 recorder 同开。", "第 3–5 天"),
        ("儿童语音数据", "隐私风险和结果偏差。", "取得同意；匿名化；最小 retention；不收集测试目的之外的数据。", "第 6 天前"),
        ("7 天范围过大", "基线丢失、汇报延迟。", "第 7 天仅关闭 P0；App Store 审核和深度优化转入后续 14 天。", "全程"),
    ]
    for index, values in enumerate(risk_rows):
        row = risks.add_row()
        if index % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col, value in enumerate(values):
            set_cell_text(row.cells[col], value, bold=col == 0, color=TEXT, size=8.85)

    add_heading(doc, "5. 每日收尾报告模板", 1)
    report_items = [
        "已完成：最多 3 项结果，附 commit/build 和 pass/fail 状态。",
        "数据：accuracy、P50/P95 latency、测试通过数/总数、成本或 blocker。",
        "未关闭 P0：重现步骤、platform/device/firmware 及等待中的人员/输入。",
        "需要管理层决策：列出选项、时间/成本/风险影响和个人建议。",
        "未来 24 小时：三项工作，每项有明确交付物和 gate。",
    ]
    for index, item in enumerate(report_items, start=1):
        paragraph = add_body(doc, f"{index}.  {item}")
        paragraph.paragraph_format.left_indent = Inches(0.18)

    add_callout(
        doc,
        "建议结尾句",
        "今天 [build/commit] 达到 [x/y] 个 gate。与基线相比，[指标] 从 […] 变为 […]。"
        "尚有 P0：[…]。我建议对 […] 采取 [go/hold]。需要在 […] 前确认 […]。",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "实施过程中需核对的官方资料", 2)
    sources = [
        "Apple Speech framework: https://developer.apple.com/documentation/speech",
        "Apple Core Bluetooth background processing: https://developer.apple.com/library/archive/documentation/NetworkingInternetWeb/Conceptual/CoreBluetooth_concepts/CoreBluetoothBackgroundProcessingForIOSApps/PerformingTasksWhileYourAppIsInTheBackground.html",
        "AVAudioSession currentRoute: https://developer.apple.com/documentation/avfaudio/avaudiosession/currentroute",
        "Bluetooth HFP audio option: https://developer.apple.com/documentation/avfaudio/avaudiosession/categoryoptions-swift.struct/allowbluetooth",
        "Apple Developer Program enrollment: https://developer.apple.com/programs/enroll/",
    ]
    for source in sources:
        paragraph = doc.add_paragraph(style="Small Note")
        add_run(paragraph, source, color=MUTED, size=8.2)

    note = doc.add_paragraph(style="Small Note")
    add_run(
        note,
        "文档范围：个人工作计划。不包含同事负责的内容/音乐/素材开发时间、ODM firmware 修改或 Apple 审核时间。"
        "文档采用 compact_reference_guide 版式、Letter 纸张、1 英寸页边距、Microsoft YaHei 11 pt 和固定宽度表格，"
        "便于打印、汇报和更新。",
        color=MUTED,
        size=8.2,
    )

    apply_chinese_fonts(doc)
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
