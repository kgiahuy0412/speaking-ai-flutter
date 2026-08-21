from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from create_personal_7day_plan_docx import (
    AMBER,
    BLUE,
    DEEP_BLUE,
    GREEN,
    LIGHT_BLUE,
    LIGHT_GRAY,
    MUTED,
    NAVY,
    PALE_AMBER,
    PALE_BLUE,
    PALE_GREEN,
    PALE_RED,
    RED,
    TEXT,
    add_body,
    add_bullet,
    add_callout,
    add_cell_lines,
    add_heading,
    add_number,
    add_page_field,
    add_run,
    configure_table,
    set_cell_margins,
    set_cell_shading,
    set_cell_text,
    set_cell_width,
    set_table_header,
    setup_styles,
)


_set_visual_table_header = set_table_header


def set_table_header(table, labels: list[str]) -> None:
    """Style the first row and expose it as a semantic header row to Word."""
    _set_visual_table_header(table, labels)
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    row_properties.append(table_header)


OUTPUT = Path(
    "output/docx/Ke_hoach_ca_nhan_7_ngay_H20_iOS_Native_cap_nhat_20-26_08_2026.docx"
)


def setup_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, width in enumerate((5800, 3560)):
        set_cell_width(table.cell(0, index), width)
        set_cell_margins(table.cell(0, index), 0, 0, 0, 0)
    set_cell_text(
        table.cell(0, 0),
        "INNOTRIK • H20 • CHỐT SẢN PHẨM & iOS NATIVE",
        bold=True,
        color=DEEP_BLUE,
        size=8.5,
    )
    set_cell_text(
        table.cell(0, 1),
        "20–26/08/2026",
        bold=True,
        color=MUTED,
        size=8.5,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    if header.paragraphs:
        header.paragraphs[0]._element.getparent().remove(header.paragraphs[0]._element)

    footer = section.footer
    ft = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    ft.autofit = False
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, width in enumerate((3120, 3120, 3120)):
        set_cell_width(ft.cell(0, index), width)
        set_cell_margins(ft.cell(0, index), 0, 0, 0, 0)
    set_cell_text(ft.cell(0, 0), "Nội bộ • Bản cập nhật 2.0", color=MUTED, size=8)
    set_cell_text(
        ft.cell(0, 1),
        "Phạm vi: công việc cá nhân",
        color=MUTED,
        size=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    right = ft.cell(0, 2)
    right.text = ""
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Trang ", color=MUTED, size=8)
    add_page_field(p)
    if footer.paragraphs:
        footer.paragraphs[0]._element.getparent().remove(footer.paragraphs[0]._element)


def add_day_block(
    doc: Document,
    day: str,
    date: str,
    title: str,
    objective: str,
    morning: list[str],
    afternoon: list[str],
    outputs: list[str],
    gate: str,
    *,
    accent: str,
    fill: str,
) -> None:
    heading = doc.add_heading(f"Ngày {day} • {date} — {title}", level=2)
    heading.paragraph_format.keep_with_next = True

    add_callout(doc, "Mục tiêu trong ngày", objective, fill=fill, accent=accent)

    phases = doc.add_table(rows=1, cols=2)
    configure_table(phases, [4680, 4680], header=False)
    set_cell_shading(phases.cell(0, 0), LIGHT_GRAY)
    set_cell_shading(phases.cell(0, 1), LIGHT_GRAY)
    add_cell_lines(
        phases.cell(0, 0),
        [("BUỔI SÁNG", True, accent)] + [(f"• {item}", False, TEXT) for item in morning],
        size=9.25,
    )
    add_cell_lines(
        phases.cell(0, 1),
        [("BUỔI CHIỀU", True, accent)] + [(f"• {item}", False, TEXT) for item in afternoon],
        size=9.25,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "Đầu ra bắt buộc: ", bold=True, color=DEEP_BLUE)
    add_run(p, "; ".join(outputs) + ".", color=TEXT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "Gate cuối ngày: ", bold=True, color=GREEN)
    add_run(p, gate, color=TEXT)


def build_document() -> Document:
    doc = Document()
    setup_styles(doc)
    setup_header_footer(doc)

    props = doc.core_properties
    props.title = "Kế hoạch cá nhân 7 ngày chốt sản phẩm H20 và chuẩn bị iOS Native"
    props.subject = "UI, trợ lý AI, độ trễ, iOS Native, BLE/HFP, kiểm thử và báo cáo"
    props.author = "INNOTRIK"
    props.keywords = "H20, Flutter, iOS Native, BLE, HFP, Speech, UI, trợ lý AI, kế hoạch 7 ngày"

    eyebrow = doc.add_paragraph(style="Eyebrow")
    add_run(eyebrow, "KẾ HOẠCH CÁ NHÂN • 7 NGÀY • ƯU TIÊN THEO RỦI RO")

    title = doc.add_paragraph(style="Title")
    add_run(title, "Chốt sản phẩm H20 & mở đường iOS Native", bold=True, color=NAVY, size=24)

    subtitle = doc.add_paragraph(style="Subtitle")
    add_run(
        subtitle,
        "Hoàn thiện luồng chức năng và UI đang tồn đọng, đo hiệu quả trợ lý AI, "
        "giảm độ trễ nói liên tục, đồng thời tạo bản iOS Native có thể kiểm thử trên iPhone.",
        color=MUTED,
        size=12,
    )

    meta = doc.add_table(rows=2, cols=4)
    configure_table(meta, [1680, 3000, 1680, 3000], header=False)
    metadata = [
        ("Thời gian", "20–26/08/2026"),
        ("Người phụ trách", "Cá nhân — không tính công việc đồng nghiệp"),
        ("Đích tuần", "UI chốt • chức năng đo được • iOS Native PoC"),
        ("Nguyên tắc", "Mỗi ngày có đầu ra, số liệu và gate pass/fail"),
    ]
    for idx, (label, value) in enumerate(metadata):
        row = idx // 2
        col = (idx % 2) * 2
        set_cell_shading(meta.cell(row, col), LIGHT_BLUE)
        set_cell_shading(meta.cell(row, col + 1), PALE_BLUE)
        set_cell_text(meta.cell(row, col), label, bold=True, color=DEEP_BLUE, size=9)
        set_cell_text(meta.cell(row, col + 1), value, color=TEXT, size=9)

    add_heading(doc, "Kết luận điều hành", 1)
    add_callout(
        doc,
        "Thứ tự ưu tiên đã chốt",
        "Khẩn cấp và quan trọng: (1) chức năng, sau đó (2) UI. Không khẩn cấp nhưng quan trọng: "
        "(3) trải nghiệm người dùng và iOS Native, vì đã có nền code cơ bản. Tuy nhiên, iOS Native phải "
        "được khởi động trong tuần để không trở thành rủi ro phát hành muộn.",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_bullet(
        doc,
        "Chức năng phải được đo bằng độ chính xác, độ trễ P50/P95, chi phí/1.000 lượt và tỷ lệ fallback; không đánh giá bằng cảm giác.",
    )
    add_bullet(
        doc,
        "UI được đóng băng cuối Ngày 2. Sau mốc này chỉ sửa lỗi chức năng, khả năng tiếp cận hoặc lỗi hiển thị P0/P1.",
    )
    add_bullet(
        doc,
        "“Training thêm” được hiểu là mở rộng từ/cụm từ/biến thể intent và bộ kiểm thử có nhãn; chưa fine-tune mô hình khi chưa có dữ liệu và đồng ý sử dụng dữ liệu trẻ em.",
    )
    add_bullet(
        doc,
        "Bản iOS trong 7 ngày là PoC/bản nội bộ có bằng chứng. Thời gian Apple duyệt App Store là phụ thuộc ngoài kế hoạch, không dùng làm cam kết hoàn tất tuần.",
    )

    add_heading(doc, "1. Danh sách tồn đọng đã chuẩn hóa", 1)
    backlog = doc.add_table(rows=1, cols=4)
    configure_table(backlog, [1180, 2100, 3860, 2220])
    set_table_header(backlog, ["Nhóm", "Mức ưu tiên", "Vấn đề cần xử lý", "Kết quả cần đạt"])
    backlog_rows = [
        (
            "2 • Chức năng",
            "P0 — khẩn cấp",
            "Chưa tối ưu tổng chi phí; trợ lý AI thiếu bộ biến thể nhận biết; độ trễ nói liên tục còn cao; fallback câu ngoài phạm vi chưa rõ.",
            "Có baseline, bản tối ưu, số liệu trước/sau và luật fallback 2 bước.",
        ),
        (
            "1 • UI",
            "P0 — sau chức năng",
            "Chưa chốt background, icon, vị trí bài học, luồng nội dung, luồng sử dụng và luồng vào bài hát.",
            "UI spec v1.0 + prototype/screen list + quyết định đóng băng.",
        ),
        (
            "3 • Trải nghiệm",
            "P1 — quan trọng",
            "iOS chưa native hoàn toàn; BLE/HFP và quản lý âm thanh chưa chứng minh; lọc tiếng ồn gần/xa chưa đo; chưa test trẻ Bắc–Trung–Nam.",
            "PoC iOS Native + ma trận âm thanh/vùng miền + issue ưu tiên.",
        ),
    ]
    for idx, values in enumerate(backlog_rows):
        row = backlog.add_row()
        fill = PALE_RED if idx == 0 else PALE_AMBER if idx == 1 else PALE_GREEN
        accent = RED if idx == 0 else AMBER if idx == 1 else GREEN
        set_cell_shading(row.cells[0], fill)
        set_cell_shading(row.cells[1], fill)
        for col, value in enumerate(values):
            set_cell_text(
                row.cells[col],
                value,
                bold=col in (0, 1),
                color=accent if col in (0, 1) else TEXT,
                size=9.1,
                align=WD_ALIGN_PARAGRAPH.CENTER if col == 1 else None,
            )

    add_heading(doc, "Luật fallback ngoài phạm vi cần chốt", 2)
    add_number(doc, "Lần 1 trẻ nói câu có âm thanh rõ nhưng ngoài intent hiện tại: phản hồi ngắn, nêu lại các lựa chọn hợp lệ và mở mic lại.")
    add_number(doc, "Lần 2 tiếp tục ngoài phạm vi: nói lời tạm biệt, kết thúc phiên sạch và trở về trạng thái chờ MAIN.")
    add_number(doc, "Không tính im lặng, ASR độ tin cậy thấp, lỗi mic/HFP, mất mạng hoặc timeout backend là câu ngoài phạm vi.")
    add_number(doc, "Ghi log riêng: no_speech, low_confidence, out_of_scope, network_error, audio_route_error và cancelled.")

    add_heading(doc, "2. Kế hoạch chi tiết 7 ngày", 1)
    add_body(
        doc,
        "Kế hoạch giữ nguyên ý Ngày 1–3 đã đề xuất và bổ sung Ngày 4–7 theo hướng iOS Native, H20, kiểm thử vùng miền và cổng phát hành.",
    )

    add_day_block(
        doc,
        "1",
        "20/08",
        "Chốt thiết kế & mở rộng bộ nhận biết",
        "Khóa toàn bộ quyết định sản phẩm có ảnh hưởng trực tiếp đến triển khai; tạo baseline đo trợ lý AI trước khi tối ưu.",
        [
            "Chốt background, bộ icon, typography, màu và các trạng thái loading/error/success.",
            "Chốt vị trí bài học, luồng vào chủ đề, từ vựng, dịch và bài hát; lập screen/flow map.",
            "Chốt câu dẫn, nội dung trên màn hình và điểm vào/ra của từng luồng.",
        ],
        [
            "Mở rộng biến thể từ/câu cho các intent đang có: học chủ đề, học từ mới, dịch tiếng Anh, câu tiếp/câu trước/dừng.",
            "Tạo tập test có nhãn gồm câu chuẩn, câu vùng miền, câu thiếu từ và câu ngoài phạm vi.",
            "Nếu chị Huyền đã chốt luồng: triển khai trên branch riêng và test; nếu chưa, tiếp tục kiểm thử trợ lý AI, không chặn UI spec.",
        ],
        ["UI/UX spec v1.0", "intent lexicon v1", "báo cáo baseline accuracy/fallback"],
        "Không còn quyết định UI P0 bỏ ngỏ; mọi intent chủ lực có tối thiểu một tập câu chuẩn và biến thể để chạy lại tự động/thủ công.",
        accent=BLUE,
        fill=PALE_BLUE,
    )

    add_day_block(
        doc,
        "2",
        "21/08",
        "Giảm độ trễ & đóng băng UI",
        "Giảm thời gian chờ sau khi trẻ nói xong mà không làm giảm độ chính xác hoặc gây mic tự tắt sớm.",
        [
            "Gắn mốc đo: mic_start, speech_start, end_of_speech, ASR_final, intent_done, translation_done và first_audio_play.",
            "Đo P50/P95 trên câu ngắn, câu dài, mạng tốt/yếu; xác định nút thắt ASR, mạng, backend, TTS hay state machine.",
            "Tối ưu có kiểm soát: prewarm, upload sớm, cache, chạy song song; không chỉ giảm timeout một cách cảm tính.",
        ],
        [
            "Áp dụng UI cuối cùng và đóng băng UI sau khi chạy smoke test các luồng chính.",
            "Tìm hiểu thuê/mượn Mac, cấu hình Xcode/iPhone cần có, khả năng nâng cấp và phương án backup.",
            "Xác minh điều kiện Apple Developer cá nhân/tổ chức, quyền pháp lý, Apple Account 2FA, D‑U‑N‑S (nếu tổ chức) và chi phí hiện hành.",
            "Báo cáo sếp: UI, chức năng, số liệu latency và tình trạng chuẩn bị iOS.",
        ],
        ["bảng latency trước/sau", "UI release candidate", "memo Mac + Apple Developer", "báo cáo sếp 1 trang"],
        "P50 giảm tối thiểu 20% so với baseline hoặc đạt mục tiêu nội bộ ≤2,5 giây; độ chính xác không giảm quá 3 điểm phần trăm; UI smoke test xanh.",
        accent=RED,
        fill=PALE_RED,
    )

    add_day_block(
        doc,
        "3",
        "22/08",
        "Khởi tạo iOS Native & Apple Speech",
        "Chạy Flutter native trên iPhone thật, tái sử dụng UI/domain/backend nhưng loại phụ thuộc Safari khỏi đường nhận dạng chính.",
        [
            "Thiết lập Mac/Xcode/CocoaPods, signing debug và iPhone test; ghi rõ blocker nếu thiếu tài khoản hoặc thiết bị.",
            "Tạo branch iOS Native; chạy app Flutter trực tiếp trên iPhone, không đóng gói PWA/WebView làm bản chính.",
            "Rà Info.plist, quyền microphone/speech/Bluetooth và deployment target.",
        ],
        [
            "PoC Speech framework native; chọn API phù hợp deployment target, kiểm tra quyền và availability theo locale.",
            "So sánh Apple Speech native với Cloudflare Batch fallback bằng cùng tập câu test.",
            "Thiết kế một audio pipeline duy nhất để tránh hai recorder tranh mic; giữ fallback khi native service không sẵn sàng.",
        ],
        ["build debug trên iPhone", "Architecture Decision Record", "kết quả nhận dạng 20 câu", "danh sách blocker signing"],
        "App chạy được trên iPhone thật; có transcript native; không còn phụ thuộc Safari Web Speech trên đường chính.",
        accent=BLUE,
        fill=PALE_BLUE,
    )

    add_day_block(
        doc,
        "4",
        "23/08",
        "BLE MAIN & đường âm thanh H20 trên iOS",
        "Chứng minh iOS Native có thể nhận MAIN qua BLE và quản lý mic/loa H20 bằng lớp native, không dựa vào PWA.",
        [
            "Viết CoreBluetooth bridge: scan/connect service 9E3B0001, subscribe Indicate 9E3B0002, write APP state 9E3B0003.",
            "Lưu peripheral và kiểm tra reconnect/state restoration theo giới hạn iOS.",
            "Xác nhận sản phẩm chỉ dùng short MAIN; loại long press khỏi test case mới.",
        ],
        [
            "Cấu hình AVAudioSession playAndRecord + Bluetooth HFP; kiểm tra currentRoute input/output.",
            "Test record/play H20, chuyển route, ngắt/kết nối lại, cuộc gọi/thông báo và phone fallback.",
            "Lưu raw packet, route log, video và phân loại lỗi APP/OS/firmware.",
        ],
        ["native bridge tối thiểu", "MAIN 20/20", "HFP route matrix", "issue gửi ODM nếu firmware chặn"],
        "20/20 short MAIN không trùng; 10/10 phiên chọn H20 xác nhận đúng input/output hoặc có blocker firmware được chứng minh bằng log.",
        accent=GREEN,
        fill=PALE_GREEN,
    )

    add_day_block(
        doc,
        "5",
        "24/08",
        "Tích hợp end-to-end & fallback an toàn",
        "Nối MAIN → trợ lý → intent → học/dịch trên iOS Native và đóng các trạng thái kẹt thường gặp.",
        [
            "Nối CoreBluetooth/AVAudioSession/Speech vào state machine Flutter dùng chung.",
            "Chạy end-to-end học chủ đề, từ mới, dịch từng câu và nói liên tục.",
            "Bảo toàn state khi route đổi, app vào nền, màn hình khóa hoặc audio bị interruption.",
        ],
        [
            "Triển khai fallback ngoài phạm vi 2 bước; tách no_speech/low_confidence/network/audio_route.",
            "Khi mạng mất: không báo thành công, giữ trạng thái; khi mạng trở lại: cho retry đúng một lần, không nhân đôi request.",
            "Chuẩn bị signed archive/TestFlight nội bộ nếu account đã sẵn sàng; giữ review mode không bắt buộc có H20.",
        ],
        ["video demo iOS Native", "10 kịch bản end-to-end", "fallback test report", "archive nội bộ hoặc blocker"],
        "Không còn state “đang chuẩn bị” kẹt; 10/10 luồng chính kết thúc đúng; 0 điều hướng nhầm với câu ngoài phạm vi trong tập gate.",
        accent=AMBER,
        fill=PALE_AMBER,
    )

    add_day_block(
        doc,
        "6",
        "25/08",
        "Kiểm thử âm thanh, vùng miền & mô hình chi phí",
        "Đo trải nghiệm thực tế thay vì chỉ kiểm tra chức năng trong môi trường yên tĩnh.",
        [
            "Test yên tĩnh, tiếng TV/người nói xa, giọng trẻ gần mic; đo SNR/ASR/intent theo cùng kịch bản.",
            "Pilot Bắc–Trung–Nam với dữ liệu có đồng ý; tách lỗi âm học, từ địa phương và intent mapping.",
            "Chạy Android vs iOS Native vs Safari PWA để thấy chênh lệch, không trộn kết quả ba nền tảng.",
        ],
        [
            "Tổng hợp chi phí Railway/Cloudflare/TTS-ASR/cache/log và phương án server riêng theo 1.000 lượt, peak load và dữ liệu lưu.",
            "Chỉ đề xuất chuyển server khi benchmark chứng minh lợi ích chi phí/độ trễ/kiểm soát; không migrate gấp trong tuần nếu chưa đủ dữ liệu.",
            "Sửa lỗi P0 phát hiện từ ma trận; đưa lỗi còn lại vào P1/P2 với owner và hạn xử lý.",
        ],
        ["báo cáo noise + vùng miền", "ma trận nền tảng", "cost model/1.000 lượt", "danh sách P0/P1/P2"],
        "Có số liệu accuracy/latency theo điều kiện; 0 lưu/chia sẻ audio trẻ em nếu chưa có đồng ý; quyết định server có cơ sở định lượng.",
        accent=GREEN,
        fill=PALE_GREEN,
    )

    add_day_block(
        doc,
        "7",
        "26/08",
        "Regression, release gate & báo cáo sếp",
        "Đóng tuần bằng một trạng thái sản phẩm có thể ra quyết định, không để kết quả rải rác trong chat/log cá nhân.",
        [
            "Chạy lại regression Android + iOS Native: UI, MAIN, BLE reconnect, HFP route, Speech, mạng, fallback và nội dung.",
            "Đối chiếu gate; chỉ sửa P0 chặn demo/release, không mở thêm UI hoặc refactor ngoài phạm vi.",
            "Đóng branch/tag/build có thể tái tạo; lưu hash, version, thiết bị và firmware test.",
        ],
        [
            "Hoàn thiện báo cáo sếp: đã đạt, chưa đạt, chi phí, rủi ro, phụ thuộc ODM/Apple và quyết định cần duyệt.",
            "Chốt roadmap 14 ngày tiếp theo: App Store/TestFlight, noise/region mở rộng, server, privacy và analytics.",
            "Lập go/no-go: demo nội bộ, TestFlight, hoặc giữ Android baseline trong khi iOS tiếp tục hoàn thiện.",
        ],
        ["release gate", "build/tag", "báo cáo điều hành", "roadmap 14 ngày", "decision request"],
        "Không còn P0 chưa có chủ sở hữu; mọi lỗi còn lại có mức ưu tiên và hạn; sếp có đủ dữ liệu để quyết định go/no-go.",
        accent=BLUE,
        fill=PALE_BLUE,
    )

    add_heading(doc, "3. KPI và tiêu chí nghiệm thu", 1)
    kpi = doc.add_table(rows=1, cols=4)
    configure_table(kpi, [1980, 3720, 1980, 1680])
    set_table_header(kpi, ["Hạng mục", "Chỉ số đo", "Mục tiêu nội bộ", "Gate"])
    kpi_rows = [
        ("Intent chủ lực", "Accuracy theo intent; confusion matrix; false navigation.", "≥95% yên tĩnh; 0 điều hướng nhầm trong bộ P0.", "Ngày 1/5"),
        ("Ngoài phạm vi", "Lần 1 nhắc lại; lần 2 tạm biệt; không nhầm lỗi kỹ thuật.", "100% đúng trên bộ 30 câu OOS kiểm soát.", "Ngày 5"),
        ("Nói liên tục", "End-of-speech → ASR final → intent → first audio.", "P50 ≤2,5s; P95 ≤4,0s hoặc cải thiện ≥20% baseline.", "Ngày 2/7"),
        ("UI", "Screen/flow coverage; lỗi hiển thị; regression luồng chính.", "100% màn hình/luồng đã quyết định; 0 lỗi P0.", "Ngày 2"),
        ("iOS BLE", "MAIN event, duplicate, reconnect.", "20/20 MAIN; 0 duplicate; reconnect đạt ma trận.", "Ngày 4"),
        ("iOS HFP", "currentRoute input/output; record/play; interruption.", "10/10 phiên đúng route hoặc blocker có log.", "Ngày 4/5"),
        ("Vùng miền/noise", "Accuracy, phải nói lại, lỗi theo điều kiện.", "Có baseline tách Bắc–Trung–Nam và gần/xa; không suy diễn khi mẫu nhỏ.", "Ngày 6"),
        ("Chi phí", "ASR/TTS/backend/cache/log trên 1.000 lượt và peak.", "Có so sánh hiện tại vs server riêng + ngưỡng ra quyết định.", "Ngày 6"),
    ]
    for idx, values in enumerate(kpi_rows):
        row = kpi.add_row()
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col, value in enumerate(values):
            set_cell_text(
                row.cells[col],
                value,
                bold=col == 0,
                color=TEXT,
                size=8.9,
                align=WD_ALIGN_PARAGRAPH.CENTER if col == 3 else None,
            )

    add_callout(
        doc,
        "Lưu ý KPI",
        "Các ngưỡng trên là gate thử nghiệm nội bộ, không phải cam kết marketing. Nếu baseline hoặc mẫu thử chưa đủ, "
        "báo cáo phải ghi rõ cỡ mẫu, thiết bị, khoảng cách mic, môi trường và confidence interval/độ không chắc chắn.",
        fill=PALE_AMBER,
        accent=AMBER,
    )

    add_heading(doc, "4. Phụ thuộc, rủi ro và quyết định cần xin", 1)
    risks = doc.add_table(rows=1, cols=4)
    configure_table(risks, [2260, 2500, 2820, 1780])
    set_table_header(risks, ["Phụ thuộc/rủi ro", "Tác động", "Cách xử lý trong kế hoạch", "Hạn quyết định"])
    risk_rows = [
        ("Luồng của chị Huyền chưa chốt", "Có thể làm lại nội dung/state.", "Tách UI/intent nền; chỉ tích hợp flow mới khi có bản chốt/version.", "Ngày 1"),
        ("Thiếu Mac/Xcode/iPhone", "Không chạy/ký iOS Native.", "Thuê/mượn Mac; xác nhận cấu hình và thời gian sử dụng; chuẩn bị backup.", "Ngày 2"),
        ("Apple Developer chưa sẵn sàng", "Không TestFlight/App Store.", "Chọn cá nhân/tổ chức; 2FA; quyền pháp lý; D‑U‑N‑S nếu tổ chức; dự trù 99 USD/năm theo giá Apple hiện hành.", "Ngày 2–3"),
        ("ODM protocol/firmware chưa rõ", "MAIN/reconnect có thể sai.", "Lưu raw packet; dispatch chỉ packet đã biết; gửi issue có video/log.", "Ngày 4"),
        ("Speech native không available", "Nhận dạng gián đoạn theo locale/mạng/OS.", "Kiểm tra availability; giữ Cloudflare Batch fallback; không mở hai recorder.", "Ngày 3–5"),
        ("Dữ liệu giọng trẻ", "Rủi ro riêng tư và kết quả sai lệch.", "Có đồng ý; ẩn danh; retention tối thiểu; không thu thập ngoài mục đích test.", "Trước Ngày 6"),
        ("Scope quá rộng trong 7 ngày", "Dễ mất baseline và trễ báo cáo.", "Ngày 7 chỉ đóng P0; App Store approval và tối ưu sâu chuyển roadmap 14 ngày.", "Xuyên suốt"),
    ]
    for idx, values in enumerate(risk_rows):
        row = risks.add_row()
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col, value in enumerate(values):
            set_cell_text(row.cells[col], value, bold=col == 0, color=TEXT, size=8.85)

    add_heading(doc, "5. Mẫu báo cáo cuối ngày", 1)
    for index, item in enumerate(
        [
            "Đã hoàn thành: tối đa 3 kết quả, kèm commit/build và trạng thái pass/fail.",
            "Số liệu: accuracy, latency P50/P95, số test pass/total, chi phí hoặc blocker.",
            "P0 còn mở: mô tả bước tái hiện, platform/device/firmware và người/đầu vào đang chờ.",
            "Quyết định cần sếp: nêu lựa chọn, tác động thời gian/chi phí/rủi ro và đề xuất cá nhân.",
            "24 giờ tiếp theo: ba việc có đầu ra và gate rõ ràng.",
        ],
        start=1,
    ):
        paragraph = add_body(doc, f"{index}.  {item}")
        paragraph.paragraph_format.left_indent = Inches(0.18)

    add_callout(
        doc,
        "Mẫu câu chốt",
        "Hôm nay [build/commit] đạt [x/y] gate. So với baseline, [chỉ số] thay đổi từ […] thành […]. "
        "P0 còn […]. Tôi đề xuất [go/hold] cho […]. Cần sếp chốt […] trước […].",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "Nguồn chính thức cần đối chiếu trong quá trình triển khai", 2)
    sources = [
        "Apple Speech framework: https://developer.apple.com/documentation/speech",
        "Apple Core Bluetooth background processing: https://developer.apple.com/library/archive/documentation/NetworkingInternetWeb/Conceptual/CoreBluetooth_concepts/CoreBluetoothBackgroundProcessingForIOSApps/PerformingTasksWhileYourAppIsInTheBackground.html",
        "AVAudioSession currentRoute: https://developer.apple.com/documentation/avfaudio/avaudiosession/currentroute",
        "Bluetooth HFP audio option: https://developer.apple.com/documentation/avfaudio/avaudiosession/categoryoptions-swift.struct/allowbluetooth",
        "Apple Developer Program enrollment: https://developer.apple.com/programs/enroll/",
    ]
    for source in sources:
        p = doc.add_paragraph(style="Small Note")
        add_run(p, source, color=MUTED, size=8.2)

    note = doc.add_paragraph(style="Small Note")
    add_run(
        note,
        "Phạm vi tài liệu: kế hoạch công việc cá nhân. Không bao gồm thời lượng phát triển nội dung/nhạc/asset của đồng nghiệp, "
        "sửa firmware ODM hoặc thời gian Apple duyệt. Tài liệu dùng preset compact_reference_guide, Letter 1-inch margins, "
        "Calibri 11 pt và bảng fixed-width để thuận tiện in, báo cáo và cập nhật.",
        color=MUTED,
        size=8.2,
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
