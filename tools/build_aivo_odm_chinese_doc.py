from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

import build_aivo_odm_bilingual_doc as base


OUTPUT = Path(__file__).resolve().parents[1] / "output" / "docs" / "AIVO_V1_ODM_INTEGRATION_REQUIREMENTS_ZH_CN.docx"


def set_chinese_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = True
    section.different_first_page_header_footer = True

    for header_part in [section.header, section.even_page_header, section.first_page_header]:
        header = header_part.paragraphs[0]
        header.clear()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        base.set_run(header.add_run("AIVO V1 | ODM 集成要求"), size=8.5, color=base.MUTED, lang="zh-CN")

    for footer_part in [section.footer, section.even_page_footer, section.first_page_footer]:
        footer = footer_part.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        base.set_run(footer.add_run("AIVO - 技术草案  |  "), size=8, color=base.MUTED, lang="zh-CN")
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)


def add_title_page(doc: Document) -> None:
    for _ in range(4):
        base.add_text(doc, "", after=8, lang="zh-CN")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    base.set_run(p.add_run("AIVO V1"), size=24, bold=True, lang="zh-CN")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    base.set_run(p.add_run("硬件与 APP 协议集成要求"), size=18, bold=True, lang="zh-CN")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    base.set_run(p.add_run("供 ODM 打样、固件开发与联合调试使用"), size=12.5, lang="zh-CN")

    base.add_table(
        doc,
        ["项目", "内容"],
        [
            ["接收方", "新 ODM 制造商"],
            ["范围", "Android 与 iOS 原生应用"],
            ["状态", "技术草案 - 待 ODM 确认"],
            ["版本", "V0.2 - 2026/09/03"],
        ],
        [2600, 7260],
        lang="zh-CN",
        font_size=10,
    )
    base.add_text(
        doc,
        "注意：本文件已提供 APP 方建议通信协议。HFP、BLE 与 OTA 能力仍须由 ODM 通过固件和实物样机确认后方可定版生产。",
        bold=True,
        size=10,
        after=4,
        lang="zh-CN",
    )
    doc.add_page_break()


def audit_docx(path: Path) -> None:
    check = Document(path)
    text = "\n".join(p.text for p in check.paragraphs)
    text += "\n" + "\n".join(
        cell.text for table in check.tables for row in table.rows for cell in row.cells
    )
    required = [
        "硬件与 APP 协议集成要求",
        "7. 本文件已由 APP 方提供的内容",
        "9E3B0001-4A7C-4D6F-8B21-5C17A2D94010",
        "MAIN_LONG",
        "APP PAUSED ACK",
        "OTA 要求",
    ]
    for token in required:
        assert token in text, token
    assert len(check.tables) >= 8


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = "AIVO V1 - ODM 集成要求"
    doc.core_properties.subject = "HFP、BLE 控制、按键协议、OTA 与包装要求"
    doc.core_properties.author = "AIVO"
    doc.core_properties.keywords = "AIVO, ODM, HFP, BLE, OTA, Android, iOS"
    base.style_document(doc)
    set_chinese_header_footer(doc)
    bullet_num = base.create_bullet_numbering(doc)
    add_title_page(doc)
    base.add_chinese(
        doc,
        bullet_num,
        include_part_heading=False,
        include_page_break=False,
    )
    in_odm_requirements = False
    for paragraph in doc.paragraphs:
        if paragraph.text == "6. ODM 必须提供或确认的信息":
            in_odm_requirements = True
            continue
        if paragraph.text == "7. 本文件已由 APP 方提供的内容":
            break
        if in_odm_requirements and paragraph.text.strip():
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.size = Pt(9.8)
    doc.save(OUTPUT)
    audit_docx(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
