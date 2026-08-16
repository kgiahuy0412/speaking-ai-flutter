from pathlib import Path
from zipfile import ZipFile

from docx import Document


path = next(Path("output/docx").glob("H20-*.docx"))
document = Document(path)
text = "\n".join(paragraph.text for paragraph in document.paragraphs)
text += "\n" + "\n".join(
    cell.text for table in document.tables for row in table.rows for cell in row.cells
)
section = document.sections[0]
checks = {
    "title": "H20 在 Android APK 后台运行的可行性简报" in text,
    "conclusion": "结论：可行性高" in text,
    "support": "有条件支持" in text,
    "risk": "Force Stop" in text,
    "recommendation": "建议：批准实施" in text,
}
xml = ZipFile(path).read("word/document.xml").decode("utf-8")

print("file", path.name)
print("bytes", path.stat().st_size)
print("sections", len(document.sections), "tables", len(document.tables))
print("table-rows", len(document.tables[0].rows))
print("page-inches", round(section.page_width.inches, 2), round(section.page_height.inches, 2))
print(
    "margins",
    *[
        round(value.inches, 2)
        for value in (
            section.top_margin,
            section.right_margin,
            section.bottom_margin,
            section.left_margin,
        )
    ],
)
print("checks", checks)
print("header-row", "<w:tblHeader" in xml)
print("font-eastAsia", 'w:eastAsia="Microsoft YaHei"' in xml)
print("replacement-char", "\ufffd" in text)
