from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUTPUT = ROOT / "output" / "docx" / "H20-Android后台运行可行性简报-中文.docx"

FONT = "Microsoft YaHei"
NAVY = "10233F"
BLUE = "2563EB"
GREEN = "0F9D68"
AMBER = "B7791F"
RED = "B42318"
INK = "172033"
MUTED = "5F6B7A"
LINE = "D8E0EA"
LIGHT = "F6F8FB"
PALE_GREEN = "ECFDF5"
PALE_BLUE = "EFF6FF"
WHITE = "FFFFFF"


def set_run_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT)


def configure_style(style, size, bold=False, color=INK, before=0, after=6, line=1.10):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_paragraph_fill(paragraph, fill, border_color=None):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    if border_color:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)
        ppr.append(p_bdr)


def set_paragraph_padding(paragraph, before=7, after=7, left=8, right=8):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.left_indent = Pt(left)
    pf.right_indent = Pt(right)


def make_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    rpr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    rpr.append(color)
    lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def add_bullet(doc, num_id, text, after=3.5):
    paragraph = doc.add_paragraph()
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_el)
    ppr.append(num_pr)
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.10
    set_run_font(paragraph.add_run(text), size=10.5)
    return paragraph


def add_heading(doc, text):
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(text), size=13, bold=True, color=NAVY)
    return paragraph


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)
    set_run_font(run, size=8, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    configure_style(styles["Normal"], 10.5, after=5, line=1.10)
    configure_style(styles["Heading 1"], 16, bold=True, color=NAVY, before=10, after=6)
    configure_style(styles["Heading 2"], 13, bold=True, color=NAVY, before=7, after=4)
    configure_style(styles["Heading 3"], 12, bold=True, color=NAVY, before=6, after=3)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("H20  |  ANDROID APK  |  内部简报"), size=8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.65))
    set_run_font(footer.add_run("2026年8月14日  |  V1范围：BLE控制 + 双向HFP"), size=8, color=MUTED)
    footer.add_run("\t")
    add_page_field(footer)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True
    set_run_font(title.add_run("H20 在 Android APK 后台运行的可行性简报"), size=20, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(7)
    set_run_font(subtitle.add_run("管理层一页简版  |  ODM 已确认架构：BLE 仅用于控制，HFP 用于双向音频"), size=9.5, color=MUTED)

    lead = doc.add_paragraph()
    set_paragraph_padding(lead, before=6, after=6, left=8, right=8)
    add_paragraph_fill(lead, PALE_GREEN, GREEN)
    set_run_font(lead.add_run("结论：可行性高。"), size=11, bold=True, color=GREEN)
    set_run_font(
        lead.add_run(
            " 用户先打开 APP 并启用“H20 后台学习模式”后，即使切换到 Facebook/其他应用或锁屏，孩子仍可按 MAIN 继续学习。实现方式类似 Be/Grab：通过带常驻通知的 Android Foreground Service 维持后台运行。"
        ),
        size=10.5,
    )

    add_heading(doc, "使用状态评估")
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [6900, 2460], indent_dxa=120)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    headers = ("使用场景", "评估")
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(text), size=9.5, bold=True, color=WHITE)
    scenarios = [
        ("已打开 APP 并启用后台学习，然后切换应用或锁屏", "支持", GREEN),
        ("从最近任务中划掉界面，但后台 Service 仍在运行", "有条件支持", AMBER),
        ("未启用后台学习、手机刚重启或用户已 Force Stop", "不保证", RED),
    ]
    for row_index, (scenario, status, color) in enumerate(scenarios, start=1):
        cells = table.add_row().cells
        set_table_geometry(table, [6900, 2460], indent_dxa=120)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 0:
                set_cell_shading(cell, LIGHT)
        p_left = cells[0].paragraphs[0]
        p_left.paragraph_format.space_after = Pt(0)
        set_run_font(p_left.add_run(scenario), size=9.5)
        p_right = cells[1].paragraphs[0]
        p_right.paragraph_format.space_after = Pt(0)
        set_run_font(p_right.add_run(status), size=9.5, bold=True, color=color)

    num_id = make_bullet_numbering(doc)

    add_heading(doc, "V1 后台工作流程")
    add_bullet(doc, num_id, "BLE Control：接收 MAIN，并读取电量、固件和状态。")
    add_bullet(doc, num_id, "双向 HFP：H20 麦克风 → APP/云端；英文音频 → H20 扬声器。")
    add_bullet(doc, num_id, "Foreground Service：在后台保持 BLE、HFP、学习会话、网络及自动重连。")

    add_heading(doc, "APP 侧实施要求")
    requirement = doc.add_paragraph()
    set_paragraph_padding(requirement, before=5, after=5, left=8, right=8)
    add_paragraph_fill(requirement, PALE_BLUE, BLUE)
    set_run_font(
        requirement.add_run(
            "新增 connectedDevice + microphone 类型的 Foreground Service（必要时增加 mediaPlayback）；显示“H20 已就绪”常驻通知；将 BLE/HFP 从 MainActivity 生命周期迁移到 Service；处理音频焦点、重连及无界面时的 MAIN 状态机。"
        ),
        size=10,
    )

    add_heading(doc, "需要提前说明的限制")
    add_bullet(doc, num_id, "Facebook 播放有声视频时，Android 可能降低/暂停其音量，或改变 HFP 音频路由；必须使用 H20 实机验证。")
    add_bullet(doc, num_id, "Force Stop 或 Android“正在运行的应用”中的 Stop 会终止 Service，之后必须重新打开 APP。")
    add_bullet(doc, num_id, "云端功能需要网络；离线 Diagnostic APK 仍应本地验证 BLE、HFP、MAIN、麦克风和扬声器。")
    add_bullet(doc, num_id, "样机到达越南前无需 ODM 增加新功能；当前改动范围属于 APP 端。")

    recommendation = doc.add_paragraph()
    set_paragraph_padding(recommendation, before=6, after=6, left=8, right=8)
    add_paragraph_fill(recommendation, NAVY)
    set_run_font(recommendation.add_run("建议：批准实施。"), size=10.5, bold=True, color=WHITE)
    set_run_font(
        recommendation.add_run(
            " 第一阶段验证离线后台链路（MAIN + HFP 录放音）；第二阶段接入 AI 后端；第三阶段测试锁屏、Facebook 及不同品牌 Android 手机。"
        ),
        size=10,
        color=WHITE,
    )

    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(3)
    source.paragraph_format.space_after = Pt(0)
    set_run_font(
        source.add_run(
            "技术依据：Android Developers - BLE 后台通信及 Foreground Service 启动限制；Google Play - Foreground Service 申报要求。"
        ),
        size=7.8,
        color=MUTED,
    )

    doc.core_properties.title = "H20 在 Android APK 后台运行的可行性简报"
    doc.core_properties.subject = "H20 V1 - BLE Control + 双向 HFP"
    doc.core_properties.author = "AI Speaking Project"
    doc.core_properties.keywords = "H20, Android, Foreground Service, BLE, HFP"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print("created-docx")


if __name__ == "__main__":
    build()
