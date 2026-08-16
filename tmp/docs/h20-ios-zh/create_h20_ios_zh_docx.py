import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[3]
HELPER_DIR = ROOT / "tmp" / "docs" / "h20-android-zh"
sys.path.insert(0, str(HELPER_DIR))

from create_h20_android_zh_docx import (  # noqa: E402
    AMBER,
    BLUE,
    GREEN,
    INK,
    LIGHT,
    LINE,
    MUTED,
    NAVY,
    RED,
    WHITE,
    add_bullet,
    add_heading,
    add_page_field,
    add_paragraph_fill,
    configure_style,
    make_bullet_numbering,
    mark_header_row,
    set_cell_shading,
    set_cell_margins,
    set_paragraph_padding,
    set_run_font,
    set_table_borders,
    set_table_geometry,
)


OUTPUT = ROOT / "output" / "docx" / "H20-iOS原生后台运行可行性简报-中文.docx"
PALE_AMBER = "FFF8E6"
PALE_BLUE = "EFF6FF"

# Named override for the requested one-page CJK executive brief. The underlying
# design system remains standard_business_brief + memo_masthead.
ONE_PAGE_CJK = {
    "top": 0.55,
    "bottom": 0.50,
    "left": 0.82,
    "right": 0.82,
    "header": 0.32,
    "footer": 0.32,
    "body_size": 10.0,
    "body_after": 4,
    "body_line": 1.05,
}


def add_status_table(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [6900, 2460], indent_dxa=120)
    set_table_borders(table, color=LINE)
    mark_header_row(table.rows[0])

    for idx, text in enumerate(("使用场景", "评估")):
        cell = table.rows[0].cells[idx]
        set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(text), size=9.5, bold=True, color=WHITE)

    scenarios = [
        ("已启动后台学习与 HFP 会话后锁屏", "有条件支持", AMBER),
        ("切换 Facebook/其他应用（无声）", "有条件支持", AMBER),
        ("Facebook 播放有声视频并同时使用 HFP", "不保证", RED),
        ("未提前启动，或用户已 Force Quit", "不支持", RED),
    ]
    for row_index, (scenario, status, color) in enumerate(scenarios, start=1):
        cells = table.add_row().cells
        set_table_geometry(table, [6900, 2460], indent_dxa=120)
        for cell in cells:
            set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 0:
                set_cell_shading(cell, LIGHT)
        left = cells[0].paragraphs[0]
        left.paragraph_format.space_after = Pt(0)
        set_run_font(left.add_run(scenario), size=9.3)
        right = cells[1].paragraphs[0]
        right.paragraph_format.space_after = Pt(0)
        right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(right.add_run(status), size=9.3, bold=True, color=color)

    # set_table_geometry applies the shared helper's legacy 100-DXA vertical
    # margin; normalize every cell to the preset's exact 80/120/80/120 tokens.
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=80, start=120, bottom=80, end=120)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(ONE_PAGE_CJK["top"])
    section.bottom_margin = Inches(ONE_PAGE_CJK["bottom"])
    section.left_margin = Inches(ONE_PAGE_CJK["left"])
    section.right_margin = Inches(ONE_PAGE_CJK["right"])
    section.header_distance = Inches(ONE_PAGE_CJK["header"])
    section.footer_distance = Inches(ONE_PAGE_CJK["footer"])

    styles = doc.styles
    configure_style(
        styles["Normal"],
        ONE_PAGE_CJK["body_size"],
        after=ONE_PAGE_CJK["body_after"],
        line=ONE_PAGE_CJK["body_line"],
    )
    configure_style(styles["Heading 1"], 15.5, bold=True, color=NAVY, before=8, after=5)
    configure_style(styles["Heading 2"], 12.5, bold=True, color=NAVY, before=5, after=2.5)
    configure_style(styles["Heading 3"], 11.5, bold=True, color=NAVY, before=5, after=2.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("H20  |  iOS NATIVE  |  内部简报"), size=8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.65))
    set_run_font(footer.add_run("2026年8月14日  |  V1范围：BLE控制 + 双向HFP"), size=8, color=MUTED)
    footer.add_run("\t")
    add_page_field(footer)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True
    set_run_font(title.add_run("H20 在 iOS 原生 APP 后台运行的可行性简报"), size=20, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(7)
    set_run_font(
        subtitle.add_run("管理层一页简版  |  ODM 已确认架构：BLE 仅用于控制，HFP 用于双向音频"),
        size=9.5,
        color=MUTED,
    )

    lead = doc.add_paragraph()
    set_paragraph_padding(lead, before=4, after=4, left=7, right=7)
    add_paragraph_fill(lead, PALE_AMBER, AMBER)
    set_run_font(lead.add_run("结论：有条件可行，后台可靠性低于 Android。"), size=10.8, bold=True, color=AMBER)
    set_run_font(
        lead.add_run(
            " 用户先打开 APP、授权并启动“H20 后台学习”后，锁屏或切换应用时可继续 HFP；Facebook 有声播放、系统中断、挂起或 Force Quit 时不保证。"
        ),
        size=10.2,
        color=INK,
    )

    add_heading(doc, "使用状态评估")
    add_status_table(doc)

    num_id = make_bullet_numbering(doc)

    add_heading(doc, "V1 后台工作流程")
    add_bullet(doc, num_id, "CoreBluetooth 后台模式接收 MAIN，并读取电量、固件和状态。", after=2)
    add_bullet(doc, num_id, "AVAudioSession 双向 HFP：H20 麦克风 → APP/云端；英文音频 → H20 扬声器。", after=2)
    add_bullet(doc, num_id, "APP 处理网络、音频中断、路由变化、BLE 重连及状态恢复。", after=2)

    add_heading(doc, "iOS 原生 APP 实施要求")
    requirement = doc.add_paragraph()
    set_paragraph_padding(requirement, before=4, after=4, left=7, right=7)
    add_paragraph_fill(requirement, PALE_BLUE, BLUE)
    set_run_font(
        requirement.add_run(
            "Swift 原生采用 CoreBluetooth bluetooth-central + 状态恢复；AVAudioSession 使用 playAndRecord + voiceChat + allowBluetoothHFP，并启用 audio 后台模式。用户须先授权并主动开启后台学习。"
        ),
        size=9.8,
    )

    add_heading(doc, "需要提前说明的限制")
    add_bullet(doc, num_id, "APP 被挂起时，BLE MAIN 可能只短暂唤醒；未启动会话不能保证整轮云端/HFP。", after=2)
    add_bullet(doc, num_id, "Facebook 有声播放可能抢占/改变 HFP；手机与 H20 扬声器不能保证独立并行。", after=2)
    add_bullet(doc, num_id, "用户 Force Quit 后，iOS 不允许 APP 继续运行，必须重新打开。", after=2)
    add_bullet(doc, num_id, "后台音频/BLE 必须服务真实功能；App Store 不允许用静音音频保活。", after=2)

    recommendation = doc.add_paragraph()
    set_paragraph_padding(recommendation, before=4, after=4, left=7, right=7)
    add_paragraph_fill(recommendation, NAVY)
    set_run_font(recommendation.add_run("建议：批准技术验证，但不承诺与 Android 相同的可靠性。"), size=10.2, bold=True, color=WHITE)
    set_run_font(
        recommendation.add_run(
            " 先验证锁屏 + MAIN + HFP，再测试 Facebook 中断恢复及 App Store/多机型。若必须做到“无需打开 APP 也长期独立学习”，应评估 H20 自带 Wi-Fi/云端方案。"
        ),
        size=9.7,
        color=WHITE,
    )

    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(3)
    source.paragraph_format.space_after = Pt(0)
    set_run_font(
        source.add_run(
            "技术依据：Apple Developer - Core Bluetooth 后台处理、AVAudioSession playAndRecord/HFP；App Store Review Guidelines 2.5.4。"
        ),
        size=7.8,
        color=MUTED,
    )

    doc.core_properties.title = "H20 在 iOS 原生 APP 后台运行的可行性简报"
    doc.core_properties.subject = "H20 V1 - BLE Control + 双向 HFP"
    doc.core_properties.author = "AI Speaking Project"
    doc.core_properties.keywords = "H20, iOS, CoreBluetooth, AVAudioSession, BLE, HFP"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print("created-docx")


if __name__ == "__main__":
    build()
