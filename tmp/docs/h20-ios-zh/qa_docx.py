from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


PATH = Path("output/docx/H20-iOS原生后台运行可行性简报-中文.docx")
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


document = Document(PATH)
text = "\n".join(paragraph.text for paragraph in document.paragraphs)
text += "\n" + "\n".join(
    cell.text for table in document.tables for row in table.rows for cell in row.cells
)
section = document.sections[0]

with ZipFile(PATH) as archive:
    document_xml = archive.read("word/document.xml")
    styles_xml = archive.read("word/styles.xml").decode("utf-8")
    numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
    header_xml = archive.read("word/header1.xml").decode("utf-8")
    footer_xml = archive.read("word/footer1.xml").decode("utf-8")

root = etree.fromstring(document_xml)
table = root.find(".//w:tbl", NS)
table_width = table.find("./w:tblPr/w:tblW", NS).get(f"{{{W_NS}}}w")
table_indent = table.find("./w:tblPr/w:tblInd", NS).get(f"{{{W_NS}}}w")
grid_widths = [
    int(node.get(f"{{{W_NS}}}w")) for node in table.findall("./w:tblGrid/w:gridCol", NS)
]
row_widths = [
    [
        int(node.get(f"{{{W_NS}}}w"))
        for node in row.findall("./w:tc/w:tcPr/w:tcW", NS)
    ]
    for row in table.findall("./w:tr", NS)
]
cell_margins = [
    {
        edge.tag.rsplit("}", 1)[-1]: int(edge.get(f"{{{W_NS}}}w"))
        for edge in cell.findall("./w:tcPr/w:tcMar/*", NS)
    }
    for cell in table.findall(".//w:tc", NS)
]

checks = {
    "title": "H20 在 iOS 原生 APP 后台运行的可行性简报" in text,
    "conditional_conclusion": "结论：有条件可行" in text,
    "facebook_audio_risk": "Facebook 有声播放" in text or "Facebook 播放有声视频" in text,
    "force_quit": "Force Quit" in text,
    "app_store": "App Store" in text,
    "recommendation": "建议：批准技术验证" in text,
    "source": "Core Bluetooth 后台处理" in text,
    "one_section": len(document.sections) == 1,
    "one_table": len(document.tables) == 1,
    "five_table_rows": len(document.tables[0].rows) == 5,
    "letter_page": round(section.page_width.inches, 2) == 8.5
    and round(section.page_height.inches, 2) == 11.0,
    "margins": [
        round(section.top_margin.inches, 2),
        round(section.right_margin.inches, 2),
        round(section.bottom_margin.inches, 2),
        round(section.left_margin.inches, 2),
    ]
    == [0.55, 0.82, 0.5, 0.82],
    "style_tokens": round(document.styles["Normal"].font.size.pt, 1) == 10.0
    and round(document.styles["Heading 2"].font.size.pt, 1) == 12.5,
    "header_row": b"<w:tblHeader" in document_xml,
    "east_asia_font": 'w:eastAsia="Microsoft YaHei"' in styles_xml
    and b'w:eastAsia="Microsoft YaHei"' in document_xml,
    "real_bullets": 'w:numFmt w:val="bullet"' in numbering_xml,
    "no_fake_bullets": not any(
        paragraph.text.startswith(("•", "- ")) for paragraph in document.paragraphs
    ),
    "table_geometry": table_width == "9360"
    and table_indent == "120"
    and grid_widths == [6900, 2460]
    and all(widths == [6900, 2460] for widths in row_widths),
    "cell_margins": all(
        margins == {"top": 80, "start": 120, "bottom": 80, "end": 120}
        for margins in cell_margins
    ),
    "header": "H20  |  iOS NATIVE  |  内部简报" in header_xml,
    "footer": "V1范围：BLE控制 + 双向HFP" in footer_xml,
    "no_replacement_char": "\ufffd" not in text,
}

print("file", PATH.name)
print("bytes", PATH.stat().st_size)
print("paragraphs", len(document.paragraphs), "tables", len(document.tables))
print("table-rows", len(document.tables[0].rows))
print("page-inches", round(section.page_width.inches, 2), round(section.page_height.inches, 2))
print(
    "margins",
    round(section.top_margin.inches, 2),
    round(section.right_margin.inches, 2),
    round(section.bottom_margin.inches, 2),
    round(section.left_margin.inches, 2),
)
print("table-width", table_width, "indent", table_indent, "grid", grid_widths)
print("first-cell-margins", cell_margins[0])
print("checks", checks)

failed = [name for name, passed in checks.items() if not passed]
if failed:
    raise SystemExit(f"QA failed: {failed}")
