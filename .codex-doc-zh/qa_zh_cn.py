from __future__ import annotations

import hashlib
import re
import sys
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

from docx import Document


VIETNAMESE_RE = re.compile(
    r"[ăâđêôơưĂÂĐÊÔƠƯàáảãạằắẳẵặầấẩẫậèéẻẽẹềếểễệ"
    r"ìíỉĩịòóỏõọồốổỗộờớởỡợùúủũụừứửữựỳýỷỹỵ]",
    re.IGNORECASE,
)


def sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def relationship_set(data: bytes) -> set[tuple[str, str, str, str]]:
    root = ET.fromstring(data)
    return {
        (
            rel.attrib.get("Id", ""),
            rel.attrib.get("Type", ""),
            rel.attrib.get("Target", ""),
            rel.attrib.get("TargetMode", ""),
        )
        for rel in root
    }


def visible_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    for section in document.sections:
        for container in (section.header, section.footer):
            values.extend(paragraph.text for paragraph in container.paragraphs)
    return "\n".join(values)


def main() -> None:
    source_path = Path(sys.argv[1])
    final_path = Path(sys.argv[2])
    source = Document(source_path)
    final = Document(final_path)

    assert len(source.paragraphs) == len(final.paragraphs) == 244
    assert len(source.tables) == len(final.tables) == 1
    assert len(source.sections) == len(final.sections) == 1
    assert len(source.inline_shapes) == len(final.inline_shapes) == 1
    assert len(source.tables[0].rows) == len(final.tables[0].rows) == 5
    assert all(len(row.cells) == 3 for row in final.tables[0].rows)

    text = visible_text(final)
    leftovers = sorted(set(VIETNAMESE_RE.findall(text)))
    assert not leftovers, f"Vietnamese characters remain: {leftovers}"

    required = [
        "V1版本实施计划",
        "忠实翻译儿童原话含义",
        "OpenAI Realtime成本规划",
        "V1整体验收标准",
        "V1最终交付资料",
        "发言时长",
        "页码 ",
    ]
    for marker in required:
        assert marker in text, f"Missing marker: {marker}"

    with ZipFile(source_path) as src_zip, ZipFile(final_path) as out_zip:
        assert set(src_zip.namelist()) == set(out_zip.namelist())
        source_media = {
            name: sha256(src_zip.read(name))
            for name in src_zip.namelist()
            if name.startswith("word/media/")
        }
        final_media = {
            name: sha256(out_zip.read(name))
            for name in out_zip.namelist()
            if name.startswith("word/media/")
        }
        assert source_media == final_media

        source_rels = {
            name: relationship_set(src_zip.read(name))
            for name in src_zip.namelist()
            if name.endswith(".rels")
        }
        final_rels = {
            name: relationship_set(out_zip.read(name))
            for name in out_zip.namelist()
            if name.endswith(".rels")
        }
        assert source_rels == final_rels

        source_page_fields = src_zip.read("word/footer1.xml").count(b"PAGE")
        final_page_fields = out_zip.read("word/footer1.xml").count(b"PAGE")
        assert source_page_fields == final_page_fields == 1

    print("QA_OK")
    print(f"output={final_path}")
    print(f"bytes={final_path.stat().st_size}")
    print(f"paragraphs={len(final.paragraphs)} tables={len(final.tables)} images={len(final.inline_shapes)}")
    print("package_parts_and_relationships_preserved=true")
    print("vietnamese_diacritics_remaining=0")


if __name__ == "__main__":
    main()
