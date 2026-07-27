from __future__ import annotations

import json
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document


def main() -> None:
    source = Path(sys.argv[1])
    document = Document(source)

    print("MULTI-RUN PARAGRAPHS")
    for index, paragraph in enumerate(document.paragraphs):
        runs = [run for run in paragraph.runs if run.text]
        if len(runs) <= 1:
            continue
        print(f"{index:03d} style={paragraph.style.name!r} text={paragraph.text!r}")
        for run_index, run in enumerate(runs):
            size = run.font.size.pt if run.font.size else None
            print(
                f"  r{run_index}: {run.text!r} bold={run.bold} "
                f"italic={run.italic} size={size}"
            )

    print("\nDOCUMENT INVENTORY")
    print(f"paragraphs={len(document.paragraphs)}")
    print(f"tables={len(document.tables)}")
    print(f"sections={len(document.sections)}")
    print(f"inline_shapes={len(document.inline_shapes)}")

    for section_index, section in enumerate(document.sections):
        for kind, container in (
            ("header", section.header),
            ("first_page_header", section.first_page_header),
            ("even_page_header", section.even_page_header),
            ("footer", section.footer),
            ("first_page_footer", section.first_page_footer),
            ("even_page_footer", section.even_page_footer),
        ):
            texts = [p.text for p in container.paragraphs if p.text.strip()]
            if texts:
                print(f"section={section_index} {kind}={json.dumps(texts, ensure_ascii=False)}")

    with ZipFile(source) as archive:
        print(f"package_parts={len(archive.namelist())}")
        for name in archive.namelist():
            if name.startswith("word/") and (
                name.endswith(".xml") or name.endswith(".rels")
            ):
                print(name)


if __name__ == "__main__":
    main()
